# -*- coding: utf-8 -*-

import os
import shutil
import logging
import tempfile
import csv
import re
import win32com.client as win32
import pythoncom
import win32print
import subprocess
import struct
import time
import xml.etree.ElementTree as ET
import sys
import functools
import tkinter as tk
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook
from PIL import Image, ImageEnhance
from datetime import datetime
from tkinter import messagebox
from win32com.client import constants

# Настройка логирования
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Декоратор для автоматического перестроения кэша win32com при ошибках CLSIDToClassMap
def rebuild_win32com_cache_on_error(func):
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except AttributeError as e:
            if "CLSIDToClassMap" in str(e):
                logging.warning("Обнаружена ошибка кэша win32com. Автоматическое перестроение кэша...")
                try:
                    # Перестраиваем кэш win32com
                    import win32com.client
                    win32com.client.gencache.is_readonly = False
                    win32com.client.gencache.Rebuild()
                    logging.info("Кэш win32com успешно перестроен. Повторная попытка выполнения функции...")
                    
                    # Повторяем вызов функции после перестроения кэша
                    return func(*args, **kwargs)
                except Exception as rebuild_error:
                    logging.error(f"Ошибка при перестроении кэша win32com: {rebuild_error}")
                    logging.error("Не удалось автоматически исправить ошибку кэша win32com. Завершение программы.")
                    sys.exit(1)
            # Если это другая AttributeError, пробрасываем её дальше
            raise
    return wrapper

# Константы для управления режимом дуплекса принтера
DM_DUPLEX = 0x00001000
DMDUP_SIMPLEX = 1      # Односторонняя печать
DMDUP_VERTICAL = 2     # Двусторонняя по длинной стороне
DMDUP_HORIZONTAL = 3   # Двусторонняя по короткой стороне

# Определяем порядок сортировки локусов (в верхнем регистре)
LOCUS_ORDER = [
    "D3S1358", "VWA", "D16S539", "CSF1PO", "TPOX", "YINDEL", "AMEL", 
    "D8S1179", "D21S11", "D18S51", "DYS391", "D2S441", "D19S433", 
    "TH01", "FGA", "D22S1045", "D5S818", "D13S317", "D7S820", "SE33", 
    "D10S1248", "D1S1656", "D12S391", "D2S1338", "SRY", "D6S1043"
]

# Функция для отображения сообщений об ошибках
def show_error_message(message):
    root = tk.Tk()
    root.withdraw()  # Скрываем главное окно Tkinter
    messagebox.showerror("Ошибка", message)
    root.destroy()

def get_default_printer_name(word_app=None):
    """Получает имя принтера по умолчанию через Word или напрямую."""
    need_cleanup = False
    try:
        if word_app is None:
            need_cleanup = True
            pythoncom.CoInitialize()
            word = win32.gencache.EnsureDispatch('Word.Application')
            word.Visible = False
        else:
            word = word_app

        printer_name = word.ActivePrinter

        if need_cleanup:
            word.Quit()
            pythoncom.CoUninitialize()

        logging.info(f"Определен принтер по умолчанию: {printer_name}")
        return printer_name
    except Exception as e:
        logging.error(f"Ошибка при определении принтера: {e}")
        if need_cleanup:
            try:
                word.Quit()
                pythoncom.CoUninitialize()
            except:
                pass
        return ""

def is_pantum_printer(printer_name):
    """Определяет, является ли принтер Pantum (по имени)."""
    if not printer_name:
        return False
    printer_name_lower = printer_name.lower()
    is_pantum = "pantum" in printer_name_lower
    is_samsung = "samsung" in printer_name_lower

    if is_pantum:
        logging.info(f"Обнаружен принтер Pantum: {printer_name}")
        return True
    elif is_samsung:
        logging.info(f"Обнаружен принтер Samsung: {printer_name}")
        return False
    else:
        logging.warning(f"Неизвестный тип принтера: {printer_name}. Используем логику Samsung по умолчанию.")
        return False

def set_printer_duplex_mode(printer_name, duplex_mode):
    """
    Устанавливает режим дуплекса принтера через win32print.
    duplex_mode: 1 = односторонняя, 2 = длинная сторона, 3 = короткая сторона.
    Возвращает (успех, оригинальный_режим).
    """
    try:
        printer_handle = win32print.OpenPrinter(printer_name)
        try:
            printer_info = win32print.GetPrinter(printer_handle, 2)
            devmode = printer_info["pDevMode"]
            if devmode:
                original_duplex = getattr(devmode, 'Duplex', None)
                devmode.Duplex = duplex_mode
                devmode.Fields |= DM_DUPLEX
                win32print.SetPrinter(printer_handle, 2, printer_info, 0)
                logging.info(f"Установлен режим дуплекса: {duplex_mode} (1=одност., 2=длинная, 3=короткая)")
                return True, original_duplex
            else:
                logging.warning("Не удалось получить DEVMODE принтера")
                return False, None
        finally:
            win32print.ClosePrinter(printer_handle)
    except Exception as e:
        logging.error(f"Ошибка при установке режима дуплекса через win32print: {e}")
        return False, None

def set_printer_duplex_powershell(printer_name, enable_duplex):
    """Управление принтером через PowerShell (запасной метод)."""
    try:
        duplex_setting = "TwoSidedLongEdge" if enable_duplex else "OneSided"
        ps_command = f'Set-PrintConfiguration -PrinterName "{printer_name}" -DuplexingMode {duplex_setting}'
        result = subprocess.run(
            ["powershell", "-Command", ps_command],
            capture_output=True,
            text=True,
            shell=True
        )
        if result.returncode == 0:
            logging.info(f"Настройка дуплекса изменена через PowerShell: {duplex_setting}")
            return True
        else:
            logging.warning(f"PowerShell предупреждение: {result.stderr}")
            return True  # Не считаем критическим, т.к. может не хватать прав
    except Exception as e:
        logging.error(f"Ошибка при выполнении PowerShell: {e}")
        return False

def restore_printer_duplex_mode(printer_name, original_duplex):
    """Восстанавливает оригинальный режим дуплекса принтера."""
    if original_duplex is not None:
        try:
            set_printer_duplex_mode(printer_name, original_duplex)
            logging.info(f"Восстановлен оригинальный режим дуплекса: {original_duplex}")
        except Exception as e:
            logging.warning(f"Не удалось восстановить оригинальный режим дуплекса: {e}")
            # Пробуем PowerShell как запасной вариант (упрощённо)
            try:
                set_printer_duplex_powershell(printer_name, original_duplex == DMDUP_VERTICAL)
            except:
                pass

def show_fio_warning(card_data):
    """Показывает предупреждение о проверке ФИО для определенных шаблонов 
    Возвращает кортеж: (было_ли_показано_предупреждение, нажал_ли_пользователь_OK)"""
    
    target_templates = [
        "СВО_Молов_образец_родственники", "СВО_Ростов_образец_родственники",
        "СВО_Молов_образец_прямая идентификация", "СВО_Ростов_образец_прямая идентификация",
        "СВО_Молов_образец_родственники_нет результата_RT", "СВО_Молов_образец_родственники_нет результата_форез",
        "СВО_Ростов_образец_прямая идентификация_нет результата_RT", "СВО_Ростов_образец_прямая идентификация_нет результата_форез",
        "СВО_Ростов_образец_родственники_нет результата_RT", "СВО_Ростов_образец_родственники_нет результата_форез",
        "ЭКЦ", "ЭКЦ_нет результата", 
        "СВО_кость", "СВО_кость_нет результата"
    ]
    
    current_template = card_data.get("1")
    is_ekc_obrazec = current_template and current_template.startswith("ЭКЦ_образец_")
    is_tobolsk_obrazec = current_template and current_template.startswith("Тобольск_образец_")
    
    if current_template in target_templates or is_ekc_obrazec or is_tobolsk_obrazec:
        logging.info(f"Обнаружен целевой шаблон: '{current_template}'")
        
        # Определяем базовое сообщение в зависимости от типа шаблона
        if current_template in ["ЭКЦ", "ЭКЦ_нет результата"] or is_ekc_obrazec or is_tobolsk_obrazec:
            # Для ЭКЦ, ЭКЦ_образец_ и Тобольск_образец_
            base_message = f"НОМЕР ЗАКЛЮЧЕНИЯ - {card_data.get('НОМ', 'N/A')}\n\nПроверь ФИО!"
            need_check_ind = False  # ИНД не проверяем
        
        elif current_template == "СВО_Молов_образец_прямая идентификация":
            # Для прямой идентификации
            base_message = f"НОМЕР ЗАКЛЮЧЕНИЯ - {card_data.get('НОМ', 'N/A')}\n\nПрямая идентификация!"
            need_check_ind = True   # ИНД нужно проверить
        
        elif current_template in [
            "СВО_Ростов_образец_прямая идентификация",
            "СВО_Ростов_образец_прямая идентификация_нет результата_RT", "СВО_Ростов_образец_прямая идентификация_нет результата_форез",
            ]:
            # Для прямой идентификации
            base_message = f"НОМЕР ЗАКЛЮЧЕНИЯ - {card_data.get('НОМ', 'N/A')}\n\nПрямая идентификация!"
            need_check_ind = False   # ИНД не проверяем

        elif current_template in ["СВО_Ростов_образец_родственники_нет результата_RT", "СВО_Ростов_образец_родственники_нет результата_форез"]:
            # Для родственников по Ростову
            base_message = f"НОМЕР ЗАКЛЮЧЕНИЯ - {card_data.get('НОМ', 'N/A')}"
            need_check_ind = False   # ИНД не проверяем

        elif current_template in ["СВО_кость", "СВО_кость_нет результата"]:
            # Для шаблонов с костями
            base_message = f"НОМЕР ЗАКЛЮЧЕНИЯ - {card_data.get('НОМ', 'N/A')}\n\nНомер кости - {card_data.get('ТР', 'N/A')}\n\nПанель - {card_data.get('ПАН', 'N/A')}"
            need_check_ind = False  # ИНД не проверяем
        
        else:
            # Остальные шаблоны из списка (родственники и т.п.)
            ind_value = card_data.get("ИНД")
            if ind_value is not None and str(ind_value).strip():
                base_message = f"НОМЕР ЗАКЛЮЧЕНИЯ - {card_data.get('НОМ', 'N/A')}"
                logging.info(f"Ячейка 'ИНД' содержит значение '{ind_value}'.")
            else:
                base_message = f"НОМЕР ЗАКЛЮЧЕНИЯ - {card_data.get('НОМ', 'N/A')}\n\nЯчейка индекса пуста!"
                logging.info("Ячейка 'ИНД' пуста или отсутствует.")
            need_check_ind = False  # ИНД уже учтена
        
        # Если для данного типа требуется дополнительная проверка ИНД (только для прямой идентификации)
        if need_check_ind:
            ind_value = card_data.get("ИНД")
            if ind_value is None or not str(ind_value).strip():
                logging.info("Ячейка 'ИНД' пуста или отсутствует. Добавляем предупреждение.")
                base_message += "\n\nЯчейка индекса пуста!"
            else:
                logging.info(f"Ячейка 'ИНД' содержит значение '{ind_value}'.")
        
        warning_text = base_message

        root = tk.Tk()
        root.withdraw()
        message_window = tk.Toplevel(root)
        message_window.title("Внимание!")
        message_window.geometry("500x250")  # Увеличили высоту для отображения номера
        message_window.protocol("WM_DELETE_WINDOW", lambda: (message_window.destroy(), root.quit()))
        message_window.bind('<Escape>', lambda e: (message_window.destroy(), root.quit()))
        message_window.bind('<Return>', lambda e: (message_window.destroy(), root.quit()))
        
        result = [False]  # Используем список для изменения вложенной переменной

        label = tk.Label(message_window, text=warning_text, font=("Arial", 16, "bold"), fg="red")
        label.pack(expand=True, fill="both", padx=20, pady=20)

        ok_button = tk.Button(
            message_window, 
            text="OK (Enter)", 
            command=lambda: (result.__setitem__(0, True), message_window.destroy(), root.quit()),
            font=("Arial", 12)
        )
        ok_button.pack(pady=10)

        message_window.after(100, lambda: (
            ok_button.focus_force(),
            message_window.attributes('-topmost', True)
        ))
        message_window.update_idletasks()
        width = message_window.winfo_width()
        height = message_window.winfo_height()
        x = (message_window.winfo_screenwidth() // 2) - (width // 2)
        y = (message_window.winfo_screenheight() // 2) - (height // 2)
        message_window.geometry(f'+{x}+{y}')

        root.mainloop()
        
        # Всегда возвращаем True для целевых шаблонов, указывая что предупреждение было показано
        return (True, result[0])  # Предупреждение показывалось, возвращаем результат нажатия OK
    
    logging.info(f"Шаблон '{card_data.get('1')}' не требует показа предупреждения.")
    return (False, False)  # Предупреждение не показывалось, пользователь не нажимал OK

def show_final_reminder():
    """Показывает финальное напоминание о необходимых действиях"""
    root = None
    message_window = None
    
    try:
        root = tk.Tk()
        root.withdraw()
        
        message_window = tk.Toplevel(root)
        message_window.title("ВАЖНО!")
        message_window.geometry("600x300")
        message_window.protocol("WM_DELETE_WINDOW", lambda: (message_window.destroy(), root.quit()))
        
        message_window.bind('<Escape>', lambda e: (message_window.destroy(), root.quit()))
        message_window.bind('<Return>', lambda e: (message_window.destroy(), root.quit()))
        
        label = tk.Label(
            message_window, 
            text="""1. ВНЕСИ ДАННЫЕ В 1С
    2. ЗАПОЛНИ ЖУРНАЛ
    3. ЗАПОЛНИ ПОДПИСКУ
    4. СДАЙ ОБЪЕКТЫ!""", 
            font=("Arial", 20, "bold"), 
            fg="red"
        )
        label.pack(expand=True, fill="both", padx=20, pady=50)
        
        ok_button = tk.Button(
            message_window, 
            text="OK (Enter)", 
            command=lambda: (message_window.destroy(), root.quit()),
            font=("Arial", 14)
        )
        ok_button.pack(pady=10)
        
        message_window.after(100, lambda: (
            ok_button.focus_force(),
            message_window.attributes('-topmost', True)
        ))
        
        message_window.update_idletasks()
        width = message_window.winfo_width()
        height = message_window.winfo_height()
        x = (message_window.winfo_screenwidth() // 2) - (width // 2)
        y = (message_window.winfo_screenheight() // 2) - (height // 2)
        message_window.geometry(f'+{x}+{y}')
        
        root.mainloop()
    except Exception as e:
        logging.error(f"Ошибка при создании окна напоминания: {e}")
        raise
    finally:
        # Очистка оконных ресурсов
        if message_window is not None:
            try:
                message_window.destroy()
            except Exception as e:
                logging.debug(f"Ошибка при закрытии message_window: {e}")
        
        if root is not None:
            try:
                root.quit()
                root.destroy()
            except Exception as e:
                logging.debug(f"Ошибка при закрытии root: {e}")

# Функция для сортировки данных по локусам
def sort_data_by_locus(data):
    def get_locus_index(locus):
        locus_upper = locus.strip().upper()
        try:
            return LOCUS_ORDER.index(locus_upper)
        except ValueError:
            return len(LOCUS_ORDER) + ord(locus_upper[0])
    return sorted(data, key=lambda x: get_locus_index(x[3]))

# Функция для проверки обязательных полей
def check_card_data(card_data):
    required_fields = ["НОМ", "2", "3", "4", "7", "ДО", "МО"]
    missing_fields = [field for field in required_fields if field not in card_data or not card_data[field]]
    if missing_fields:
        raise ValueError(f"В карте отсутствуют значения для следующих полей: {missing_fields}")

# Функция для чтения данных из Карты (Excel)
def read_card(file_path):
    data = {}
    wb = load_workbook(file_path)
    sheet = wb.active
    for row in sheet.iter_rows(min_row=2, values_only=True):
        code = str(row[0])  # Преобразуем код в строку
        value = row[2]
        data[code] = value
    return data

# Функция для копирования шаблона Word
def copy_template(card_data):
    template_folder = Path(card_data["3"])  # Путь к папке с шаблонами
    template_name = card_data["1"]  # Название шаблона из строки с кодом 1
    for ext in [".doc", ".docx"]:
        template_path = template_folder / f"{template_name}{ext}"
        if template_path.exists():
            output_path = Path(card_data["2"]) / f"{card_data['НОМ']}-26{ext}"
            shutil.copy(template_path, output_path)
            return output_path
    raise FileNotFoundError(f"Файл шаблона '{template_name}' не найден в папке {template_folder}.")

def clean_text_from_none_and_spaces(text, max_allowed_spaces=3):
    """
    Универсальная функция для очистки текста от меток с None и лишних пробелов.
    Заменяет последовательности пробелов длиной меньше max_allowed_spaces на один пробел,
    но сохраняет более длинные последовательности (предполагая, что это намеренное форматирование).
    
    :param text: Исходный текст с метками
    :param max_allowed_spaces: Максимальное количество пробелов, которые считаются "случайными" и подлежат очистке
    :return: Очищенный текст
    """
    logging.debug(f"Очистка текста от None и лишних пробелов: '{text}' (макс. пробелов: {max_allowed_spaces})")
    
    # Шаг 1: Заменяем последовательности пробелов короче max_allowed_spaces на один пробел
    # Используем функцию замены с условием
    def replace_short_spaces(match):
        spaces = match.group(0)
        if len(spaces) < max_allowed_spaces:
            return ' '
        else:
            return spaces  # Сохраняем длинные последовательности пробелов
    
    # Применяем замену только к последовательностям пробелов (не табуляциям и другим whitespace)
    cleaned_text = re.sub(r' +', replace_short_spaces, text)
    
    # Шаг 2: Убираем пробелы в начале и конце строки
    cleaned_text = cleaned_text.strip()
    
    # Шаг 3: Убираем пробелы перед знаками препинания (только одиночные пробелы)
    cleaned_text = re.sub(r'\s+([.,)])', r'\1', cleaned_text)
    cleaned_text = re.sub(r'([(])\s+', r'\1', cleaned_text)
    
    logging.debug(f"Текст после очистки: '{cleaned_text}'")
    return cleaned_text

# Функция для замены меток в Word
def replace_in_doc(doc_path, replacements):
    doc = Document(doc_path)
    found_keys = set()  
    
    def replace_text(element):
        nonlocal found_keys
        if hasattr(element, "text"):
            original_text = ''.join(run.text for run in element.runs)
            modified_text = original_text
            
            # Сначала заменяем все найденные метки
            for key, value in replacements.items():
                for pattern in [f"{{{key}}}", f"{{ {key} }}", f"{{{key} }}", f"{{ {key}}}"]:
                    if pattern in modified_text:
                        # Обрезаем пробелы перед вставкой
                        trimmed_value = str(value).strip()
                        modified_text = modified_text.replace(pattern, trimmed_value)
                        found_keys.add(key)
                        logging.info(f"Замена метки '{pattern}' на '{trimmed_value}' в тексте: "
                                     f"'{original_text}' -> '{modified_text}'")
            
            # Затем удаляем все оставшиеся метки (для которых нет значений в карте)
            import re
            pattern = r"\{\s*[A-ZА-Я0-9]+\s*\}"
            modified_text = re.sub(pattern, "", modified_text)
            
            # Очистка от None и лишних пробелов
            modified_text = clean_text_from_none_and_spaces(modified_text)
            
            if original_text != modified_text:
                for run in element.runs:
                    run.text = ""
                element.runs[0].text = modified_text
    
    for p in doc.paragraphs:
        replace_text(p)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text(p)
    
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_text(p)
        for p in section.footer.paragraphs:
            replace_text(p)
    
    missing_keys = set(replacements.keys()) - found_keys
    if missing_keys:
        logging.warning(f"Предупреждение: следующие метки не найдены в документе: {missing_keys}")
    
    doc.save(doc_path)

# Функция для вставки изображений
def insert_images(doc_path, images_folder, brightness_factor=2.0, contrast_factor=3.0):
    doc = Document(doc_path)
    images_folder = Path(images_folder)
    special_images = {"1": "[IMAGE_1]", "1_2": "[IMAGE_1_2]", "1_3": "[IMAGE_1_3]"}
    for image_name, placeholder in special_images.items():
        image_path = images_folder / f"{image_name}.jpg"
        if not image_path.exists():
            for ext in [".jpg", ".jpeg", ".png"]:
                image_path = images_folder / f"{image_name}{ext}"
                if image_path.exists():
                    break
        if image_path.exists():
            print(f"Обработка специального изображения: {image_path}")
            try:
                image = Image.open(image_path)
                temp_image_path = images_folder / "temp_adjusted_image.jpg"
                image.save(temp_image_path, format="JPEG")
                print(f"Временное изображение сохранено: {temp_image_path}")
                for p in doc.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, "")
                        run = p.add_run()
                        run.add_picture(str(temp_image_path), width=Inches(6.4))
                        print(f"Изображение вставлено в метку {placeholder}")
                        break
                temp_image_path.unlink()
                print(f"Временное изображение удалено: {temp_image_path}")
            except Exception as e:
                print(f"Ошибка при обработке изображения {image_path}: {e}")
    doc.save(doc_path)
    print("Все изображения обработаны и вставлены.")

# Функция для поиска файла .txt в целевой папке
def find_txt_file(folder_path):
    txt_files = list(Path(folder_path).glob("*.txt"))
    logging.info(f"Поиск файлов .txt в папке: {folder_path}")
    if not txt_files:
        logging.error(f"В папке {folder_path} не найдено файлов .txt.")
        raise FileNotFoundError(f"В папке {folder_path} не найдено файлов .txt.")
    if len(txt_files) > 1:
        logging.warning(f"В папке {folder_path} найдено несколько TXT-файлов:")
        for i, file in enumerate(txt_files, start=1):
            logging.warning(f"{i}. {file.name}")
        while True:
            try:
                choice = int(input("Выберите номер файла для использования: "))
                if 1 <= choice <= len(txt_files):
                    selected_file = txt_files[choice - 1]
                    logging.info(f"Выбран файл: {selected_file}")
                    return selected_file
                else:
                    logging.warning("Неверный номер. Пожалуйста, выберите существующий вариант.")
            except ValueError:
                logging.warning("Пожалуйста, введите число.")
    selected_file = txt_files[0]
    logging.info(f"Выбран файл: {selected_file}")
    return selected_file

# Функция для чтения и фильтрации данных из .txt файла
def read_and_filter_txt_data(file_path):
    # Значения, которые нужно исключить (в разных вариантах регистра)
    exclude_values = {"AL", "Al", "al", "K+", "К+", "K-", "К-", "KV", "kv", "KF", "kf"}
    logging.info(f"Чтение и фильтрация данных из файла: {file_path}")

    # Чтение файла с указанной кодировкой
    with open(file_path, 'r', encoding='windows-1251') as file:
        lines = file.readlines()

    # Убираем заголовок и разбиваем строки на столбцы
    data = [line.strip().split('\t') for line in lines[1:]]

    # Получаем все уникальные идентификаторы объектов (первая колонка)
    # исключая значения из exclude_values
    object_ids = set()
    for line in data:
        if line[0] and line[0].strip() not in exclude_values:
            object_ids.add(line[0].strip())

    # Проверяем наличие профилей, заканчивающихся на -26
    profiles_with_26 = [obj_id for obj_id in object_ids if obj_id.endswith('-26')]

    # Если есть ровно один профиль с -26, выбираем его автоматически
    if len(profiles_with_26) == 1:
        selected_object_id = profiles_with_26[0]
        logging.info(f"Найден только один профиль с '-26', выбран автоматически: {selected_object_id}")
    else:
        # В остальных случаях используем старую логику выбора
        if len(profiles_with_26) > 1:
            logging.info(f"Найдено несколько профилей с '-26': {profiles_with_26}")
        else:
            logging.info("Профилей с '-26' не найдено, используется стандартный выбор")

        # Сортируем идентификаторы объектов по порядку номеров или по алфавиту
        sorted_object_ids = sorted(object_ids, key=lambda x: (x.isdigit(), x))

        # Выводим найденные идентификаторы для отладки
        logging.info(f"Найденные идентификаторы объектов: {sorted_object_ids}")

        # Проверяем, что есть хотя бы один уникальный идентификатор
        if len(sorted_object_ids) == 0:
            logging.error("В файле не найдено ни одного допустимого объекта.")
            raise ValueError("В файле не найдено ни одного допустимого объекта.")

        # Если найдено несколько уникальных идентификаторов, спрашиваем пользователя
        if len(sorted_object_ids) > 1:
            logging.warning("В файле найдено несколько различных объектов:")
            for i, obj_id in enumerate(sorted_object_ids, start=1):
                logging.warning(f"{i}. {obj_id}")
            while True:
                try:
                    choice = int(input("Выберите номер объекта для загрузки: ")) - 1
                    if 0 <= choice < len(sorted_object_ids):
                        selected_object_id = sorted_object_ids[choice]
                        break
                    else:
                        logging.warning("Некорректный выбор. Пожалуйста, введите номер из списка.")
                except ValueError:
                    logging.warning("Пожалуйста, введите числовое значение.")
        else:
            selected_object_id = sorted_object_ids[0]

    # Фильтруем данные, оставляя только строки с выбранным идентификатором объекта
    filtered_data = [line for line in data if line[0].strip() == selected_object_id]

    # Подготовка к сбору локусов с f
    loci_with_f = set()

    # Парсим заголовок для поиска колонок Marker и аллелей
    if lines:
        header = lines[0].strip().split('\t')

        # Поиск индекса колонки Marker
        marker_idx = None
        for i, col in enumerate(header):
            if col.lower().strip() in ('marker', 'маркер'):
                marker_idx = i
                break

        # Поиск индексов аллелей
        allele_indices = []
        first_allele_index = None
        for i, col in enumerate(header):
            col_lower = col.lower().strip()
            if col_lower.startswith('allele') or col_lower.startswith('аллель'):
                allele_indices.append(i)
                if first_allele_index is None:
                    first_allele_index = i
        if not allele_indices:
            # Если не нашли по заголовку, используем предположительные индексы (4..15)
            logging.warning("Не найдены столбцы аллелей по заголовку, используются индексы по умолчанию (4..15)")
            max_cols = max(len(line) for line in filtered_data) if filtered_data else 0
            allele_indices = list(range(4, min(16, max_cols)))
            first_allele_index = 4 if 4 in allele_indices else None

        # Сбор локусов, где встречается f (до замены)
        if marker_idx is not None and allele_indices:
            for row in filtered_data:
                marker = row[marker_idx].strip() if marker_idx < len(row) else ''
                # Проверяем все аллели на наличие f
                for idx in allele_indices:
                    if idx < len(row) and row[idx].strip().upper() == 'F':
                        if marker:
                            loci_with_f.add(marker)
                        break   # достаточно одной f в строке

        # Замена f на первый аллель
        if first_allele_index is not None and allele_indices:
            for row in filtered_data:
                first_val = row[first_allele_index].strip() if first_allele_index < len(row) else ""
                if not first_val or first_val.upper() == 'F':
                    continue
                for idx in allele_indices:
                    if idx < len(row) and row[idx].strip().upper() == 'F':
                        old_val = row[idx]
                        row[idx] = first_val
                        logging.debug(f"Замена f на {first_val} в строке (образец {selected_object_id}): {old_val} -> {row[idx]}")
        else:
            logging.warning("Не удалось определить индексы аллелей, замена f не производится.")

    # Выводим количество отфильтрованных строк данных для отладки
    logging.info(f"Количество отфильтрованных строк данных: {len(filtered_data)}")

    # Логируем структуру данных для отладки
    if filtered_data:
        first_line = filtered_data[0]
        logging.info(f"Первая строка данных содержит {len(first_line)} столбцов")
        logging.info(f"Содержимое первой строки: {first_line}")

    # Возвращаем кортеж (отфильтрованные данные, множество локусов с f)
    return filtered_data, loci_with_f

# Функция для чтения и фильтрации данных из .txt файла ДЛЯ WORD (БЕЗ замены f)
def read_and_filter_txt_data_for_word(file_path):
    """
    Читает и фильтрует данные из txt файла БЕЗ замены f на первый аллель.
    Используется для заполнения таблиц в документах Word, где f должно остаться как есть.
    """
    # Значения, которые нужно исключить (в разных вариантах регистра)
    exclude_values = {"AL", "Al", "al", "K+", "К+", "K-", "К-", "KV", "kv", "KF", "kf"}
    logging.info(f"Чтение и фильтрация данных для Word (БЕЗ замены f) из файла: {file_path}")

    # Чтение файла с указанной кодировкой
    with open(file_path, 'r', encoding='windows-1251') as file:
        lines = file.readlines()

    # Убираем заголовок и разбиваем строки на столбцы
    data = [line.strip().split('\t') for line in lines[1:]]

    # Получаем все уникальные идентификаторы объектов (первая колонка)
    # исключая значения из exclude_values
    object_ids = set()
    for line in data:
        if line[0] and line[0].strip() not in exclude_values:
            object_ids.add(line[0].strip())

    # Проверяем наличие профилей, заканчивающихся на -26
    profiles_with_26 = [obj_id for obj_id in object_ids if obj_id.endswith('-26')]

    # Если есть ровно один профиль с -26, выбираем его автоматически
    if len(profiles_with_26) == 1:
        selected_object_id = profiles_with_26[0]
        logging.info(f"Для Word: найден только один профиль с '-26', выбран автоматически: {selected_object_id}")
    else:
        # В остальных случаях используем старую логику выбора
        if len(profiles_with_26) > 1:
            logging.info(f"Для Word: найдено несколько профилей с '-26': {profiles_with_26}")
        else:
            logging.info("Для Word: профилей с '-26' не найдено, используется стандартный выбор")

        # Сортируем идентификаторы объектов по порядку номеров или по алфавиту
        sorted_object_ids = sorted(object_ids, key=lambda x: (x.isdigit(), x))

        # Выводим найденные идентификаторы для отладки
        logging.info(f"Для Word: найденные идентификаторы объектов: {sorted_object_ids}")

        # Проверяем, что есть хотя бы один уникальный идентификатор
        if len(sorted_object_ids) == 0:
            logging.error("Для Word: в файле не найдено ни одного допустимого объекта.")
            raise ValueError("В файле не найдено ни одного допустимого объекта.")

        # Если найдено несколько уникальных идентификаторов, используем тот же выбор, что и в основной функции
        if len(sorted_object_ids) > 1:
            logging.info("Для Word: используется тот же выбор объекта, что и для основной обработки")
            # Здесь мы должны использовать тот же selected_object_id, что был выбран в основной функции
            # Для этого мы вызовем основную функцию и получим её выбор
            temp_data, temp_loci = read_and_filter_txt_data(file_path)
            if temp_data:
                selected_object_id = temp_data[0][0]  # Берём ID из первой строки
                logging.info(f"Для Word: используется ID из основной функции: {selected_object_id}")
            else:
                raise ValueError("Не удалось определить выбранный объект для Word")
        else:
            selected_object_id = sorted_object_ids[0]

    # Фильтруем данные, оставляя только строки с выбранным идентификатором объекта
    # ВАЖНО: НЕ заменяем f на первый аллель!
    filtered_data = [line for line in data if line[0].strip() == selected_object_id]

    # Выводим количество отфильтрованных строк данных для отладки
    logging.info(f"Для Word: количество отфильтрованных строк данных (БЕЗ замены f): {len(filtered_data)}")

    # Логируем структуру данных для отладки
    if filtered_data:
        first_line = filtered_data[0]
        logging.info(f"Для Word: первая строка данных содержит {len(first_line)} столбцов")
        logging.info(f"Для Word: содержимое первой строки (БЕЗ замены f): {first_line}")

    # Проверяем наличие значений f для информирования
    f_found = False
    for row in filtered_data:
        # Проверяем колонки аллелей (обычно это столбцы 5-8 или около того)
        for idx in range(4, min(len(row), 16)):  # Проверяем до 16 колонки
            if idx < len(row) and row[idx].strip().upper() == 'F':
                f_found = True
                marker = row[3] if len(row) > 3 else 'неизвестный маркер'
                logging.info(f"Для Word: найдено значение 'f' в маркере {marker}, колонка {idx + 1}. Значение остаётся неизменным.")
                break

    if not f_found:
        logging.info("Для Word: значений 'f' в данных не обнаружено")

    return filtered_data

# Функция для вставки данных в таблицу Word
def insert_table_data(doc_path, txt_data, card_data, third_col_width=Inches(6)):
    doc = Document(doc_path)
    has_y_in_amel = False
    sorted_data = sort_data_by_locus(txt_data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "[TABLE_DATA]" in cell.text:
                    cell.text = ""
                    col1_width = table.columns[0].width
                    col2_width = table.columns[1].width
                    unique_markers = set()
                    for line in sorted_data:
                        marker = line[3].strip()
                        if marker.upper() in unique_markers:
                            continue
                        unique_markers.add(marker.upper())
                        
                        # Читаем все аллели (до 4)
                        alleles_list = []
                        exclude_values = {"AL", "Al", "al", "K+", "К+", "K-", "К-", "KV", "kv", "KF", "kf"}
                        for col_index in range(5, 9):  # Столбцы 5-8: Allele 1-4
                            if col_index < len(line) and line[col_index] and line[col_index].strip():
                                allele_value = line[col_index].strip()
                                if allele_value not in exclude_values:
                                    alleles_list.append(allele_value)
                        alleles = ",".join(alleles_list) if alleles_list else "-"
                        
                        # Проверка AMEL (оставляем старую логику для обратной совместимости)
                        if marker.upper() == "AMEL":
                            allele1 = line[5] if len(line) > 5 and line[5] else ""
                            allele2 = line[6] if len(line) > 6 and line[6] else ""
                            if "Y" in [allele1.upper(), allele2.upper()]:
                                has_y_in_amel = True
                                print(f"Найдена буква Y в маркере AMEL: Allele1={allele1}, Allele2={allele2}")
                        
                        new_row = table.add_row()
                        new_row.cells[0].text = marker
                        new_row.cells[1].text = alleles
                        run = new_row.cells[0].paragraphs[0].runs[0]
                        run.font.name = 'Times New Roman'
                        r = run._element
                        rPr = r.get_or_add_rPr()
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        rPr.append(rFonts)
                        run = new_row.cells[1].paragraphs[0].runs[0]
                        run.font.name = 'Times New Roman'
                        r = run._element
                        rPr = r.get_or_add_rPr()
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        rPr.append(rFonts)
                    rows_to_remove = []
                    for row in table.rows:
                        if all(cell.text.strip() == "" for cell in row.cells):
                            rows_to_remove.append(row)
                    for row in rows_to_remove:
                        table._tbl.remove(row._tr)

                    # Устанавливаем границы для каждой ячейки таблицы
                    for row in table.rows:
                        for cell in row.cells:
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcBorders = OxmlElement('w:tcBorders')
                            for border_name in ['top', 'left', 'bottom', 'right']:
                                border = OxmlElement(f'w:{border_name}')
                                border.set(qn('w:val'), 'single')
                                border.set(qn('w:sz'), '4')
                                border.set(qn('w:space'), '0')
                                border.set(qn('w:color'), '000000')
                                tcBorders.append(border)
                            tcPr.append(tcBorders)

                    for row in table.rows:
                        if len(row.cells) > 1:
                            cell = row.cells[1]
                            cell.paragraphs[0].alignment = 1
                    if not has_y_in_amel:
                        print("Буква Y не найдена в маркере AMEL. Добавляем примечание.")
                        p = doc.add_paragraph()
                        run = p.add_run("Примечание: ")
                        run.underline = True
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
                        p.add_run("(-) - прочерк означает отсутствие продуктов амплификации.").font.size = Pt(10)
                        table_element = table._tbl
                        parent_element = table_element.getparent()
                        parent_element.insert(parent_element.index(table_element) + 1, p._element)
                        next_element = p._element.getnext()
                        if next_element is not None and next_element.tag.endswith('p'):
                            next_paragraph = doc._body._element.index(next_element)
                            if not next_element.text.strip():
                                doc._body._element.remove(next_element)
                                print("Пустая строка после примечания удалена.")
                    table.add_column(third_col_width)
                    table.columns[0].width = col1_width
                    table.columns[1].width = col2_width
                    first_cell = table.cell(0, len(table.columns) - 1)
                    last_cell = table.cell(len(table.rows) - 1, len(table.columns) - 1)
                    tc = first_cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    vMerge = OxmlElement('w:vMerge')
                    vMerge.set(qn('w:val'), 'restart')
                    tcPr.append(vMerge)
                    for row in table.rows[1:]:
                        tc = row.cells[len(table.columns) - 1]._tc
                        tcPr = tc.get_or_add_tcPr()
                        vMerge = OxmlElement('w:vMerge')
                        vMerge.set(qn('w:val'), 'continue')
                        tcPr.append(vMerge)
                    for row in table.rows:
                        tr = row._tr
                        trPr = tr.get_or_add_trPr()
                        spacing = OxmlElement('w:trHeight')
                        spacing.set(qn('w:val'), "240")
                        spacing.set(qn('w:hRule'), "auto")
                        trPr.append(spacing)
                    images_folder = Path(card_data["4"])
                    nom_folder_name = f"{card_data['НОМ']}-26"
                    nom_folder = images_folder / nom_folder_name
                    image_path = nom_folder / "1.jpg"
                    if image_path.exists():
                        with Image.open(image_path) as img:
                            img_width, img_height = img.size
                            aspect_ratio = img_height / img_width
                            new_width = third_col_width
                            new_height = new_width * aspect_ratio
                            run = first_cell.paragraphs[0].add_run()
                            run.add_picture(str(image_path), width=new_width, height=new_height)
                    else:
                        logging.warning(f"Изображение {image_path} не найдено.")
                    tblPr = table._tblPr
                    tblW = tblPr.first_child_found_in("w:tblW")
                    if tblW is not None:
                        tblPr.remove(tblW)
                    break
    doc.save(doc_path)
    print("Примечание добавлено после таблицы.")

# Функция для поиска файла Excel с любым именем в текущей директории
def find_excel_file():
    """
    Ищет единственный Excel-файл в текущей директории.
    Возвращает путь к файлу или вызывает исключение, если файл не найден или их несколько.
    """
    # Получаем список всех Excel-файлов в текущей директории
    excel_files = list(Path('.').glob('*.xlsx')) + list(Path('.').glob('*.xls'))
    
    # Проверяем количество найденных файлов
    if not excel_files:
        raise FileNotFoundError("В текущей директории не найдено Excel-файлов (.xlsx или .xls)")
    if len(excel_files) > 1:
        raise ValueError(f"В директории найдено несколько Excel-файлов: {[f.name for f in excel_files]}. "
                         "Должен быть только один файл.")
    
    return excel_files[0]

# Функция для поиска файла Excel с VBA, указанного в Карте
def find_vba_file(card_data):
    vba_folder_path = Path(card_data["3"])  # Путь к папке с шаблонами
    vba_file_name = card_data["7"].strip()  # Имя файла VBA

    # Особый случай для ЭКЦ или ЭКЦ_образец_
    if card_data.get("1") == "ЭКЦ" or (card_data.get("1") and card_data.get("1").startswith("ЭКЦ_образец_")):
        vba_folder_path = Path(r"U:\ШАБЛОНЫ\Заключения\ЭКЦ")
        vba_file_name = "ИКЛ ЭКЦ.xlsm"  # Жестко заданное имя файла для этого случая

    logging.info(f"Путь к папке с шаблонами: {vba_folder_path}")
    logging.info(f"Имя файла VBA: {vba_file_name}")
    # Добавляем расширение файла, если оно отсутствует
    if not vba_file_name.endswith('.xlsm'):
        vba_file_name += '.xlsm'
    vba_file_path = vba_folder_path / vba_file_name
    logging.info(f"Полный путь к файлу VBA: {vba_file_path}")
    if not vba_file_path.exists():
        logging.error(f"Файл VBA не найден по пути: {vba_file_path}")
        raise FileNotFoundError(f"Файл VBA не найден по пути: {vba_file_path}")
    if not vba_file_path.is_file():
        logging.error(f"Путь {vba_file_path} не является файлом.")
        raise ValueError(f"Путь {vba_file_path} не является файлом.")
    logging.info(f"Найден файл VBA: {vba_file_path}")
    return vba_file_path

def is_file_open(file_path):
    """Проверяет, открыт ли файл другим процессом"""
    try:
        os.rename(file_path, file_path)
        return False
    except OSError:
        return True

def create_temp_txt_for_universal_table(original_txt_path, temp_dir_from_card):
    """
    Создает временный .txt файл нужного формата для "Универсальной_таблицы_для_генотипов"
    на основе исходного .txt файла. Копируются только строки, где Sample Name заканчивается на '-26'.

    Args:
        original_txt_path (str): Путь к исходному .txt файлу.
        temp_dir_from_card (str или Path): Путь к временной директории из card_data["4"].

    Returns:
        str: Путь к созданному временному файлу.
    """
    logger = logging.getLogger(__name__) # Используем именованный логгер
    logger.info(f"Создание временного txt файла для универсальной таблицы из: {original_txt_path}")
    logger.info(f"Временный файл будет сохранен в: {temp_dir_from_card}")

     # Словарь для сопоставления русских и английских заголовков
    header_mapping = {
        "название образца": ["sample name"],
        "название запуска": ["run name"],
        "панель": ["panel"],
        "маркер": ["marker"],
        "краситель": ["dye"],
        "аллель 1": ["allele 1"],
        "аллель 2": ["allele 2"],
        "аллель 3": ["allele 3"],
        "аллель 4": ["allele 4"],
        "аллель 5": ["allele 5"],
        "аллель 6": ["allele 6"],
        "аллель 7": ["allele 7"],
        "аллель 8": ["allele 8"],
        "аллель 9": ["allele 9"],
        "аллель 10": ["allele 10"],
        "аллель 11": ["allele 11"],
        "аллель 12": ["allele 12"],
        "высота 1": ["height 1"],
        "высота 2": ["height 2"],
        "высота 3": ["height 3"],
        "высота 4": ["height 4"],
        "высота 5": ["height 5"],
        "высота 6": ["height 6"],
        "высота 7": ["height 7"],
        "высота 8": ["height 8"],
        "высота 9": ["height 9"],
        "высота 10": ["height 10"],
        "высота 11": ["height 11"],
        "высота 12": ["height 12"],
        "os": ["os"],
        "bin": ["bin"],
        "phr": ["phr"],
        "lph": ["lph"],
        "mph": ["mph"],
        "spk": ["spk"],
        "an": ["an"],
        "acc": ["acc"],
        "bd": ["bd"],
        "cc": ["cc"],
        "ovl": ["ovl"],
        "gq": ["gq"],
    }

    # Определяем заголовки (всегда на английском)
    new_headers = [
        "Sample Name", "Marker",
        "Allele 1", "Allele 2", "Allele 3", "Allele 4", "Allele 5", "Allele 6",
        "Allele 7", "Allele 8", "Allele 9", "Allele 10", "Allele 11", "Allele 12",
        "Height 1", "Height 2", "Height 3", "Height 4", "Height 5", "Height 6",
        "Height 7", "Height 8", "Height 9", "Height 10", "Height 11", "Height 12",
        "OS", "PHR", "LPH", "MPH", "SPK", "AN", "ACC", "BD", "CC", "OVL", "GQ"
    ]

    # Создаем временный файл в указанной директории
    temp_txt_path = os.path.join(temp_dir_from_card, "temp_universal_input.txt")
    logger.debug(f"Путь к временному файлу: {temp_txt_path}")

    try:
        # Открываем файл для записи (всегда используем табуляцию)
        with open(original_txt_path, 'r', encoding='utf-8') as infile, \
                open(temp_txt_path, 'w', encoding='utf-8', newline='') as outfile:
            
            delimiter = '\t'  # Жестко задаем табуляцию
            logger.info(f"Используемый разделитель: '{repr(delimiter)}'") # repr для отладки \t
            reader = csv.reader(infile, delimiter=delimiter)
            writer = csv.writer(outfile, delimiter='\t')  # Новый файл всегда использует табуляцию

            # Читаем заголовки исходного файла
            try:
                original_headers = next(reader)
                # Нормализуем заголовки: убираем пробелы по краям и приводим к нижнему регистру
                original_headers_normalized = [h.strip().lower() for h in original_headers]
                logger.debug(f"Заголовки исходного файла (оригинал): {original_headers}")
                logger.info(f"Заголовки исходного файла (нормализованные): {original_headers_normalized}")
            except StopIteration:
                raise ValueError("Исходный txt файл пуст или не содержит заголовков.")

            # Создаем словарь для сопоставления нормализованных заголовков с индексами
            header_to_index = {header: i for i, header in enumerate(original_headers_normalized)}
            logger.debug(f"Словарь заголовков исходного файла (нормализованный -> индекс): {header_to_index}")

            # Индексы для аллелей и высот (ищем по имени, учитывая оба языка)
            allele_indices = []
            height_indices = []
            # Также будем искать индекс Dye для потенциального использования Height
            dye_index = None

            # Сначала найдем индекс Dye, если он есть
            dye_keys = header_mapping.get("краситель", []) # Русский ключ
            dye_keys_normalized = [k.lower().strip() for k in dye_keys]
            for key in ["краситель"] + dye_keys_normalized: # Проверяем и русский, и английский
                dye_idx = header_to_index.get(key)
                if dye_idx is not None:
                    dye_index = dye_idx
                    logger.info(f"Найден столбец 'Dye'/'Краситель' с индексом: {dye_index}")
                    break

            # Поиск индексов Allele и Height
            for i in range(1, 13):  # 1 до 12
                # Поиск Allele i
                allele_header_ru = f"аллель {i}"
                a_idx_ru = header_to_index.get(allele_header_ru)
                a_idx_en = None
                en_keys = header_mapping.get(allele_header_ru, []) # Получаем список английских эквивалентов
                for en_key in en_keys:
                     a_idx_en = header_to_index.get(en_key.lower().strip()) # Нормализация английского ключа
                     if a_idx_en is not None:
                         break # Нашли английский вариант

                final_allele_idx = a_idx_ru if a_idx_ru is not None else a_idx_en
                if final_allele_idx is not None:
                    logger.info(f"Найден столбец '{'Аллель' if a_idx_ru is not None else 'Allele'} {i}' с индексом: {final_allele_idx}")
                else:
                    logger.warning(f"Столбец 'Аллель {i}'/'Allele {i}' НЕ НАЙДЕН.")
                allele_indices.append(final_allele_idx)

                # Поиск Height i
                height_header_ru = f"высота {i}"
                h_idx_ru = header_to_index.get(height_header_ru)
                h_idx_en = None
                en_keys_h = header_mapping.get(height_header_ru, []) # Получаем список английских эквивалентов
                for en_key_h in en_keys_h:
                     h_idx_en = header_to_index.get(en_key_h.lower().strip()) # Нормализация английского ключа
                     if h_idx_en is not None:
                         break # Нашли английский вариант

                final_height_idx = h_idx_ru if h_idx_ru is not None else h_idx_en
                if final_height_idx is not None:
                    logger.info(f"Найден столбец '{'Высота' if h_idx_ru is not None else 'Height'} {i}' с индексом: {final_height_idx}")
                else:
                    logger.warning(f"Столбец 'Высота {i}'/'Height {i}' НЕ НАЙДЕН.")
                height_indices.append(final_height_idx)

            # Индексы для других фиксированных колонок (Sample Name, Marker, OS и т.д.)
            # Создаем обратный словарь для удобства поиска
            inverted_header_mapping = {}
            for file_key, equivalent_keys in header_mapping.items():
                for eq_key in equivalent_keys:
                     eq_key_norm = eq_key.strip().lower()
                     # inverted_header_mapping[английский_эквивалент] = русский_ключ_из_файла
                     inverted_header_mapping[eq_key_norm] = file_key
            logger.debug(f"Инвертированный словарь сопоставлений (англ. эквивалент -> русский ключ): {inverted_header_mapping}")

            # Заполняем fixed_columns_map с английским ключом
            fixed_columns_map = {}
            for en_key in ["sample name", "run name", "panel", "marker", "dye", "os", "bin", "phr", "lph", "mph", "spk", "an", "acc", "bd", "cc", "ovl", "gq"]:
                 # Найти русский ключ в файле, соответствующий этому английскому ключу
                 file_key = inverted_header_mapping.get(en_key)
                 idx_file = None
                 idx_en = None
                 if file_key:
                     idx_file = header_to_index.get(file_key) # Индекс русского заголовка
                 idx_en = header_to_index.get(en_key) # Индекс английского заголовка (если он есть в файле)
                 final_idx = idx_file if idx_file is not None else idx_en
                 fixed_columns_map[en_key] = final_idx # Сохраняем индекс под английским ключом
                 if final_idx is not None:
                     logger.info(f"Найден столбец '{en_key}' с индексом: {final_idx} (файл: {file_key}, англ: {idx_en})")
                 else:
                     logger.warning(f"Столбец '{en_key}' НЕ НАЙДЕН.")

            # Проверяем, есть ли колонка Sample Name (на любом языке)
            sample_name_idx = fixed_columns_map.get("sample name")
            if sample_name_idx is None:
                logger.warning("Колонка 'Sample Name'/'Название образца' не найдена в исходном файле. Будут скопированы все строки.")

            # Записываем новые заголовки (всегда на английском)
            writer.writerow(new_headers)
            logger.info(f"Записаны новые заголовки: {new_headers}")

            # Обрабатываем строки данных
            for row_num, row in enumerate(reader):
                # Проверяем, нужно ли копировать строку (на основе Sample Name)
                should_copy = True
                if sample_name_idx is not None and sample_name_idx < len(row):
                    sample_name_value = row[sample_name_idx].strip()
                    if not sample_name_value.endswith('-26'):
                        should_copy = False
                        logger.debug(f"Строка {row_num + 2} пропущена: Sample Name '{sample_name_value}' не заканчивается на '-26'")
                if not should_copy:
                    continue  # Пропускаем эту строку и переходим к следующей

                # Замена f: определяем значение первого аллеля для текущей строки
                first_allele_value = ""
                if allele_indices and allele_indices[0] is not None and allele_indices[0] < len(row):
                    first_allele_value = row[allele_indices[0]].strip()

                # Обработка строки для копирования
                new_row = []
                for header in new_headers:
                    cell_value = ""  # Значение по умолчанию
                    if header == "Marker":
                        original_marker_idx = fixed_columns_map.get("marker")
                        if original_marker_idx is not None and original_marker_idx < len(row):
                            marker_val = row[original_marker_idx].strip()
                            # Логика обработки Marker (остаётся без изменений)
                            if marker_val.upper() in ["AMEL", "AMELOGENIN"]:
                                cell_value = "Amel"
                            else:
                                cell_value = marker_val

                    elif header.startswith("Allele "):
                        # Используем предварительно найденные индексы
                        idx_in_header_list = int(header.split(" ")[1]) - 1
                        if idx_in_header_list < len(allele_indices):
                            orig_col_idx = allele_indices[idx_in_header_list]
                            if orig_col_idx is not None and orig_col_idx < len(row):
                                val = row[orig_col_idx].strip()
                                # === ЗАМЕНА f ===
                                if val.upper() == 'F' and first_allele_value and first_allele_value.upper() != 'F':
                                    logger.debug(f"Замена f на {first_allele_value} в строке {row_num+2}, аллель {idx_in_header_list+1}")
                                    val = first_allele_value
                                cell_value = val
                    elif header.startswith("Height "):
                        # Используем предварительно найденные индексы
                        idx_in_header_list = int(header.split(" ")[1]) - 1
                        if idx_in_header_list < len(height_indices):
                            orig_col_idx = height_indices[idx_in_header_list]
                            if orig_col_idx is not None and orig_col_idx < len(row):
                                cell_value = row[orig_col_idx].strip()
                    else: # Для других фиксированных колонок (OS, PHR и т.д.)
                        en_key_for_fixed = header.lower().strip()
                        orig_col_idx = fixed_columns_map.get(en_key_for_fixed)
                        if orig_col_idx is not None and orig_col_idx < len(row):
                            cell_value = row[orig_col_idx].strip()
                    new_row.append(cell_value)
                writer.writerow(new_row)

        logger.info(f"Временный txt файл успешно создан: {temp_txt_path}")
        return temp_txt_path

    except Exception as e:
        if os.path.exists(temp_txt_path):
            try:
                os.remove(temp_txt_path)
                logger.info(f"Временный файл {temp_txt_path} удален из-за ошибки.")
            except Exception as rm_e:
                logger.warning(f"Не удалось удалить временный файл {temp_txt_path} после ошибки: {rm_e}")
        logger.error(f"Ошибка при создании временного txt файла: {e}")
        raise

def replace_f_in_txt_file(original_txt_path, output_txt_path):
    """
    Читает исходный txt-файл, заменяет все значения 'f' (регистронезависимо) в колонках аллелей
    на значение первого аллеля в той же строке. Сохраняет результат в output_txt_path.
    Возвращает output_txt_path.
    """
    logger = logging.getLogger(__name__)
    logger.info(f"Замена f в файле {original_txt_path} -> {output_txt_path}")

    with open(original_txt_path, 'r', encoding='utf-8') as infile:
        lines = infile.readlines()

    if not lines:
        raise ValueError("Файл пуст")

    header = lines[0].strip().split('\t')
    # Определяем индексы колонок аллелей
    allele_indices = []
    for i, col in enumerate(header):
        if col.lower().startswith('allele') or col.lower().startswith('аллель'):
            allele_indices.append(i)
    if not allele_indices:
        logger.warning("Столбцы аллелей не найдены по заголовку, замена f не производится. Файл будет скопирован без изменений.")
        shutil.copy2(original_txt_path, output_txt_path)
        return output_txt_path

    first_allele_idx = allele_indices[0]

    # Обрабатываем строки данных (начиная со второй строки)
    processed_lines = [lines[0]]  # заголовок
    for line_num, line in enumerate(lines[1:], start=2):
        row = line.strip().split('\t')
        # Замена f
        first_val = row[first_allele_idx].strip() if first_allele_idx < len(row) else ""
        if first_val and first_val.upper() != 'F':
            for idx in allele_indices:
                if idx < len(row) and row[idx].strip().upper() == 'F':
                    row[idx] = first_val
                    logger.debug(f"Строка {line_num}: замена f на {first_val} в колонке {idx+1}")
        # Собираем строку обратно
        processed_lines.append('\t'.join(row) + '\n')

    with open(output_txt_path, 'w', encoding='utf-8') as outfile:
        outfile.writelines(processed_lines)

    logger.info(f"Файл с заменёнными f сохранён: {output_txt_path}")
    return output_txt_path

# Функция для вызова макроса VBA 2-ДНК
def run_vba_macro_2_DNK(vba_path, txt_path, output_folder, card_data, sheet_name="2-ДНК", loci_with_f=None):
    """
    Запускает макрос VBA в Excel файле с разными настройками в зависимости от листа.
    """
    try:
        # Преобразуем пути в абсолютные
        vba_path = os.path.abspath(vba_path)
        txt_path = os.path.abspath(txt_path)
        output_folder = os.path.abspath(output_folder)

        # Проверка существования файлов и папок
        if not os.path.exists(vba_path):
            logging.error(f"VBA файл не найден: {vba_path}")
            raise FileNotFoundError(f"VBA файл не найден: {vba_path}")
        if not os.path.exists(txt_path):
            logging.error(f"TXT файл не найден: {txt_path}")
            raise FileNotFoundError(f"TXT файл не найден: {txt_path}")
        if not os.path.exists(output_folder):
            logging.error(f"Папка назначения не найдена: {output_folder}")
            raise FileNotFoundError(f"Папка назначения не найдена: {output_folder}")

        # Создаем путь для файла
        if sheet_name == "ИКЛ":
            # Безопасное извлечение и преобразование значений
            fi = str(card_data.get('ФИ', '')).strip() if card_data.get('ФИ') is not None else ''
            ii = str(card_data.get('ИИ', '')).strip() if card_data.get('ИИ') is not None else ''
            oi = str(card_data.get('ОИ', '')).strip() if card_data.get('ОИ') is not None else ''
            nom = str(card_data.get('НОМ', '')).strip() if card_data.get('НОМ') is not None else ''
            
            # Формируем ФИО строку, исключая пустые части
            fio_parts = [part for part in [fi, ii, oi] if part and part != 'None']
            fio_string = ' '.join(fio_parts)
            
            # Формируем имя файла
            if fio_string:
                output_filename = f"{fio_string} {nom}-26.xlsm"
            else:
                output_filename = f"{nom}-26.xlsm"
                
            logging.info(f"Сформировано имя файла для ИКЛ: {output_filename}")
        else:
            # Аналогичная обработка для другого случая
            fi = str(card_data.get('ФИ', '')).strip() if card_data.get('ФИ') is not None else ''
            ii = str(card_data.get('ИИ', '')).strip() if card_data.get('ИИ') is not None else ''
            oi = str(card_data.get('ОИ', '')).strip() if card_data.get('ОИ') is not None else ''
            dr = str(card_data.get('ДР', '')).strip() if card_data.get('ДР') is not None else ''
            
            fio_parts = [part for part in [fi, ii, oi] if part and part != 'None']
            fio_string = ' '.join(fio_parts)
            
            if fio_string:
                output_filename = f"{fio_string} {dr}.xlsm"
            else:
                output_filename = f"{dr}.xlsm"
                
            logging.info(f"Сформировано имя файла: {output_filename}")

        output_path = os.path.join(output_folder, output_filename)

        logging.info(f"Начало выполнения макроса:")
        logging.info(f"VBA файл: {vba_path}")
        logging.info(f"TXT файл: {txt_path}")
        logging.info(f"Результат будет сохранен как: {output_path}")

        # Инициализация COM-объектов
        pythoncom.CoInitialize()
        excel = None
        workbook = None

        max_retries = 3
        retry_delay = 1

        try:
            for attempt in range(max_retries):
                try:
                    # Создаем экземпляр Excel
                    excel = win32.Dispatch("Excel.Application")
                    excel.Visible = False
                    excel.DisplayAlerts = False
                    excel.EnableEvents = False
                    excel.ScreenUpdating = False

                    logging.info(f"Excel запущен (попытка {attempt + 1})")

                    # Копируем файл VBA в выходную папку
                    shutil.copy2(vba_path, output_path)
                    logging.info(f"Создана копия шаблона: {output_path}")

                    # Открываем копию файла
                    workbook = excel.Workbooks.Open(output_path)
                    logging.info("Файл открыт успешно")

                    # Очистка данных
                    try:
                        workbook.Application.Run("ClearData")
                        logging.info("Данные очищены")
                    except Exception as e:
                        logging.warning(f"Не удалось очистить данные: {e}")

                    # Получение листа
                    try:
                        dnk_sheet = workbook.Sheets(sheet_name)
                    except Exception as e:
                        logging.error(f"Не удалось найти лист '{sheet_name}' в файле: {e}")
                        raise ValueError(f"Лист '{sheet_name}' не найден в Excel-файле") from e

                    # Определение ячеек в зависимости от листа
                    if sheet_name == "ИКЛ":
                        target_cell = "A25"
                        path_cell = "X11"
                    else:
                        target_cell = "A28"
                        path_cell = "X14"

                    # Устанавливаем путь к txt файлу
                    dnk_sheet.Range(path_cell).Value = txt_path
                    logging.info(f"Установлен путь к txt файлу в ячейку {path_cell}")

                    # Активируем нужный лист и ячейку
                    dnk_sheet.Activate()
                    dnk_sheet.Range(target_cell).Select()

                    # Выполняем макрос cmdToDo
                    logging.info("Запуск макроса cmdToDo...")
                    workbook.Application.Run("mdlTemplate.cmdToDo")

                    # Обработка диалогового окна
                    time.sleep(0.5)
                    try:
                        excel.SendKeys("{DOWN}")
                        time.sleep(0.1)
                        excel.SendKeys("{ENTER}")
                        logging.info("Диалоговое окно обработано")
                    except Exception as e:
                        logging.warning(f"Ошибка при обработке диалога: {e}")

                    # Проверяем результат
                    time.sleep(0.5)
                    if dnk_sheet.Range(target_cell).Value:
                        logging.info("Макрос выполнен успешно")
                        break
                    else:
                        raise ValueError(f"Данные не записаны в ячейку {target_cell}")

                except Exception as e:
                    if attempt < max_retries - 1:
                        logging.warning(f"Попытка {attempt + 1} не удалась: {e}")
                        time.sleep(retry_delay)
                        if workbook:
                            try:
                                workbook.Close(False)
                            except:
                                pass
                        if excel:
                            try:
                                excel.Quit()
                            except:
                                pass
                        continue
                    raise

            # Запись примечания для f
            if sheet_name == "ИКЛ" and loci_with_f:
                try:
                    ik_sheet = workbook.Sheets("ИКЛ")
                    loci_list = sorted(loci_with_f)
                    text = f"В локусе(-ах) {', '.join(loci_list)} возможно выпадение гетерозиготного аллеля"
                    ik_sheet.Range("A22").Value = text
                    logging.info(f"В ячейку A22 записан текст: {text}")
                except Exception as e:
                    logging.warning(f"Не удалось записать текст в A22: {e}")

            # Замена меток в Excel
            replace_in_excel_com(workbook, card_data)

            # Сохраняем результат
            workbook.Save()
            logging.info(f"Файл сохранен: {output_path}")

            # Закрываем workbook
            workbook.Close(True)
            logging.info("Workbook закрыт")

            return output_path

        except Exception as e:
            logging.error(f"Критическая ошибка при выполнении макроса: {e}")
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)
                    logging.info("Удален временный файл после ошибки")
                except:
                    pass
            raise

        finally:
            # Закрываем все объекты
            if workbook:
                try:
                    workbook.Close(True)
                except:
                    pass
            if excel:
                try:
                    excel.Quit()
                except:
                    pass
            pythoncom.CoUninitialize()

    except Exception as e:
        logging.error(f"Ошибка при выполнении макроса VBA: {e}")
        raise

# Функция для вызова макроса VBA ИК-2
def run_vba_macro_IK_2(vba_path, txt_path, output_folder, card_data):
    """
    Запускает макрос VBA IK-2 в Excel файле с оптимизированными настройками и улучшенной обработкой ошибок.
    Args:
    vba_path (str): Путь к файлу Excel с макросом
    txt_path (str): Путь к txt файлу с данными
    output_folder (str): Путь к папке для сохранения результата
    card_data (dict): Данные из файла Карта
    Returns:
    str: Путь к созданному файлу
    Raises:
    FileNotFoundError: Если файлы не найдены
    ValueError: При ошибках в данных или выполнении макроса
    """
    try:
        # Преобразуем пути в абсолютные
        vba_path = os.path.abspath(vba_path)
        txt_path = os.path.abspath(txt_path)
        output_folder = os.path.abspath(output_folder)
        
        # Проверка существования файлов и папок
        if not os.path.exists(vba_path):
            logging.error(f"VBA файл не найден: {vba_path}")
            raise FileNotFoundError(f"VBA файл не найден: {vba_path}")
        if not os.path.exists(txt_path):
            logging.error(f"TXT файл не найден: {txt_path}")
            raise FileNotFoundError(f"TXT файл не найден: {txt_path}")
        if not os.path.exists(output_folder):
            logging.error(f"Папка назначения не найдена: {output_folder}")
            raise FileNotFoundError(f"Папка назначения не найдена: {output_folder}")
        
        # Создаем путь для файла
        tr_value = str(card_data.get('ТР', '')).replace('/', '-').replace('\\', '-')
        output_filename = f"{tr_value}.xlsm"
        output_path = os.path.join(output_folder, output_filename)
        logging.info(f"Начало выполнения макроса:")
        logging.info(f"VBA файл: {vba_path}")
        logging.info(f"TXT файл: {txt_path}")
        logging.info(f"Результат будет сохранен как: {output_path}")

        # Инициализация COM-объектов
        pythoncom.CoInitialize()
        excel = None
        workbook = None
        max_retries = 3  # Максимальное количество попыток
        retry_delay = 1  # Задержка между попытками (в секундах)
        
        try:
            for attempt in range(max_retries):
                try:
                    # Создаем экземпляр Excel
                    excel = win32.Dispatch("Excel.Application")
                    
                    # Оптимизированные настройки Excel для повышения производительности
                    excel.Visible = False  # Делаем Excel невидимым
                    excel.DisplayAlerts = False  # Отключаем всплывающие уведомления
                    excel.EnableEvents = False  # Отключаем события
                    excel.ScreenUpdating = False  # Отключаем обновление экрана
                    logging.info(f"Excel запущен (попытка {attempt + 1})")
                    
                    # Копируем файл VBA в выходную папку
                    shutil.copy2(vba_path, output_path)
                    logging.info(f"Создана копия шаблона: {output_path}")
                    
                    # Открываем копию файла
                    workbook = excel.Workbooks.Open(output_path)
                    logging.info("Файл открыт успешно")
                    
                    # Очищаем данные в Excel с помощью макроса ClearData (без указания пути)
                    try:
                        workbook.Application.Run("ClearData")
                        logging.info("Данные очищены")
                    except Exception as e:
                        logging.warning(f"Не удалось очистить данные: {e}")
                    
                    # Устанавливаем путь к txt файлу в ячейку X14 на листе "ИК-2"
                    ik_sheet = workbook.Sheets("ИК-2")
                    ik_sheet.Range("X18").Value = txt_path
                    
                    # Также устанавливаем путь в ячейке для последнего файла
                    # (в соответствии с VBA-кодом это поле PC_RANGE_LDR_LASTFILE)
                    try:
                        # Попытка установить значение напрямую в рабочую область, как в VBA
                        workbook.Names("ldrLastFile").RefersToRange.Value = txt_path
                    except Exception as e:
                        logging.warning(f"Не удалось установить ldrLastFile: {e}")
                    
                    logging.info(f"Установлен путь к txt файлу")

                    # Активируем нужный лист и ячейку
                    ik_sheet.Activate()
                    ik_sheet.Range("A31").Select()
                    
                    # Вызов макроса через разные варианты имен модулей
                    try:
                        logging.info("Запуск макроса cmdToDo...")
                        workbook.Application.Run("mdlTemplate.cmdToDo")
                    except Exception as e:
                        # Если не удалось через предполагаемый модуль, пробуем стандартные варианты
                        try:
                            logging.info(f"Ошибка при вызове через mdlTemplate: {e}, пробуем другие варианты...")
                            # Вариант 2: через имя "Module1" (стандартное имя в VBA)
                            workbook.Application.Run("Module1.cmdToDo")
                        except Exception as e2:
                            # Вариант 3: без указания модуля
                            logging.info(f"Ошибка при вызове через Module1: {e2}, пробуем без модуля...")
                            workbook.Application.Run("cmdToDo")
                    
                    # Более длительная пауза для обработки диалога
                    time.sleep(0.7)
                    
                    # Улучшенная обработка диалогового окна (если оно появляется)
                    try:
                        # Последовательные нажатия клавиш для взаимодействия с диалогом выбора
                        excel.SendKeys("{DOWN}")  # Нажимаем стрелку вниз
                        time.sleep(0.2)
                        excel.SendKeys("{DOWN}")  # Нажимаем стрелку вниз еще раз
                        time.sleep(0.2)
                        excel.SendKeys("{ENTER}")  # Нажимаем Enter
                        time.sleep(0.3)
                        # В случае появления второго диалогового окна
                        excel.SendKeys("{ENTER}")  # Нажимаем Enter снова
                        logging.info("Диалоговые окна обработаны")
                    except Exception as e:
                        logging.warning(f"Ошибка при обработке диалога: {e}")
                    
                    # Проверяем результат выполнения макроса с более длительным ожиданием
                    time.sleep(0.3)
                    if ik_sheet.Range("A28").Value:
                        logging.info("Макрос выполнен успешно")
                        break  # Выходим из цикла попыток
                    else:
                        # Возможно, данные были сохранены в другой ячейке или результат нужно проверять иначе
                        # Все равно считаем успешным и продолжаем
                        logging.warning("Данные не записаны в ячейку A28, но продолжаем работу")
                        break

                except Exception as e:
                    if attempt < max_retries - 1:
                        logging.warning(f"Попытка {attempt + 1} не удалась: {e}")
                        time.sleep(retry_delay)  # Ждем перед повторной попыткой
                        if workbook:
                            try:
                                workbook.Close(False)  # Закрываем workbook без сохранения
                            except:
                                pass
                        if excel:
                            try:
                                excel.Quit()  # Закрываем Excel
                            except:
                                pass
                        continue  # Повторяем попытку
                    raise  # Если попытки закончились, выбрасываем исключение
            
            # Заменяем метки в Excel файле через COM-интерфейс
            replace_in_excel_com(workbook, card_data)
            
            # Сохраняем результат
            workbook.Save()
            logging.info(f"Файл сохранен: {output_path}")
            
            # Закрываем workbook
            workbook.Close(True)
            logging.info("Workbook закрыт")
            
            return output_path

        except Exception as e:
            logging.error(f"Критическая ошибка при выполнении макроса: {e}")
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)  # Удаляем временный файл после ошибки
                    logging.info("Удален временный файл после ошибки")
                except:
                    pass
            raise
        finally:
            # Закрываем все объекты
            if workbook:
                try:
                    workbook.Close(True)
                    logging.info("Workbook закрыт")
                except:
                    pass
            if excel:
                try:
                    excel.Quit()
                    logging.info("Excel закрыт")
                except:
                    pass
            pythoncom.CoUninitialize()  # Освобождаем COM-объекты
    except Exception as e:
        logging.error(f"Ошибка при выполнении макроса VBA IK-2: {e}")
        raise

def run_vba_macro_universal_table(vba_path, txt_path, output_folder, card_data):
    """
    Запускает макрос VBA в файле "Универсальная_таблица_для_генотипов".
    Открывает файл и запускает StartFormStart (пользователь вручную выбирает режим).
    """
    logging.info(f"Начало выполнения макроса универсальной таблицы (DIRECT):")
    logging.info(f"VBA файл (шаблон): {vba_path}")
    logging.info(f"TXT файл: {txt_path}")
    logging.info(f"Папка вывода: {output_folder}")
    # Упрощаем логирование карты
    # logging.debug(f"Данные карты (первые 5 элементов): {dict(list(card_data.items())[:5])}")

    pythoncom.CoInitialize()
    excel = None
    workbook = None
    final_created_file_path = None
    template_working_copy_path = None
    max_retries = 2 # Уменьшаем до 2 попыток
    retry_delay = 1

    try:
        # Подготовка пути к шаблону и рабочей копии
        template_path = Path(vba_path)
        if not template_path.exists():
            raise FileNotFoundError(f"Файл шаблона VBA не найден: {template_path}")

        working_copy_path = Path(output_folder) / template_path.name
        template_working_copy_path = str(working_copy_path)
        logging.info(f"Рабочая копия шаблона будет: {working_copy_path}")

        for attempt in range(max_retries):
            try:
                logging.info(f"Попытка {attempt + 1} открытия и работы с универсальной таблицей")

                # Копирование шаблона для работы
                if working_copy_path.exists():
                    try:
                        os.remove(working_copy_path)
                        logging.debug(f"Удалена предыдущая рабочая копия: {working_copy_path}")
                        time.sleep(0.5)
                    except PermissionError as pe:
                        logging.error(f"Не удалось удалить предыдущую рабочую копию {working_copy_path}: {pe}")
                        time.sleep(retry_delay)
                        if working_copy_path.exists():
                            raise pe

                shutil.copy2(template_path, working_copy_path)
                logging.info(f"Шаблон скопирован в рабочую папку: {working_copy_path}")

                # Открытие рабочей копии
                excel = win32.Dispatch("Excel.Application")
                excel.Visible = False
                excel.DisplayAlerts = False
                excel.EnableEvents = True
                excel.ScreenUpdating = True
                logging.info(f"Excel запущен (попытка {attempt + 1})")

                workbook = excel.Workbooks.Open(str(working_copy_path))
                logging.info("Рабочая копия файла открыта успешно")
                
                # Запуск StartFormStart для инициализации
                logging.info("Запуск макроса StartFormStart...")
                workbook.Application.Run("StartFormStart")
                logging.info("Макрос StartFormStart выполнен. Ожидание действий пользователя...")

                # Автозаполнение по меткам
                try:
                    logging.info("Начало автозаполнения универсальной таблицы данными из Карты...")
                    
                    # Получаем лист для заполнения
                    try:
                        target_sheet = workbook.Sheets("Основная_таблица") 
                        logging.info(f"Целевой лист для автозаполнения: 'Основная_таблица'")
                    except Exception as sheet_e:
                        error_msg = f"Лист 'Основная_таблица' не найден в рабочей книге: {sheet_e}"
                        logging.error(error_msg)
                        raise RuntimeError(error_msg)

                    # а) Заполнение ячейки C2
                    try:
                        value_c2 = ""
                        column_1_value = card_data.get("1", "").strip()
                        logging.info(f"Значение в поле '1' Карты для C2: '{column_1_value}'")

                        if column_1_value in ["СВО_Молов_образец_родственники", "СВО_Ростов_образец_родственники"]:
                            # Собираем ФИО родственника, исключая пустые и "None"
                            fio_parts = []
                            for part_key in ["ФИ", "ИИ", "ОИ"]:
                                val = str(card_data.get(part_key, "")).strip()
                                if val and val != "None":
                                    fio_parts.append(val)
                                    logging.debug(f"Добавлена часть ФИО: '{val}'")
                                else:
                                    if val == "None":
                                        logging.debug(f"Значение {part_key} равно 'None', пропущено.")
                                    else:
                                        logging.debug(f"Значение {part_key} пустое, пропущено.")
                            
                            fio_c2 = " ".join(fio_parts)
                            logging.debug(f"Сформирована ФИО: '{fio_c2}'")
                            
                            # Дата рождения родственника (ДР)
                            dr = str(card_data.get("ДР", "")).strip()
                            if dr and dr != "None":
                                value_c2 = f"{fio_c2} {dr} г.р."
                                logging.info(f"Добавлена дата рождения родственника: '{dr}'")
                            else:
                                value_c2 = fio_c2
                                if dr == "None":
                                    logging.debug("Дата рождения (ДР) равна 'None', не добавляем.")
                                else:
                                    logging.debug("Дата рождения (ДР) отсутствует, не добавляем.")
                        
                        elif column_1_value == "СВО_кость":
                            raw_tr = card_data.get("ТР", "")
                            if raw_tr and str(raw_tr).strip() not in ("", "None"):
                                value_c2 = str(raw_tr).strip()
                                logging.info(f"Заполняем C2 для 'СВО_кость': '{value_c2}'")
                            else:
                                value_c2 = ""
                                logging.info("Значение ТР для 'СВО_кость' отсутствует или равно 'None', ячейка C2 оставлена пустой.")
                        else:
                            logging.info(f"Значение в поле '1' ('{column_1_value}') не соответствует условиям для заполнения C2. C2 останется пустым или без изменений.")

                        if value_c2:  # Устанавливаем значение, только если оно не пустое
                            target_sheet.Range("C2").Value = value_c2
                            logging.info(f"Значение '{value_c2}' записано в ячейку C2.")
                        else:
                            # Явно устанавливаем пустую строку, чтобы очистить возможное предыдущее значение
                            target_sheet.Range("C2").Value = ""
                            logging.info("Ячейка C2 очищена (нет данных для заполнения).")

                    except Exception as e:
                        logging.error(f"Ошибка при заполнении ячейки C2: {e}", exc_info=True)

                    # б) Заполнение ячейки D2
                    try:
                        column_1_value = card_data.get("1", "").strip()
                        logging.info(f"Значение в поле '1' Карты для заполнения D2: '{column_1_value}'")

                        if column_1_value in ("СВО_Молов_образец_прямая идентификация", "СВО_Ростов_образец_прямая идентификация"):
                            logging.info("Тип 'Прямая идентификация' - ячейка D2 будет оставлена пустой.")
                            target_sheet.Range("D2").Value = ""
                        else:
                            raw_value_d2 = card_data.get("РОД", "")
                            logging.info(f"Исходное значение РОД из Карты: '{raw_value_d2}'")

                            if raw_value_d2:
                                replacements = {
                                    "матери": "мать",
                                    "отца": "отец",
                                    "сестры": "сестра",
                                    "брата": "брат",
                                    "сына": "сын",
                                    "дочери": "дочь"
                                }
                                value_d2 = str(raw_value_d2)
                                for old_suffix, new_word in replacements.items():
                                    if raw_value_d2.endswith(old_suffix):
                                        value_d2 = raw_value_d2[:-len(old_suffix)] + new_word
                                        logging.info(f"Применена замена окончания: '{raw_value_d2}' -> '{value_d2}'")
                                        break
                                target_sheet.Range("D2").Value = value_d2
                                logging.info(f"Значение '{value_d2}' записано в ячейку D2.")
                            else:
                                logging.info("Значение РОД в Карте отсутствует или пусто. D2 не будет изменено (или останется пустым).")

                    except Exception as e:
                        logging.error(f"Ошибка при заполнении ячейки D2: {e}", exc_info=True)

                    # в) Заполнение ячейки E2
                    try:
                        logging.info("Начало формирования значения для ячейки E2...")
                        
                        fpi = str(card_data.get("ФПИ", "")).strip()
                        ipi = str(card_data.get("ИПИ", "")).strip()
                        opi = str(card_data.get("ОПИ", "")).strip()
                        drp = str(card_data.get("ДРП", "")).strip()
                        vch = str(card_data.get("ВЧ", "")).strip()
                        ln = str(card_data.get("ЛН", "")).strip()

                        e2_parts = []
                        
                        if fpi and fpi != "None":
                            e2_parts.append(fpi)
                            logging.debug(f"Добавлена фамилия пропавшего: '{fpi}'")
                        else:
                            if fpi == "None":
                                logging.debug("Значение ФПИ равно 'None', пропущено.")
                            else:
                                logging.debug("Значение ФПИ пустое, пропущено.")
                        
                        if ipi and ipi != "None":
                            e2_parts.append(ipi)
                            logging.debug(f"Добавлено имя пропавшего: '{ipi}'")
                        else:
                            if ipi == "None":
                                logging.debug("Значение ИПИ равно 'None', пропущено.")
                            else:
                                logging.debug("Значение ИПИ пустое, пропущено.")
                        
                        if opi and opi != "None":
                            e2_parts.append(opi)
                            logging.debug(f"Добавлено отчество пропавшего: '{opi}'")
                        else:
                            if opi == "None":
                                logging.debug("Значение ОПИ равно 'None', пропущено.")
                            else:
                                logging.debug("Значение ОПИ пустое, пропущено.")
                        
                        if drp and drp != "None":
                            e2_parts.append(f"{drp} г.р.")
                            logging.info(f"Добавлена дата рождения пропавшего: '{drp}'")
                        else:
                            if drp == "None":
                                logging.debug("Значение ДРП равно 'None', пропущено.")
                            else:
                                logging.debug("Дата рождения пропавшего (ДРП) отсутствует, пропускаем")
                        
                        if vch and vch != "None":
                            e2_parts.append(f"в/ч {vch}")
                            logging.debug(f"Добавлена воинская часть: '{vch}'")
                        else:
                            if vch == "None":
                                logging.debug("Значение ВЧ равно 'None', пропущено.")
                            else:
                                logging.debug("Значение ВЧ отсутствует, пропускаем")
                        
                        if ln and ln != "None":
                            e2_parts.append(f"л/н {ln}")
                            logging.debug(f"Добавлен личный номер: '{ln}'")
                        else:
                            if ln == "None":
                                logging.debug("Значение ЛН равно 'None', пропущено.")
                            else:
                                logging.debug("Значение ЛН отсутствует, пропускаем")

                        value_e2 = " ".join(e2_parts)
                        logging.info(f"Сформированное значение для E2: '{value_e2}'")
                        target_sheet.Range("E2").Value = value_e2
                        logging.info(f"Значение '{value_e2}' записано в ячейку E2.")

                    except Exception as e:
                        logging.error(f"Ошибка при заполнении ячейки E2: {e}", exc_info=True)

                    logging.info("Автозаполнение универсальной таблицы завершено.")

                    # Сохраняем изменения в рабочей копии
                    try:
                        workbook.Save()
                        logging.info("Изменения в рабочей копии успешно сохранены.")
                    except Exception as save_e:
                        error_msg = f"Ошибка при сохранении рабочей копии после автозаполнения: {save_e}"
                        logging.error(error_msg)
                        raise RuntimeError(error_msg) from save_e

                except Exception as fill_error:
                    logging.error(f"Критическая ошибка во время автозаполнения универсальной таблицы: {fill_error}", exc_info=True)

                # Поиск созданного файла
                logging.info("Поиск созданного файла Excel...")
                
                all_xlsm_files = list(Path(output_folder).glob("*.xlsm"))
                # Исключаем служебные файлы и исходный шаблон
                created_files = [f for f in all_xlsm_files if not f.name.startswith("~$") and f.name != template_path.name]
                
                created_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
                
                if not created_files:
                    logging.warning("Файл .xlsm, созданный VBA, не найден.")
                    logging.info(f"Содержимое папки {output_folder}:")
                    for item in Path(output_folder).iterdir():
                        if item.is_file():
                            logging.info(f"  - {item.name} (modified: {time.ctime(item.stat().st_mtime)})")
                else:
                    final_created_file_path = str(created_files[0])
                    logging.info(f"Найден созданный файл: {final_created_file_path}")

                # Закрытие workbook
                if workbook:
                    try:
                        workbook.Close(SaveChanges=False)
                        logging.info("Рабочая копия шаблона закрыта без сохранения.")
                    except Exception as close_e:
                        logging.warning(f"Не удалось закрыть рабочую копию шаблона: {close_e}")
                    workbook = None

                if excel:
                    try:
                        excel.Quit()  # Закрываем Excel
                        logging.info("Excel полностью закрыт.")
                    except Exception as quit_e:
                        logging.warning(f"Не удалось закрыть Excel: {quit_e}")
                    excel = None

                pythoncom.CoUninitialize()

                # Переименование файла после закрытия Excel
                if final_created_file_path and Path(final_created_file_path).exists():
                    try:
                        # Определяем новое имя файла в зависимости от значения в карте
                        if card_data.get("1") == "СВО_кость":
                            tr_value = str(card_data.get('ТР', '')).strip() if card_data.get('ТР') is not None else ''
                            new_name = f"{tr_value}.xlsm"
                            logging.info("Формат имени файла для 'СВО_кость'")
                        else:
                            # Безопасное извлечение и преобразование значений
                            fi = str(card_data.get('ФИ', '')).strip() if card_data.get('ФИ') is not None else ''
                            ii = str(card_data.get('ИИ', '')).strip() if card_data.get('ИИ') is not None else ''
                            oi = str(card_data.get('ОИ', '')).strip() if card_data.get('ОИ') is not None else ''
                            dr = str(card_data.get('ДР', '')).strip() if card_data.get('ДР') is not None else ''
                            
                            # Формируем ФИО строку, исключая пустые части
                            fio_parts = [part for part in [fi, ii, oi] if part and part != 'None']
                            fio_string = ' '.join(fio_parts)
                            
                            # Формируем имя файла
                            if fio_string:
                                new_name = f"{fio_string} {dr}.xlsm"
                            else:
                                new_name = f"{dr}.xlsm"
                                
                            logging.info("Стандартный формат имени файла")
                        
                        new_path = Path(output_folder) / new_name
                        
                        time.sleep(1)
                        
                        os.rename(final_created_file_path, new_path)
                        logging.info(f"Файл успешно переименован: {new_path}")
                        
                        # Обновляем путь к файлу
                        final_created_file_path = str(new_path)
                    except Exception as rename_error:
                        logging.error(f"Ошибка при переименовании файла: {rename_error}")

                logging.info("Макрос универсальной таблицы (DIRECT) выполнен.")
                return final_created_file_path

            except Exception as e:
                logging.error(f"Ошибка при выполнении макроса универсальной таблицы (DIRECT) на попытке {attempt + 1}: {e}")
                logging.exception(e)
                
                # Закрываем объекты в случае ошибки
                if workbook:
                    try:
                        workbook.Close(SaveChanges=False)
                    except:
                        pass
                    workbook = None
                if excel:
                    try:
                        excel.Visible = False # Оставляем видимым даже при ошибке
                    except:
                        pass
                    excel = None
                
                # Освобождаем COM в случае ошибки
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass

                if attempt < max_retries - 1:
                    logging.info(f"Повторная попытка через {retry_delay} секунд...")
                    time.sleep(retry_delay)
                else:
                    logging.error("Все попытки выполнения макроса универсальной таблицы (DIRECT) исчерпаны.")
                    raise

    except Exception as e:
        logging.error(f"Критическая ошибка при выполнении макроса VBA для универсальной таблицы (DIRECT): {e}")
        logging.exception(e)
        raise
    finally:
        try:
            pythoncom.CoUninitialize()
        except:
            pass
        
        # Удаление рабочей копии шаблона
        if template_working_copy_path and os.path.exists(template_working_copy_path):
            try:
                # Небольшая пауза перед удалением, если файл ещё используется
                time.sleep(1)
                os.remove(template_working_copy_path)
                logging.info(f"Рабочая копия шаблона удалена: {template_working_copy_path}")
            except Exception as rm_e:
                logging.warning(f"Не удалось удалить рабочую копию шаблона {template_working_copy_path}: {rm_e}")
        elif template_working_copy_path:
            logging.info(f"Рабочая копия шаблона уже удалена или не существует: {template_working_copy_path}")

# Заполняем Excel по меткам
def replace_in_excel_com(workbook, replacements):
    try:
        # ВРЕМЕННЫЙ БЛОК - на время перехода с ИК на УТдГ: Добавляем словарь для перевода значений РОД
        rod_translation = {
            "матери": "сын",
            "отца": "сын",
            "сестры": "брат",
            "брата": "брат",
            "дочери": "отец",
            "сына": "отец",
            "родство не указано": "родственник (степень родства не указана)"
        }
        
        # Словарь для преобразования названий месяцев в числа
        month_translation = {
            "января": "01",
            "февраля": "02",
            "марта": "03",
            "апреля": "04",
            "мая": "05",
            "июня": "06",
            "июля": "07",
            "августа": "08",
            "сентября": "09",
            "октября": "10",
            "ноября": "11",
            "декабря": "12"
        }
        
        # Преобразование значений для исключения None
        processed_replacements = {}
        for key, value in replacements.items():
            if value is None:
                processed_replacements[key] = ""
            else:
                # Преобразуем в строку и убираем пробелы
                str_value = str(value).strip()
                # Дополнительная проверка на "None" как строку
                if str_value == "None":
                    processed_replacements[key] = ""
                else:
                    processed_replacements[key] = str_value
        
        # Если значение РОД присутствует в словаре, заменяем его
        if 'РОД' in processed_replacements:
            rod_value = processed_replacements['РОД']
            processed_replacements['РОД'] = rod_translation.get(rod_value, rod_value)
        
        # Если значение месяца присутствует в словаре, заменяем его
        if 'МО' in processed_replacements:
            mo_value = processed_replacements['МО'].lower()
            processed_replacements['МО'] = month_translation.get(mo_value, processed_replacements['МО'])
        
        # Импортируем модуль регулярных выражений
        import re
        
        for sheet in workbook.Sheets:
            for cell in sheet.UsedRange:
                if cell.Value and isinstance(cell.Value, str):
                    original_value = cell.Value.strip()
                    modified_value = original_value
                    
                    # Сначала заменяем все найденные метки
                    for key, value in processed_replacements.items():
                        for pattern in [f"{{{key}}}", f"{{ {key} }}", f"{{{key} }}", f"{{ {key}}}"]:
                            if pattern in modified_value:
                                modified_value = modified_value.replace(pattern, value)
                                logging.info(f"Замена метки '{pattern}' на '{value}' в ячейке: "
                                             f"'{original_value}' -> '{modified_value}'")
                    
                    # Затем удаляем все оставшиеся метки (для которых нет значений в карте)
                    pattern = r"\{\s*[A-ZА-Я0-9]+\s*\}"
                    modified_value = re.sub(pattern, "", modified_value)
                    
                    # Очистка от None и лишних пробелов
                    modified_value = clean_text_from_none_and_spaces(modified_value)
                    
                    if original_value != modified_value:
                        cell.Value = modified_value
                        logging.info(f"Ячейка обновлена: '{original_value}' -> '{cell.Value}'")
                    else:
                        logging.debug(f"Метки не найдены в ячейке: '{original_value}'")
        
        logging.info("Замены меток выполнены в Excel файле")
    
    except Exception as e:
        logging.error(f"Ошибка при замене меток в Excel файле: {e}")
        raise

# Конвертируем txt в xml
def convert_txt_to_xml(txt_file_path, xml_file_path, object_name_suffix):
    """
    Конвертирует txt файл в xml формат с улучшенной диагностикой ошибок.
    Сохраняет всю оригинальную логику, добавляя только обработку кодировок.
    """
    logging.info(f"Начало конвертации txt в xml: {txt_file_path}")
    
    # Попытка чтения файла: сначала utf-8-sig (автоматически удаляет BOM), затем fallback на windows-1251
    lines = None
    used_encoding = None
    
    try:
        with open(txt_file_path, 'r', encoding='utf-8-sig') as file:
            lines = file.readlines()
        used_encoding = 'utf-8-sig'
        logging.info(f"Файл успешно прочитан с кодировкой: utf-8-sig (BOM автоматически удалён)")
    except UnicodeDecodeError as e:
        logging.warning(f"Не удалось прочитать файл с кодировкой utf-8-sig: {e}")
        # Fallback на windows-1251
        try:
            with open(txt_file_path, 'r', encoding='windows-1251') as file:
                lines = file.readlines()
            used_encoding = 'windows-1251'
            logging.info(f"Файл успешно прочитан с fallback кодировкой: windows-1251")
        except UnicodeDecodeError as e2:
            logging.error(f"Не удалось прочитать файл ни с utf-8-sig, ни с windows-1251. Ошибки: utf-8-sig: {e}, windows-1251: {e2}")
            raise UnicodeDecodeError(f"Не удалось прочитать файл ни с utf-8-sig, ни с windows-1251. Ошибки: utf-8-sig: {e}, windows-1251: {e2}")
    
    if not lines:
        raise ValueError("Файл пуст или не содержит данных")
    
    logging.info(f"Использована кодировка: {used_encoding}, прочитано строк: {len(lines)}")
    
    # Парсинг заголовков
    header_line = lines[0].strip()
    headers = header_line.split('\t')
    
    # Добавляем диагностику только для отладки
    logging.info(f"Найдено колонок заголовков: {len(headers)}")
    logging.debug(f"Заголовки: {headers}")
    
    # Логика определения индексов
    sample_name_idx = None
    marker_idx = None
    allele_indices = []
    
    # Проверка наличия заголовков на английском языке
    if 'Sample Name' in headers:
        sample_name_idx = headers.index('Sample Name')
    elif 'Название образца' in headers:
        sample_name_idx = headers.index('Название образца')
    else:
        # Добавляем диагностику при ошибке
        logging.error(f"=== ДИАГНОСТИКА ОШИБКИ ===")
        logging.error(f"Файл: {txt_file_path}")
        logging.error(f"Кодировка: {used_encoding}")
        logging.error(f"Заголовки ({len(headers)}): {headers}")
        logging.error(f"Исходная строка заголовков: '{header_line}'")
        if len(header_line) != len(header_line.strip()):
            logging.error(f"Обнаружены пробелы в начале/конце строки")
        logging.error("=== КОНЕЦ ДИАГНОСТИКИ ===")
        raise ValueError("Заголовок 'Sample Name' или 'Название образца' не найден в файле.")
    
    if 'Marker' in headers:
        marker_idx = headers.index('Marker')
    elif 'Маркер' in headers:
        marker_idx = headers.index('Маркер')
    else:
        logging.error(f"Заголовок 'Marker' или 'Маркер' не найден. Доступные заголовки: {headers}")
        raise ValueError("Заголовок 'Marker' или 'Маркер' не найден в файле.")
    
    # Ищем все колонки с аллелями (до 4)
    for i, h in enumerate(headers):
        if h.startswith('Allele') or h.startswith('Аллель'):
            allele_indices.append(i)
    
    # Ограничиваем до 4 аллелей
    allele_indices = allele_indices[:4]
    
    if not allele_indices:
        logging.error(f"Заголовки 'Allele' или 'Аллель' не найдены. Доступные заголовки: {headers}")
        raise ValueError("Заголовки 'Allele' или 'Аллель' не найдены в файле.")
    
    logging.info(f"Найдены индексы: Sample Name={sample_name_idx}, Marker={marker_idx}, Alleles={allele_indices}")
    
    # Логика создания XML
    # Создание корневого элемента XML
    root = ET.Element("CODISImportFile", xmlns="urn:CODISImportFile-schema")
    ET.SubElement(root, "HEADERVERSION").text = "3.2"
    ET.SubElement(root, "MESSAGETYPE").text = "Import"
    ET.SubElement(root, "DESTINATIONORI").text = "destlab"
    ET.SubElement(root, "SOURCELAB").text = "srclab"
    ET.SubElement(root, "SUBMITBYUSERID").text = "gmidx"
    ET.SubElement(root, "SUBMITDATETIME").text = datetime.now().strftime("%Y-%m-%dT%H:%M:%S.%f%z")
    
    # Создание элемента SPECIMEN
    specimen = ET.SubElement(root, "SPECIMEN")
    ET.SubElement(specimen, "SPECIMENID").text = object_name_suffix
    ET.SubElement(specimen, "SPECIMENCATEGORY").text = "Other"
    
    # Логика чтения данных
    first_allele_index = allele_indices[0] if allele_indices else None
    processed_rows = 0

    for line in lines[1:]:
        data = line.strip().split('\t')
        sample_name = data[sample_name_idx]
        if sample_name.strip() not in {"AL", "Al", "K+", "К+", "K-", "К-"} and sample_name.strip() == object_name_suffix:
            marker = data[marker_idx]
            reading_by = "gmidx"
            reading_datetime = datetime.now().strftime("%Y-%m-%dT%H:%M:%S.%f%z")

            # Замена f: получаем первый аллель для этой строки
            first_allele_value = ""
            if first_allele_index is not None and first_allele_index < len(data):
                first_allele_value = data[first_allele_index].strip()
            
            # Создание элемента LOCUS
            locus = ET.SubElement(specimen, "LOCUS")
            ET.SubElement(locus, "LOCUSNAME").text = marker
            ET.SubElement(locus, "READINGBY").text = reading_by
            ET.SubElement(locus, "READINGDATETIME").text = reading_datetime
            
            # Добавление аллелей (до 4)
            for idx in allele_indices:
                if idx < len(data):
                    allele_value = data[idx].strip()
                    # === ЗАМЕНА f ===
                    if allele_value.upper() == 'F' and first_allele_value and first_allele_value.upper() != 'F':
                        logging.debug(f"В XML замена f на {first_allele_value} для маркера {marker}")
                        allele_value = first_allele_value
                    if allele_value:
                        allele = ET.SubElement(locus, "ALLELE")
                        ET.SubElement(allele, "ALLELEVALUE").text = allele_value
            
            processed_rows += 1
    
    logging.info(f"Обработано строк данных: {processed_rows}")
    
    # Логика сохранения
    tree = ET.ElementTree(root)
    tree.write(xml_file_path, encoding='utf-8', xml_declaration=True)
    logging.info(f"Файл XML сохранен: {xml_file_path}")

def ask_for_print(document_type):
    """Запрашивает подтверждение печати с параметризованным текстом"""
    logging.info(f"Запрос на печать {document_type}")
    response = input(f"Хотите распечатать {document_type}? (д/н) ").strip().lower()
    logging.info(f"Ответ пользователя на запрос '{document_type}': '{response}'")
    return response == "д"

def create_and_fill_envelope(card_data):
    """
    Создает временный файл конверта на основе шаблона и заполняет его данными из Карты.
    """
    try:
        template_folder = Path(card_data["3"])  # Путь к папке с шаблонами
        template_name = "Конверт"
        template_path = template_folder / f"{template_name}.docx"

        logging.info(f"Путь к шаблону конверта: {template_path}")
        if not template_path.exists():
            raise FileNotFoundError(f"Шаблон конверта '{template_name}' не найден в папке {template_folder}.")

        # Сохраняем во временную папку (код 4)
        temp_folder = Path(card_data["4"])
        temp_folder.mkdir(parents=True, exist_ok=True)
        logging.info(f"Временная папка для конверта: {temp_folder}")

        # Определяем имя файла
        output_filename = f"{card_data['ФИ']} {card_data['ИИ']} {card_data['ОИ']} {card_data['НОМ']}-26_конверт.docx"
        temp_path = temp_folder / output_filename

        logging.info(f"Создание временного файла конверта: {temp_path}")
        shutil.copy(template_path, temp_path)

        # Замена меток
        replacements = {code: str(value).strip() for code, value in card_data.items() if value is not None}
        replace_in_doc(temp_path, replacements)
        logging.info("Метки в конверте успешно заменены")

        return temp_path

    except Exception as e:
        logging.error(f"Ошибка при создании и заполнении конверта: {e}")
        raise

@rebuild_win32com_cache_on_error
def print_simple_copy(file_path, copies=1):
    """
    Упрощённая функция печати документа Word.
    Печатает указанный файл заданное количество раз (односторонняя печать).
    Включает черновой режим через Options.PrintDraft.
    """
    try:
        pythoncom.CoInitialize()
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False

        # --- Включение чернового режима ---
        original_print_draft = None
        try:
            # Пытаемся сохранить текущее состояние PrintDraft
            original_print_draft = word.Options.PrintDraft
            word.Options.PrintDraft = True
            logging.info("Черновой режим (Options.PrintDraft) включён.")
        except AttributeError:
            logging.warning("Свойство Options.PrintDraft не поддерживается в этой версии Word. Черновой режим не активирован.")
        except Exception as e:
            logging.warning(f"Не удалось установить черновой режим: {e}")
        # ---------------------------------

        # Определение принтера
        printer_name = get_default_printer_name(word)
        is_pantum = is_pantum_printer(printer_name)
        original_duplex = None
        logging.info(f"print_simple_copy: тип принтера {'Pantum' if is_pantum else 'Samsung/другой'}")

        if is_pantum:
            logging.info("Установка одностороннего режима для Pantum перед печатью...")
            success, original_duplex = set_printer_duplex_mode(printer_name, DMDUP_SIMPLEX)
            if not success:
                logging.warning("Не удалось установить односторонний режим через win32print, пробуем PowerShell...")
                set_printer_duplex_powershell(printer_name, False)
            time.sleep(1)

        try:
            # Открываем документ
            doc = word.Documents.Open(str(file_path))

            # Печатаем заданное количество копий
            logging.info(f"Печать документа: {file_path} ({copies} экз.)")
            for i in range(copies):
                logging.info(f"Печать экземпляра #{i+1}")
                doc.PrintOut(PrintToFile=False)
                time.sleep(1)

        except Exception as e:
            logging.error(f"Ошибка при печати документа: {e}")
            raise
        finally:
            # Закрываем документ без сохранения
            if 'doc' in locals():
                doc.Close(SaveChanges=False)

            # --- Восстановление исходного режима черновика ---
            if original_print_draft is not None:
                try:
                    word.Options.PrintDraft = original_print_draft
                    logging.info(f"Восстановлен исходный режим черновика: {original_print_draft}")
                except Exception as e:
                    logging.warning(f"Не удалось восстановить Options.PrintDraft: {e}")
            # -----------------------------------------------

            word.Quit()

    finally:
        # Восстановление оригинального режима принтера (для Pantum)
        if is_pantum and original_duplex is not None:
            logging.info("Восстановление оригинального режима принтера...")
            restore_printer_duplex_mode(printer_name, original_duplex)

        pythoncom.CoUninitialize()

@rebuild_win32com_cache_on_error
def print_document(zakl_path, card_data=None):
    """
    Функция печати документа Word с разной логикой в зависимости от типа шаблона.
    Теперь поддерживает принтеры Pantum (автоматическое переключение дуплекса)
    и включает черновой режим через Options.PrintDraft.
    """
    pythoncom.CoInitialize()
    try:
        logging.info(f"Начало печати документа: {zakl_path}")
        if card_data:
            logging.info(f"Данные карты переданы в print_document. Тип шаблона: '{card_data.get('1', 'НЕТ ДАННЫХ')}'")
        else:
            logging.warning("Данные карты НЕ переданы в print_document!")

        # Определяем, является ли документ типом ЭКЦ
        is_ekc_type = False
        template_value = ""
        if card_data and card_data.get("1"):
            template_value = str(card_data.get("1", "")).strip()
            logging.info(f"Проверка типа шаблона: '{template_value}'")
            if (template_value == "ЭКЦ" or 
                template_value == "ЭКЦ_нет результата" or 
                template_value.startswith("ЭКЦ_образец_")):
                is_ekc_type = True
                logging.info(f"Обнаружен шаблон типа ЭКЦ: '{template_value}'. Применяется упрощенная логика печати.")

        # Открываем документ в Word
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(zakl_path)

        # Включение чернового режима
        original_print_draft = None
        try:
            original_print_draft = word.Options.PrintDraft
            word.Options.PrintDraft = True
            logging.info("Черновой режим (Options.PrintDraft) включён.")
        except AttributeError:
            logging.warning("Свойство Options.PrintDraft не поддерживается в этой версии Word. Черновой режим не активирован.")
        except Exception as e:
            logging.warning(f"Не удалось установить черновой режим: {e}")

        # Определение принтера
        printer_name = get_default_printer_name(word)
        is_pantum = is_pantum_printer(printer_name)
        original_duplex = None
        logging.info(f"Тип принтера: {'Pantum' if is_pantum else 'Samsung/другой'}")

        try:
            if is_ekc_type:
                # УПРОЩЕННАЯ ЛОГИКА ДЛЯ ЭКЦ
                logging.info("1. Начало упрощенной печати для типа ЭКЦ")

                # Установка одностороннего режима для Pantum
                if is_pantum:
                    logging.info("Установка одностороннего режима для Pantum перед печатью последней страницы...")
                    success, original_duplex = set_printer_duplex_mode(printer_name, DMDUP_SIMPLEX)
                    if not success:
                        logging.warning("Не удалось установить односторонний режим через win32print, пробуем PowerShell...")
                        set_printer_duplex_powershell(printer_name, False)
                    time.sleep(1)

                # Шаг 1: Печать последней страницы с обычными полями
                doc.PageSetup.MirrorMargins = False
                last_page = doc.ComputeStatistics(2)
                logging.info(f"2. Количество страниц: {last_page}")
                logging.info(f"3. Печать последней страницы (страница {last_page})...")
                doc.PrintOut(PrintToFile=False, Range=win32.constants.wdPrintRangeOfPages, Pages=str(last_page))

                # Шаг 2: Изменение цвета шрифта колонтитулов первой страницы на черный
                logging.info("4. Изменение цвета шрифта колонтитулов первой страницы на черный...")
                try:
                    first_section = doc.Sections(1)
                    try:
                        first_page_footer = first_section.Footers(win32.constants.wdHeaderFooterFirstPage)
                        if first_page_footer.Exists:
                            first_page_footer.Range.Font.Color = win32.constants.wdColorBlack
                            logging.info("Цвет шрифта нижнего колонтитула первой страницы изменен на черный")
                        else:
                            primary_footer = first_section.Footers(win32.constants.wdHeaderFooterPrimary)
                            if primary_footer.Exists:
                                primary_footer.Range.Font.Color = win32.constants.wdColorBlack
                                logging.info("Цвет шрифта основного нижнего колонтитула изменен на черный")
                    except Exception as e:
                        logging.warning(f"Ошибка при изменении цвета нижнего колонтитула: {e}")
                except Exception as e:
                    logging.error(f"Ошибка при изменении цвета колонтитулов: {e}")

                # Шаг 3: Включаем зеркальные поля для печати всего документа
                logging.info("5. Включение зеркальных полей для печати всего документа...")
                doc.PageSetup.MirrorMargins = True

                # Установка двустороннего режима для Pantum
                if is_pantum:
                    logging.info("Установка двустороннего режима для Pantum перед печатью всего документа...")
                    success, _ = set_printer_duplex_mode(printer_name, DMDUP_VERTICAL)
                    if not success:
                        logging.warning("Не удалось установить двусторонний режим через win32print, пробуем PowerShell...")
                        set_printer_duplex_powershell(printer_name, True)
                    time.sleep(1)

                # Шаг 4: Печатаем весь документ в одном экземпляре
                last_page = doc.ComputeStatistics(2)
                logging.info(f"6. Количество страниц перед печатью всего документа: {last_page}")
                logging.info("7. Печать всего документа в одном экземпляре...")
                doc.PrintOut(PrintToFile=False, Range=win32.constants.wdPrintRangeOfPages, Pages=f"1-{last_page}", Copies=1, Collate=True)

                logging.info("8. Упрощенная печать для типа ЭКЦ завершена.")
            else:
                # СТАНДАРТНАЯ ЛОГИКА
                logging.info(f"Применяется стандартная логика печати. Тип шаблона: '{template_value}'")

                # Установка одностороннего режима для Pantum
                if is_pantum:
                    logging.info("Установка одностороннего режима для Pantum перед печатью последней страницы...")
                    success, original_duplex = set_printer_duplex_mode(printer_name, DMDUP_SIMPLEX)
                    if not success:
                        logging.warning("Не удалось установить односторонний режим через win32print, пробуем PowerShell...")
                        set_printer_duplex_powershell(printer_name, False)
                    time.sleep(1)

                # Шаг 1: Печать последней страницы с обычными полями
                logging.info("1. Отключение зеркальных полей для последней страницы...")
                doc.PageSetup.MirrorMargins = False
                last_page = doc.ComputeStatistics(2)
                logging.info(f"2. Количество страниц до печати последней страницы: {last_page}")
                logging.info(f"3. Печать последней страницы (страница {last_page})...")
                doc.PrintOut(PrintToFile=False, Range=win32.constants.wdPrintRangeOfPages, Pages=str(last_page))

                # Шаг 2: Удаляем текст после "Расписка о получении заключения"
                logging.info("4. Удаление текста после 'Расписка о получении заключения'...")
                for paragraph in doc.Paragraphs:
                    if "Расписка о получении заключения" in paragraph.Range.Text:
                        paragraph_range = paragraph.Range
                        paragraph_range.End = doc.Content.End
                        paragraph_range.Delete()
                        break

                # Шаг 3: Печатаем последнюю страницу после удаления с обычными полями
                last_page = doc.ComputeStatistics(2)
                logging.info(f"5. Количество страниц после удаления текста: {last_page}")
                logging.info(f"6. Печать последней страницы после удаления (страница {last_page})...")
                doc.PrintOut(PrintToFile=False, Range=win32.constants.wdPrintRangeOfPages, Pages=str(last_page))

                # Шаг 4: Удаляем последний раздел полностью с сохранением форматирования
                logging.info("7. Удаление последнего раздела...")
                sections = doc.Sections
                if sections.Count > 1:
                    # Определяем координаты последнего раздела
                    last_section = sections(sections.Count)
                    last_section_start = last_section.Range.Start
                    
                    # Создаем диапазон от начала последнего раздела до конца документа
                    last_section_range = doc.Range(last_section_start, doc.Content.End)
                    last_section_range.Delete()

                    # Дополнительная очистка - находим и удаляем маркеры разрыва разделов в конце
                    end_of_doc = doc.Content.End - 1
                    cleanup_range = doc.Range(max(0, end_of_doc - 20), end_of_doc + 1)
                    
                    # Заменяем разрывы разделов в конце документа на ничто
                    cleanup_range.Find.ClearFormatting()
                    cleanup_range.Find.Replacement.ClearFormatting()
                    cleanup_range.Find.Text = "^b"
                    cleanup_range.Find.Replacement.Text = ""
                    cleanup_range.Find.Execute(Replace=win32.constants.wdReplaceAll)

                    if doc.Paragraphs.Last.Range.Text.strip() == "":
                        doc.Paragraphs.Last.Range.Delete()

                doc.Repaginate()
                last_page = doc.ComputeStatistics(2)
                logging.info(f"8. Количество страниц после удаления раздела: {last_page}")

                # Шаг 5: Изменение цвета шрифта колонтитулов первой страницы на черный
                logging.info("9. Изменение цвета шрифта колонтитулов первой страницы на черный...")
                try:
                    # Получаем первый раздел документа
                    first_section = doc.Sections(1)
                    
                    # Изменяем цвет шрифта в верхнем колонтитуле первой страницы
                    #try:
                    #    first_page_header = first_section.Headers(win32.constants.wdHeaderFooterFirstPage)
                    #    if first_page_header.Exists:
                    #        first_page_header.Range.Font.Color = win32.constants.wdColorBlack
                    #        logging.info("Цвет шрифта верхнего колонтитула первой страницы изменен на черный")
                    #    else:
                    #        # Если нет специального колонтитула для первой страницы, используем основной
                    #        primary_header = first_section.Headers(win32.constants.wdHeaderFooterPrimary)
                    #        if primary_header.Exists:
                    #            primary_header.Range.Font.Color = win32.constants.wdColorBlack
                    #            logging.info("Цвет шрифта основного верхнего колонтитула изменен на черный")
                    #except Exception as e:
                    #    logging.warning(f"Ошибка при изменении цвета верхнего колонтитула: {e}")

                    # Изменяем цвет шрифта в нижнем колонтитуле первой страницы
                    try:
                        first_page_footer = first_section.Footers(win32.constants.wdHeaderFooterFirstPage)
                        if first_page_footer.Exists:
                            first_page_footer.Range.Font.Color = win32.constants.wdColorBlack
                            logging.info("Цвет шрифта нижнего колонтитула первой страницы изменен на черный")
                        else:
                            primary_footer = first_section.Footers(win32.constants.wdHeaderFooterPrimary)
                            if primary_footer.Exists:
                                primary_footer.Range.Font.Color = win32.constants.wdColorBlack
                                logging.info("Цвет шрифта основного нижнего колонтитула изменен на черный")
                    except Exception as e:
                        logging.warning(f"Ошибка при изменении цвета нижнего колонтитула: {e}")

                except Exception as e:
                    logging.error(f"Ошибка при изменении цвета колонтитулов: {e}")

                # Шаг 6: Включаем зеркальные поля для печати всего документа
                logging.info("10. Включение зеркальных полей для печати всего документа...")
                doc.PageSetup.MirrorMargins = True

                # Установка двустороннего режима для Pantum
                if is_pantum:
                    logging.info("Установка двустороннего режима для Pantum перед печатью всего документа...")
                    success, _ = set_printer_duplex_mode(printer_name, DMDUP_VERTICAL)
                    if not success:
                        logging.warning("Не удалось установить двусторонний режим через win32print, пробуем PowerShell...")
                        set_printer_duplex_powershell(printer_name, True)
                    time.sleep(1)

                # Шаг 7: Печатаем документ в двух экземплярах
                last_page = doc.ComputeStatistics(2)
                logging.info(f"11. Количество страниц перед печатью всего документа: {last_page}")
                logging.info("12. Печать документа в двух экземплярах без последней страницы...")
                doc.PrintOut(PrintToFile=False, Range=win32.constants.wdPrintRangeOfPages, Pages=f"1-{last_page}", Copies=2, Collate=True)

                # Шаг 8: Сохраняем изменённый документ в новом файле
                #original_path = Path(zakl_path)
                # Создаем новое имя файла с суффиксом "_modified"
                #modified_path = original_path.with_name(original_path.stem + "_modified" + original_path.suffix)
                #logging.info(f"8. Сохранение изменённого документа в {modified_path}...")
                #doc.SaveAs(str(modified_path))

        except Exception as e:
            logging.error(f"Произошла ошибка при печати документа: {e}")
            raise
        finally:
            # Восстановление исходного режима черновика
            if original_print_draft is not None:
                try:
                    word.Options.PrintDraft = original_print_draft
                    logging.info(f"Восстановлен исходный режим черновика: {original_print_draft}")
                except Exception as e:
                    logging.warning(f"Не удалось восстановить Options.PrintDraft: {e}")

            # Восстановление оригинального режима принтера (для Pantum)
            if is_pantum and original_duplex is not None:
                logging.info("Восстановление оригинального режима принтера...")
                restore_printer_duplex_mode(printer_name, original_duplex)

            # Закрываем документ без сохранения изменений
            doc.Close(SaveChanges=False)
            word.Quit()
            logging.info("13. Документ закрыт, Word завершён.")
    finally:
        pythoncom.CoUninitialize()

def create_and_fill_document(card_data, txt_data_for_word, template_name):
    """Создает и заполняет документ Word, сохраняя его во временную папку (код 4)"""
    try:
        template_folder = Path(card_data["3"])  # Путь к папке с шаблонами
        template_path = template_folder / f"{template_name}.docx"
        logging.info(f"Путь к папке с шаблонами: {template_folder}")
        logging.info(f"Имя файла шаблона: {template_name}.docx")

        if not template_path.exists():
            logging.error(f"Файл шаблона '{template_name}' не найден в папке {template_folder}.")
            raise FileNotFoundError(f"Файл шаблона '{template_name}' не найден в папке {template_folder}.")

        # Сохраняем во временную папку (код 4)
        temp_folder = Path(card_data["4"])
        temp_folder.mkdir(parents=True, exist_ok=True)
        logging.info(f"Временная папка для сохранения: {temp_folder}")

        logging.debug(f"Ключи в card_data перед заменой: {card_data.keys()}")
        logging.debug(f"Значения в card_data перед заменой: {card_data}")

        # Определение имени файла для сохранения
        if card_data.get("1") in ["СВО_Молов_образец_родственники", "СВО_Молов_образец_прямая идентификация", "СВО_Ростов_образец_родственники", "СВО_Ростов_образец_прямая идентификация"]:
            # Безопасное извлечение и преобразование значений
            fi = str(card_data.get('ФИ', '')).strip() if card_data.get('ФИ') is not None else ''
            ii = str(card_data.get('ИИ', '')).strip() if card_data.get('ИИ') is not None else ''
            oi = str(card_data.get('ОИ', '')).strip() if card_data.get('ОИ') is not None else ''
            nom = str(card_data.get('НОМ', '')).strip() if card_data.get('НОМ') is not None else ''
            
            # Формируем ФИО строку, исключая пустые части
            fio_parts = [part for part in [fi, ii, oi] if part and part != 'None']
            fio_string = ' '.join(fio_parts)
            
            # Формируем имя файла
            if fio_string:
                output_filename = f"{fio_string} {nom}-26.docx"
            else:
                output_filename = f"{nom}-26.docx"
                
        elif card_data.get("1") == "СВО_кость":
            ob_value = str(card_data.get('ОБ', '')).strip() if card_data.get('ОБ') is not None else ''
            tr_value = str(card_data.get('ТР', '')).strip() if card_data.get('ТР') is not None else ''
            nom_value = str(card_data.get('НОМ', '')).strip() if card_data.get('НОМ') is not None else ''
            
            ob_last_char = ob_value[-1] if ob_value else ''
            output_filename = f"{tr_value} {nom_value}-26 {ob_last_char}.docx"
        else:
            # Безопасное извлечение и преобразование значений
            fi = str(card_data.get('ФИ', '')).strip() if card_data.get('ФИ') is not None else ''
            ii = str(card_data.get('ИИ', '')).strip() if card_data.get('ИИ') is not None else ''
            oi = str(card_data.get('ОИ', '')).strip() if card_data.get('ОИ') is not None else ''
            nom = str(card_data.get('НОМ', '')).strip() if card_data.get('НОМ') is not None else ''
            
            # Формируем ФИО строку, исключая пустые части
            fio_parts = [part for part in [fi, ii, oi] if part and part != 'None']
            fio_string = ' '.join(fio_parts)
            
            # Формируем имя файла
            if fio_string:
                output_filename = f"{fio_string} {nom}-26.docx"
            else:
                output_filename = f"{nom}-26.docx"

        temp_path = temp_folder / output_filename
        logging.info(f"Временный путь к выходному файлу: {temp_path}")

        try:
            shutil.copy(template_path, temp_path)
            logging.info(f"Шаблон скопирован во временный файл: {temp_path}")
        except Exception as e:
            logging.error(f"Ошибка при копировании шаблона: {e}")
            raise

        try:
            replacements = {code: str(value).strip() for code, value in card_data.items() if value is not None}
            replace_in_doc(temp_path, replacements)
            logging.info("Замены меток выполнены в документе")
        except Exception as e:
            logging.error(f"Ошибка при замене меток в документе: {e}")
            raise

        try:
            # ИЗМЕНЕНИЕ: передаём txt_data_for_word вместо txt_data
            insert_table_data(temp_path, txt_data_for_word, card_data)
            logging.info("Таблица данных вставлена в документ (с сохранением исходных значений f)")
        except Exception as e:
            logging.error(f"Ошибка при вставке таблицы данных: {e}")
            raise

        try:
            insert_images(temp_path, card_data["4"])
            logging.info("Изображения вставлены в документ")
        except Exception as e:
            logging.error(f"Ошибка при вставке изображений: {e}")
            raise

        return temp_path

    except Exception as e:
        logging.error(f"Ошибка при создании и заполнении документа: {e}")
        raise

def move_to_final_location(card_data, word_file_path, excel_file_path=None):
    """Перемещает файлы из временной папки в финальные папки после создания обоих файлов"""
    try:
        logging.info("Начало перемещения файлов в финальные папки")
        
        # Преобразуем пути в объекты Path, если они строками
        word_file_path = Path(word_file_path) if isinstance(word_file_path, str) else word_file_path
        if excel_file_path:
            excel_file_path = Path(excel_file_path) if isinstance(excel_file_path, str) else excel_file_path
        
        # Определяем финальную папку на основе оригинальной логики
        if card_data.get("1") == "СВО_Молов_образец_прямая идентификация":
            final_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Личные профили 2-ДНК (прямая идентификация)")
        elif card_data.get("1") == "СВО_Молов_образец_родственники":
            final_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Родственники с 2-ДНК Тюмень")
        elif card_data.get("1") == "СВО_Ростов_образец_прямая идентификация":
            final_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Личные профили 2-ДНК (прямая идентификация)")
        elif card_data.get("1") == "СВО_Ростов_образец_родственники":
            final_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Родственники с 2-ДНК Ростов")
        elif card_data.get("1") == "СВО_кость":
            final_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Кости с ИК-2")
        else:
            final_folder = Path(card_data["6"])
        
        logging.info(f"Определена финальная папка: {final_folder}")
        
        # Проверяем существование финальной папки
        if not final_folder.exists():
            error_msg = f"Финальная папка не найдена: {final_folder}"
            logging.error(error_msg)
            raise FileNotFoundError(error_msg)

        # Перемещаем Word-файл
        final_word_path = final_folder / word_file_path.name
        logging.info(f"Перемещение Word-файла: {word_file_path} -> {final_word_path}")
        shutil.move(str(word_file_path), str(final_word_path))
        logging.info(f"Word-файл успешно перемещен: {final_word_path}")

        # Перемещаем Excel-файл, если он существует
        if excel_file_path and Path(excel_file_path).exists():
            final_excel_path = final_folder / Path(excel_file_path).name
            logging.info(f"Перемещение Excel-файла: {excel_file_path} -> {final_excel_path}")
            shutil.move(str(excel_file_path), str(final_excel_path))
            logging.info(f"Excel-файл успешно перемещен: {final_excel_path}")
        else:
            logging.info("Excel-файл для перемещения не указан или не существует")
            final_excel_path = None

        return final_word_path, final_excel_path

    except Exception as e:
        logging.error(f"Ошибка при перемещении файлов в финальную папку: {e}")
        # В случае ошибки оставляем файлы во временной папке
        logging.warning("Файлы остаются во временной папке из-за ошибки перемещения")
        raise

# Функция для копирования разделов документа
def copy_sections(src_doc, dst_doc, sections):
    """
    Копирует указанные разделы из исходного документа в целевой
    :param src_doc: Исходный документ Word
    :param dst_doc: Целевой документ Word
    :param sections: Список кортежей (начало_текст, конец_текст, включить_конец)
    """
    for start_text, end_text, include_end in sections:
        # Находим начало раздела
        start_range = src_doc.Content
        start_range.Find.Execute(FindText=start_text)
        if not start_range.Find.Found:
            logging.warning(f"Начальный текст не найден: '{start_text}'")
            continue

        # Находим конец раздела
        end_range = start_range.Duplicate
        end_range.Collapse(Direction=constants.wdCollapseEnd)
        end_range.Find.Execute(FindText=end_text)
        if not end_range.Find.Found:
            logging.warning(f"Конечный текст не найден: '{end_text}'")
            continue

        # Корректируем диапазон
        if not include_end:
            end_range.MoveEnd(constants.wdCharacter, -len(end_text))

        # Копируем раздел
        section_range = src_doc.Range(start_range.Start, end_range.End)
        section_range.Copy()

        # Вставляем в новый документ
        dst_range = dst_doc.Range(dst_doc.Content.End-1)
        dst_range.PasteAndFormat(constants.wdFormatOriginalFormatting)
        logging.info(f"Скопирован раздел: '{start_text}' > '{end_text}'")

    # Удаляем последний пустой абзац
    if dst_doc.Paragraphs.Last.Range.Text == "\r":
        dst_doc.Paragraphs.Last.Range.Delete()

def archive_and_print_document(source_path, archive_folder, ask_before_print=True, card_data=None):
    """
    Архивирует документ и печатает его.
    Использует минимальные операции копирования.
    """
    try:
        # Преобразуем в Path, если нужно
        source_path = Path(source_path)
        archive_folder = Path(archive_folder)

        # Проверяем, существует ли исходный файл
        if not source_path.exists():
            raise FileNotFoundError(f"Исходный файл для архивации не найден: {source_path}")

        logging.info(f"Начало архивации и печати файла: {source_path.name}")
        
        # Добавляем логирование данных карты
        if card_data:
            logging.info(f"Данные карты переданы в archive_and_print_document. Тип шаблона: '{card_data.get('1', 'НЕТ')}'")
        else:
            logging.warning("Данные карты НЕ переданы в archive_and_print_document!")

        # === Часть печати ===
        if source_path.exists():
            if not ask_before_print or ask_for_print("заключение"):
                try:
                    # Передаем card_data в функцию print_document
                    print_document(str(source_path), card_data)
                    logging.info(f"Заключение отправлено на печать: {source_path}")
                except Exception as e:
                    logging.error(f"Ошибка при печати заключения: {e}")
                    raise
        else:
            logging.error(f"Документ для печати не найден: {source_path}")
            raise FileNotFoundError(f"Документ не найден: {source_path}")

        # === Часть архивации ===
        destination = archive_folder / source_path.name
        logging.info(f"Целевой путь для архивации: {destination}")

        # Проверяем, не заблокирован ли исходный файл
        for attempt in range(3):
            try:
                with open(source_path, 'rb') as f:
                    f.read(1)
                break
            except (PermissionError, OSError) as e:
                logging.warning(f"Файл заблокирован (попытка {attempt + 1}/3): {e}")
                time.sleep(1)
        else:
            raise RuntimeError("Файл остаётся заблокированным после нескольких попыток.")

        # Минимальное копирование без изменения атрибутов
        try:
            with open(source_path, 'rb') as src, open(destination, 'wb') as dst:
                while chunk := src.read(64 * 1024):  # 64KB порциями
                    dst.write(chunk)
            logging.info(f"Файл успешно скопирован в архив: {destination}")
        except Exception as e:
            logging.error(f"Критическая ошибка при копировании файла: {e}")
            raise

        # Копирование в дополнительную папку для ВСК
        if card_data is not None and card_data.get("1") in ["СВО_Молов_образец_родственники", "СВО_Молов_образец_прямая идентификация"]:
            # Определяем дополнительную папку
            additional_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\ВСК Тюмень\Архив Молов")
            
            # Проверяем существование дополнительной папки
            if not additional_folder.exists():
                error_msg = f"Дополнительная папка не существует: {additional_folder}"
                logging.error(error_msg)
                raise FileNotFoundError(error_msg)
            
            # Формируем путь для копирования
            additional_destination = additional_folder / source_path.name
            
            # Копируем файл обычным способом
            try:
                shutil.copy2(source_path, additional_destination)
                logging.info(f"Файл успешно скопирован в дополнительную папку: {additional_destination}")
            except Exception as e:
                logging.error(f"Ошибка при копировании в дополнительную папку: {e}")
                # Не прерываем выполнение программы, но логируем ошибку

    except Exception as e:
        logging.error(f"Ошибка при архивации или печати документа: {e}")
        raise

def find_answer_file(card_data):
    """Ищет файл 'Ответ на обращение заявителя' по ФИО из Карты"""
    try:
        # Формируем имя файла из данных Карты
        fio = f"{card_data['ФОР']} {card_data['ИОР']} {card_data['ООР']}"
        safe_fio = "".join(c if c.isalnum() or c in " _-" else "_" for c in fio)
        target_filename = f"Ответ на обращение заявителя - {safe_fio}.docx"
        
        # Путь к папке поиска
        answer_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\ВСК Тюмень\Сопроводительные\Распечатанные")
        
        # Проверяем существование папки
        if not answer_folder.exists():
            raise FileNotFoundError(f"Папка для поиска ответа не найдена: {answer_folder}")
        
        # Поиск файла
        answer_files = list(answer_folder.glob(f"*{fio}*"))
        
        # Логируем результаты поиска
        logging.info(f"Поиск файла '{target_filename}' в папке '{answer_folder}':")
        for i, file in enumerate(answer_files, 1):
            logging.info(f"{i}. Найден файл: {file.name}")
        
        # Обработка результатов
        if not answer_files:
            raise FileNotFoundError(f"Файл '{target_filename}' не найден в папке '{answer_folder}'")
        
        if len(answer_files) > 1:
            logging.warning("Найдено несколько подходящих файлов:")
            for i, file in enumerate(answer_files):
                logging.warning(f"{i+1}. {file.name}")
            
            # Пытаемся найти наиболее точное совпадение
            exact_match = next((f for f in answer_files if f.name == target_filename), None)
            if exact_match:
                logging.info(f"Найдено точное совпадение: {exact_match.name}")
                return exact_match
            
            # Запрашиваем выбор у пользователя
            while True:
                try:
                    choice = int(input("Выберите номер файла для печати: "))
                    if 1 <= choice <= len(answer_files):
                        selected_file = answer_files[choice-1]
                        logging.info(f"Выбран файл: {selected_file}")
                        return selected_file
                    else:
                        logging.warning("Неверный номер. Пожалуйста, выберите существующий вариант.")
                except ValueError:
                    logging.warning("Пожалуйста, введите число.")
        
        # Возвращаем единственный найденный файл
        logging.info(f"Найден файл: {answer_files[0]}")
        return answer_files[0]
    
    except Exception as e:
        logging.error(f"Ошибка при поиске файла ответа: {e}")
        raise

def print_with_table_removal(source_path, card_data):
    """
    Печатает документ в 3 экземплярах с таблицей, затем без таблицы (1 экз.)
    Использует временную копию для удаления таблицы
    """
    try:
        source_path = Path(source_path)
        temp_dir = Path(card_data["4"])
        temp_file = temp_dir / f"{source_path.stem}_temp{source_path.suffix}"
        
        logging.info(f"1. Создание временной копии документа: {temp_file}")
        shutil.copy2(source_path, temp_file)
        
        # Печать оригинала 3 экземпляра
        logging.info("2. Печать документа с таблицей (3 экземпляра)")
        print_simple_copy(source_path, 3)
        
        # Удаление таблицы из временного файла
        logging.info("3. Подготовка документа без таблицы")
        pythoncom.CoInitialize()
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False
        
        try:
            doc = word.Documents.Open(str(temp_file))
            # Удаляем первую таблицу (предполагаем, что таблица одна)
            if doc.Tables.Count > 0:
                doc.Tables(1).Delete()
                logging.info("Таблица успешно удалена из временного документа")
            else:
                logging.warning("В документе не найдено таблиц для удаления")
            
            # Сохраняем изменения
            doc.Save()
            doc.Close(SaveChanges=True)
            
            # Печать документа без таблицы
            logging.info("4. Печать документа без таблицы (1 экземпляр)")
            print_simple_copy(temp_file, 1)
            
        except Exception as e:
            logging.error(f"Ошибка при работе с Word: {e}")
            raise
        finally:
            try:
                word.Quit()
            finally:
                pythoncom.CoUninitialize()
        
        # Очистка временных файлов
        try:
            if temp_file.exists():
                temp_file.unlink()
                logging.info("6. Временный файл успешно удален")
        except Exception as e:
            logging.warning(f"Не удалось удалить временный файл: {e}")
            
    except Exception as e:
        logging.error(f"Ошибка при печати документа в разных вариантах: {e}")
        raise

# Основная функция
def main():
    try:
        # Вывод текущей рабочей директории
        current_directory = os.getcwd()
        logging.info(f"Текущая рабочая директория: {current_directory}")

        # Находим файл Excel
        try:
            card_path = find_excel_file()
            logging.info(f"Найден файл Карта: {card_path}")
            print(f"Найден файл Карта: {card_path}")
        except Exception as e:
            logging.error(f"Ошибка при поиске файла Excel: {e}")
            raise

        # Первичное чтение данных из найденного файла Excel
        try:
            card_data = read_card(card_path)
            print("Данные из файла Карта:", card_data) 
        except Exception as e:
            logging.error(f"Ошибка при чтении данных из файла Excel: {e}")
            raise

        # После чтения карты сохраняем номер
        nom_value = card_data.get('НОМ', 'N/A')

        # Проверяем, что все обязательные поля заполнены
        try:
            check_card_data(card_data)
        except Exception as e:
            logging.error(f"Ошибка при проверке обязательных полей: {e}")
            raise

        # Показываем предупреждение о ФИО, если нужно
        warning_was_shown, user_clicked_ok = show_fio_warning(card_data)

        # Перечитываем данные из Карты если было показано предупреждение
        # (независимо от того, нажал пользователь OK или отклонил)
        if warning_was_shown:
            logging.info("Предупреждение было показано. Перечитываем данные из Карты.")
            
            try:
                # Добавим небольшую задержку, чтобы дать пользователю время сохранить файл
                time.sleep(1)
                
                # Повторно проверяем существование файла перед чтением
                if not Path(card_path).exists():
                    raise FileNotFoundError(f"Файл Карты не найден по пути: {card_path}")
                
                # Перечитываем данные из Карты
                card_data = read_card(card_path)
                logging.info("Данные из Карты успешно перечитаны после предупреждения.")
                
                # Повторно проверяем обязательные поля после перечитывания
                check_card_data(card_data)
                
            except Exception as e:
                logging.error(f"Ошибка при повторном чтении Карты: {e}")
                show_error_message(f"Ошибка при обновлении данных из Карты: {e}")
        else:
            logging.info("Предупреждение не было показано. Карта не перечитывается.")

        # Проверяем особые случаи
        is_special_case = card_data.get("1") == "СВО_кость_нет результата"
        
        is_no_result_case = card_data.get("1") in [
            "СВО_Молов_образец_родственники_нет результата_RT", "СВО_Молов_образец_родственники_нет результата_форез",
            "СВО_Ростов_образец_прямая идентификация_нет результата_RT", "СВО_Ростов_образец_прямая идентификация_нет результата_форез",
            "СВО_Ростов_образец_родственники_нет результата_RT", "СВО_Ростов_образец_родственники_нет результата_форез",
        ]
        
        # Проверка на шаблоны ЭКЦ
        is_ekc_case = card_data.get("1") == "ЭКЦ" or (card_data.get("1") and card_data.get("1").startswith("ЭКЦ_образец_") and not card_data.get("1").endswith("_нет результата"))
        
        # Проверка на шаблоны ЭКЦ, где нет результата
        is_ekc_no_result_case = card_data.get("1") == "ЭКЦ_нет результата" or (card_data.get("1") and card_data.get("1").startswith("ЭКЦ_образец_") and card_data.get("1").endswith("_нет результата"))

        if is_ekc_case:
            logging.info(f"Обнаружен особый тип файла ЭКЦ: '{card_data.get('1')}'. Выполняется обработка.")
            try:
                # Пути к папкам
                base_working_materials_folder = Path(card_data["4"])
                nom_folder_name = f"{card_data['НОМ']}-26"
                working_materials_folder = base_working_materials_folder / nom_folder_name
                txt_path = find_txt_file(working_materials_folder)
                
                # Читаем данные и фильтруем (с заменой f) для XML
                txt_data, loci_with_f = read_and_filter_txt_data(txt_path)
                
                # Создаём временный полный txt-файл с заменёнными f для Excel
                temp_txt_for_excel = os.path.join(card_data["4"], f"temp_ekc_excel_{os.path.basename(txt_path)}")
                replace_f_in_txt_file(txt_path, temp_txt_for_excel)
                
                # Создаём временный файл с отфильтрованными данными для XML
                temp_txt_for_xml = os.path.join(card_data["4"], f"temp_ekc_xml_{os.path.basename(txt_path)}")
                with open(txt_path, 'r', encoding='utf-8') as f:
                    header_line = f.readline()
                with open(temp_txt_for_xml, 'w', encoding='utf-8') as f:
                    f.write(header_line)
                    for row in txt_data:
                        f.write('\t'.join(row) + '\n')
                
                # Генерируем XML
                xml_output_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\BASE")
                xml_path = xml_output_folder / f"{card_data['НОМ']}-26.xml"
                convert_txt_to_xml(temp_txt_for_xml, xml_path, object_name_suffix=txt_data[0][0])
                logging.info(f"XML файл создан: {xml_path}")

                # Заполняем Excel файл ИКЛ ЭКЦ, используя временный txt с заменёнными f
                vba_path = find_vba_file(card_data)
                ekc_output_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\ЭКЦ ИКЛ")
                excel_file_path = run_vba_macro_2_DNK(vba_path, temp_txt_for_excel, ekc_output_folder, card_data, sheet_name="ИКЛ", loci_with_f=loci_with_f)
                logging.info(f"Excel файл создан: {excel_file_path}")

                # Архивируем и печатаем оригинальный Word файл
                source_folder = Path(card_data["2"])
                source_file = source_folder / f"{card_data['НОМ']}-26.docx"
                archive_folder = Path(card_data["5"])
                archive_and_print_document(source_file, archive_folder, card_data=card_data)

                # Удаляем временные файлы
                for f in [temp_txt_for_excel, temp_txt_for_xml]:
                    try:
                        os.remove(f)
                        logging.debug(f"Временный файл {f} удалён")
                    except Exception as e:
                        logging.warning(f"Не удалось удалить {f}: {e}")

                sys.exit(0)
            except Exception as e:
                logging.error(f"Ошибка при обработке особого случая ЭКЦ: {e}")
                raise

        if is_ekc_no_result_case:
            logging.info(f"Обнаружен особый тип файла ЭКЦ (нет результата): '{card_data.get('1')}'. Выполняется обработка.")
            try:
                
                # Вместо Excel создаем пустой TXT файл
                # Файл должен попадать в ту же папку, куда попадает Excel в случае ekc_case
                ekc_output_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\ЭКЦ ИКЛ")
                ekc_output_folder.mkdir(parents=True, exist_ok=True)
                
                # Формируем имя файла как для Excel ИКЛ ЭКЦ, но с расширением .txt
                excel_filename = f"{card_data['ФИ']} {card_data['ИИ']} {card_data['ОИ']} {card_data['НОМ']}-26 нет профиля.xlsm"
                txt_filename = excel_filename.replace('.xlsm', '.txt')
                txt_file_path = ekc_output_folder / txt_filename
                
                # Создаем пустой TXT файл
                with open(txt_file_path, 'w', encoding='utf-8') as f:
                    f.write("")  # Пустой файл
                
                logging.info(f"Создан пустой TXT файл вместо Excel: {txt_file_path}")
                logging.info(f"Файл сохранен в папку: {ekc_output_folder}")

                # Архивируем и печатаем оригинальный Word файл
                source_folder = Path(card_data["2"])
                source_file = source_folder / f"{card_data['НОМ']}-26.docx"
                archive_folder = Path(card_data["5"])
                archive_and_print_document(source_file, archive_folder, card_data=card_data)

                sys.exit(0)
            except Exception as e:
                logging.error(f"Ошибка при обработке особого случая ЭКЦ_нет результата: {e}")
                raise

        if is_special_case:
            logging.info("Обнаружен особый тип файла: 'СВО_кость_нет результата'. Выполняется обработка документа.")
            try:
                # Формируем пути к файлам
                source_folder = Path(card_data["2"])
                source_file = source_folder / f"{card_data['НОМ']}-26.docx"
                ob_last_char = card_data['ОБ'].strip()[-1] if card_data['ОБ'].strip() else ''
                target_file = source_folder / f"{card_data['ТР']} {card_data['НОМ']}-26 {ob_last_char}.docx"
                archive_folder = Path(card_data["5"])
                logging.info(f"Исходный файл: {source_file}")
                logging.info(f"Целевой файл: {target_file}")

                # Проверяем существование исходного файла
                if not source_file.exists():
                    raise FileNotFoundError(f"Исходный файл не найден: {source_file}")

                # Добавляем задержку перед открытием файла
                time.sleep(1)

                # Инициализируем Word
                pythoncom.CoInitialize()
                word = win32.gencache.EnsureDispatch('Word.Application')
                word.Visible = False

                try:
                    # Пытаемся открыть файл с несколькими попытками
                    max_attempts = 3
                    for attempt in range(max_attempts):
                        try:
                            src_doc = word.Documents.Open(str(source_file))
                            break
                        except Exception as e:
                            if attempt == max_attempts - 1:
                                raise
                            time.sleep(1)
                            continue

                    # Создаём новый документ
                    dst_doc = word.Documents.Add()

                    # Определяем разделы для копирования
                    sections_to_copy = [
                        ("Описание вещественных доказательств", "Вещественные доказательства", False),
                        ("РЕЗУЛЬТАТЫ", "Государственный судебно-медицинский эксперт", False),
                        ("ВЫВОДЫ", "Государственный судебно-медицинский эксперт", False)
                    ]

                    # Используем существующую функцию copy_sections
                    copy_sections(src_doc, dst_doc, sections_to_copy)

                    # Сохраняем результирующий файл
                    for attempt in range(max_attempts):
                        try:
                            dst_doc.SaveAs(str(target_file))
                            logging.info(f"Документ успешно сохранён: {target_file}")
                            break
                        except Exception as e:
                            if attempt == max_attempts - 1:
                                raise
                            time.sleep(1)
                            continue

                    # Закрываем документы
                    src_doc.Close(SaveChanges=constants.wdDoNotSaveChanges)
                    dst_doc.Close(SaveChanges=constants.wdDoNotSaveChanges)

                    # Перемещаем файл в финальную папку
                    final_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Кости с ИК-2")
                    final_folder.mkdir(parents=True, exist_ok=True)
                    for attempt in range(max_attempts):
                        try:
                            shutil.move(str(target_file), str(final_folder / target_file.name))
                            logging.info(f"Файл перемещён в финальную папку: {final_folder / target_file.name}")
                            break
                        except PermissionError as e:
                            if attempt == max_attempts - 1:
                                raise
                            time.sleep(1)
                            continue

                except Exception as e:
                    logging.error(f"Ошибка при обработке документа: {e}")
                    raise
                finally:
                    # Гарантированно закрываем документы и Word
                    try:
                        if 'src_doc' in locals():
                            src_doc.Close(SaveChanges=constants.wdDoNotSaveChanges)
                    except:
                        pass
                    try:
                        if 'dst_doc' in locals():
                            dst_doc.Close(SaveChanges=constants.wdDoNotSaveChanges)
                    except:
                        pass
                    try:
                        word.Quit()
                    except:
                        pass
                    pythoncom.CoUninitialize()

                # Архивируем и печатаем исходный файл
                archive_and_print_document(source_file, archive_folder, card_data=card_data)
                sys.exit(0)  # Завершаем программу
            except Exception as e:
                logging.error(f"Ошибка при обработке особого случая: {e}")
                raise

        elif is_no_result_case:
            logging.info(f"Обнаружен особый тип файла: '{card_data.get('1')}'. Выполняется создание результирующего файла.")
            try:
                source_folder = Path(card_data["2"])
                source_file = source_folder / f"{card_data['НОМ']}-26.docx"
                
                # Создаём результирующий файл
                pythoncom.CoInitialize()
                word = win32.gencache.EnsureDispatch('Word.Application')
                word.Visible = False
                
                try:
                    # Пытаемся открыть файл с несколькими попытками
                    max_attempts = 3
                    for attempt in range(max_attempts):
                        try:
                            src_doc = word.Documents.Open(str(source_file))
                            break
                        except Exception as e:
                            if attempt == max_attempts - 1:
                                raise
                            time.sleep(1)
                            continue

                    # Создаём новый документ
                    dst_doc = word.Documents.Add()

                    # Определяем разделы для копирования
                    sections_to_copy = [
                        ("В филиале Государственного фонда", "ИССЛЕДОВАТЕЛЬСКАЯ ЧАСТЬ", False),
                        ("РЕЗУЛЬТАТЫ", "Государственный судебно-медицинский эксперт", False),
                        ("В Ы В О Д Ы", "Государственный судебно-медицинский эксперт", False)
                    ]

                    # Копируем разделы
                    copy_sections(src_doc, dst_doc, sections_to_copy)

                    # Сохраняем результирующий файл
                    result_file_name = f"{card_data['ФИ']} {card_data['ИИ']} {card_data['ОИ']} {card_data['НОМ']}-26.docx"
                    result_file_path = source_folder / result_file_name
                    
                    for attempt in range(max_attempts):
                        try:
                            dst_doc.SaveAs(str(result_file_path))
                            logging.info(f"Документ успешно сохранен: {result_file_path}")
                            break
                        except Exception as e:
                            if attempt == max_attempts - 1:
                                raise
                            time.sleep(1)
                            continue

                    # Закрываем документы
                    src_doc.Close(SaveChanges=constants.wdDoNotSaveChanges)
                    dst_doc.Close(SaveChanges=constants.wdDoNotSaveChanges)

                    # Определяем финальную папку в зависимости от шаблона
                    template = card_data.get("1")
                    if template in ["СВО_Ростов_образец_прямая идентификация_нет результата_RT", "СВО_Ростов_образец_прямая идентификация_нет результата_форез"]:
                        final_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Личные профили 2-ДНК (прямая идентификация)")
                        logging.info(f"Обнаружен шаблон прямой идентификации (Ростов). Файл будет направлен в: {final_folder}")
                    elif template in ["СВО_Ростов_образец_родственники_нет результата_RT", "СВО_Ростов_образец_родственники_нет результата_форез"]:
                        final_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Родственники с 2-ДНК Ростов")
                        logging.info(f"Обнаружен шаблон родственников (Ростов). Файл будет направлен в: {final_folder}")
                    elif template in ["СВО_Молов_образец_родственники_нет результата_RT", "СВО_Молов_образец_родственники_нет результата_форез"]:
                        final_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Родственники с 2-ДНК Тюмень")
                        logging.info(f"Обнаружен шаблон родственников (Молов/Тюмень). Файл будет направлен в: {final_folder}")
                    else:
                        # Запасной вариант, если шаблон не распознан
                        final_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО")
                        logging.warning(f"Шаблон '{template}' не распознан. Файл будет направлен в папку по умолчанию: {final_folder}")
                    
                    # Создаём финальную папку, если она не существует
                    final_folder.mkdir(parents=True, exist_ok=True)
                    logging.info(f"Финальная папка готова: {final_folder}")
                    
                    # Перемещаем результирующий файл
                    for attempt in range(max_attempts):
                        try:
                            shutil.move(str(result_file_path), str(final_folder / result_file_name))
                            logging.info(f"Файл перемещен в финальную папку: {final_folder / result_file_name}")
                            break
                        except PermissionError as e:
                            if attempt == max_attempts - 1:
                                raise
                            time.sleep(1)
                            continue

                    # Архивируем и печатаем исходный файл
                    archive_folder = Path(card_data["5"])
                    archive_and_print_document(source_file, archive_folder, card_data=card_data)
                    sys.exit(0)  # Завершаем программу

                except Exception as e:
                    logging.error(f"Ошибка при обработке документа: {e}")
                    raise
                finally:
                    # Гарантированно закрываем документы и Word
                    try:
                        if 'src_doc' in locals():
                            src_doc.Close(SaveChanges=constants.wdDoNotSaveChanges)
                    except:
                        pass
                    try:
                        if 'dst_doc' in locals():
                            dst_doc.Close(SaveChanges=constants.wdDoNotSaveChanges)
                    except:
                        pass
                    try:
                        word.Quit()
                    except:
                        pass
                    pythoncom.CoUninitialize()

            except Exception as e:
                logging.error(f"Ошибка при обработке особого случая: {e}")
                raise

        # Блок для Тобольск_образец_
        is_tobolsk_case = card_data.get("1") and card_data.get("1").startswith("Тобольск_образец_")
        if is_tobolsk_case:
            logging.info(f"Обнаружен шаблон Тобольск_образец: '{card_data.get('1')}'. Выполняется обработка (XML, печать, архив).")
            try:
                # Пути к папкам
                base_working_materials_folder = Path(card_data["4"])
                nom_folder_name = f"{card_data['НОМ']}-26"
                working_materials_folder = base_working_materials_folder / nom_folder_name
                txt_path = find_txt_file(working_materials_folder)
                
                # Читаем данные и фильтруем (с заменой f) для XML
                txt_data, _ = read_and_filter_txt_data(txt_path)
                if not txt_data:
                    raise ValueError("Нет данных для выбранного образца")
                
                # Создаём временный txt-файл с отфильтрованными данными для XML
                temp_txt_for_xml = os.path.join(card_data["4"], f"temp_tobolsk_xml_{os.path.basename(txt_path)}")
                with open(txt_path, 'r', encoding='utf-8') as f:
                    header_line = f.readline()
                with open(temp_txt_for_xml, 'w', encoding='utf-8') as f:
                    f.write(header_line)
                    for row in txt_data:
                        f.write('\t'.join(row) + '\n')
                
                # Генерируем XML
                xml_output_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\BASE")
                xml_output_folder.mkdir(parents=True, exist_ok=True)
                xml_path = xml_output_folder / f"{card_data['НОМ']}-26.xml"
                convert_txt_to_xml(temp_txt_for_xml, xml_path, object_name_suffix=txt_data[0][0])
                logging.info(f"XML файл создан: {xml_path}")
                
                # Удаляем временный файл
                try:
                    os.remove(temp_txt_for_xml)
                    logging.debug(f"Временный файл {temp_txt_for_xml} удалён")
                except Exception as e:
                    logging.warning(f"Не удалось удалить {temp_txt_for_xml}: {e}")
                
                # Архивируем и печатаем оригинальный Word файл
                source_folder = Path(card_data["2"])
                source_file = source_folder / f"{card_data['НОМ']}-26.docx"
                archive_folder = Path(card_data["5"])
                archive_and_print_document(source_file, archive_folder, card_data=card_data)
                
                sys.exit(0)  # Завершаем программу, чтобы не выполнять остальную логику
            except Exception as e:
                logging.error(f"Ошибка при обработке шаблона Тобольск_образец: {e}")
                raise

        else:

            # Находим файл txt
            try:
                base_working_materials_folder = Path(card_data["4"])
                nom_folder_name = f"{card_data['НОМ']}-26"
                working_materials_folder = base_working_materials_folder / nom_folder_name
                if not working_materials_folder.exists():
                    logging.error(f"Папка '{working_materials_folder}' не существует.")
                    raise FileNotFoundError(f"Папка '{working_materials_folder}' не существует.")
                txt_path = find_txt_file(working_materials_folder)
                
                # Читаем данные для всех документов (с заменой f для XML/Excel)
                txt_data, loci_with_f = read_and_filter_txt_data(txt_path)
                
                # Читаем данные специально для Word (без замены f)
                txt_data_for_word = read_and_filter_txt_data_for_word(txt_path)
                logging.info("Данные для Word получены: значения 'f' сохранены в исходном виде")
                logging.info("Данные для XML/Excel получены: значения 'f' заменены на повторяющиеся аллели")
                
                # Проверяем условие для создания "Ответ на обращение заявителя"
                if card_data.get("1") in ["СВО_Молов_образец_родственники", "СВО_Молов_образец_прямая идентификация"]:
                    logging.info("Обнаружено условие для создания 'Ответ на обращение заявителя'")
                    
                    # Создание документа "Ответ на обращение заявителя"
                    # Определяем шаблон в зависимости от значения в поле "1"
                    if card_data.get("1") == "СВО_Молов_образец_родственники":
                        template_name = "Ответ на обращение заявителя"
                        logging.info(f"Выбран шаблон для 'СВО_Молов_образец_родственники': {template_name}")
                    elif card_data.get("1") == "СВО_Молов_образец_прямая идентификация":
                        template_name = "Ответ на обращение заявителя прямая идентификация"
                        logging.info(f"Выбран шаблон для 'СВО_Молов_образец_прямая идентификация': {template_name}")

                    # Создание документа
                    answer_word_path = create_and_fill_document(card_data, txt_data_for_word, template_name)
                    logging.info(f"Документ '{template_name}' создан: {answer_word_path}")

                    # Перемещение в специальную папку
                    vsk_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\ВСК Тюмень\Сопроводительные\Распечатанные")
                    vsk_folder.mkdir(parents=True, exist_ok=True)
                    
                    # Формирование имени файла
                    fio = f"{card_data['ФОР']} {card_data['ИОР']} {card_data['ООР']}"
                    safe_fio = "".join(c if c.isalnum() or c in " _-" else "_" for c in fio)
                    answer_filename = f"Ответ на обращение заявителя - {safe_fio}.docx"
                    final_answer_path = vsk_folder / answer_filename
                    
                    # Перемещение файла
                    shutil.move(str(answer_word_path), str(final_answer_path))
                    logging.info(f"Документ перемещён в ВСК-папку: {final_answer_path}")

                    # Блок для создания и печати конверта
                    try:
                        logging.info("Подготовка к созданию и печати конверта")
                        envelope_path = create_and_fill_envelope(card_data)

                        # Проверяем, существует ли временный файл конверта
                        if envelope_path.exists():
                            logging.info(f"Конверт создан: {envelope_path}")
                            # Запрашиваем подтверждение печати
                            if ask_for_print("конверт"):
                                logging.info("Пользователь подтвердил печать конверта")
                                print_simple_copy(envelope_path, 1)
                            else:
                                logging.info("Пользователь отменил печать конверта")
                            # Удаляем временный файл
                            if envelope_path.exists():
                                envelope_path.unlink()
                                logging.info("Временный файл конверта успешно удален")
                        else:
                            logging.warning("Конверт не был создан, пропускаем этап печати")
                    except Exception as e:
                        logging.error(f"Ошибка при обработке конверта: {e}")
                        show_error_message(f"Ошибка при создании или печати конверта: {e}")

                    # Запрашиваем подтверждение пользователя и печатаем документ
                    if ask_for_print("сопроводительные листы"):
                        try:
                            answer_file_path = find_answer_file(card_data)
                            print_with_table_removal(answer_file_path, card_data)
                            logging.info("Сопроводительные листы успешно распечатаны")
                        except Exception as e:
                            logging.error(f"Ошибка при печати сопроводительных листов: {e}")
                            show_error_message(f"Ошибка при печати сопроводительных листов: {e}")
                    else:
                        logging.info("Пользователь отменил печать сопроводительных листов")

            except Exception as e:
                logging.error(f"Ошибка при поиске или чтении файла txt: {e}")
                raise
                
            # Определяем путь для сохранения файла XML
            try:
                xml_output_folder = Path(r"U:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\BASE")
                xml_path = xml_output_folder / f"{card_data['НОМ']}-26.xml"
                # Преобразование txt в xml
                convert_txt_to_xml(txt_path, xml_path, object_name_suffix=txt_data[0][0])
            except Exception as e:
                logging.error(f"Ошибка при преобразовании txt в xml: {e}")
                raise
                
            # Определяем шаблон в зависимости от типа
            try:
                if card_data.get("1") == "СВО_Молов_образец_прямая идентификация":
                    template_name = "Для отправки личный"
                elif card_data.get("1") in ["СВО_Молов_образец_родственники", "СВО_Ростов_образец_родственники"]:
                    template_name = "Для отправки родственники"
                elif card_data.get("1") == "СВО_Ростов_образец_прямая идентификация":
                    template_name = "Для отправки личный Ростов"
                elif card_data.get("1") == "СВО_кость":
                    template_name = "Для отправки кость"
                else:
                    template_name = "Для отправки родственники"
            except Exception as e:
                logging.error(f"Ошибка при определении шаблона документа: {e}")
                raise
                
            # Создаем Word-файл во временной папке
            try:
                word_file_path = create_and_fill_document(card_data, txt_data_for_word, template_name)

                print(f"Новый документ создан и заполнен: {word_file_path}")
            except Exception as e:
                logging.error(f"Ошибка при создании и заполнении документа: {e}")
                raise
                
            # Находим файл VBA
            try:
                vba_path = find_vba_file(card_data)
            except Exception as e:
                logging.error(f"Ошибка при поиске файла VBA: {e}")
                raise
                
            # Логирование меток перед вызовом replace_in_excel
            logging.info(f"Метки для замены: {card_data}")
            
            # Проверяем значение в графе "1" и вызываем соответствующую функцию
            value_in_column_1 = card_data.get("1", "").strip()

            # Определяем, относится ли значение к типу "Универсальная_таблица_для_генотипов"
            if value_in_column_1 in [
                "СВО_Молов_образец_родственники", 
                "СВО_Молов_образец_прямая идентификация", 
                "СВО_Ростов_образец_родственники", 
                "СВО_Ростов_образец_прямая идентификация", 
                "СВО_кость"
            ]:
                logging.info(f"Значение '1' в Карте '{value_in_column_1}' соответствует новому типу: Универсальная_таблица_для_генотипов")
                temp_txt_path = None
                try:
                    # Создание временного txt файла
                    # Передаем card_data["4"] как путь для сохранения временного файла
                    temp_txt_path = create_temp_txt_for_universal_table(txt_path, card_data["4"])
                    logging.info(f"Используется временный txt файл: {temp_txt_path}")

                    # Вызов макроса с временным файлом
                    excel_file_path = run_vba_macro_universal_table(vba_path, temp_txt_path, Path(card_data["4"]), card_data)
                    logging.info(f"Файл Excel успешно создан с использованием универсальной таблицы: {excel_file_path}")
                    
                except Exception as e:
                    logging.error(f"Ошибка при выполнении макроса VBA для универсальной таблицы: {e}")
                    raise
                finally:
                    # Удаление временного txt файла
                    if temp_txt_path and os.path.exists(temp_txt_path):
                        try:
                            os.remove(temp_txt_path)
                            logging.info(f"Временный txt файл удален: {temp_txt_path}")
                        except Exception as e:
                            logging.warning(f"Не удалось удалить временный txt файл {temp_txt_path}: {e}")

            # Все другие значения, включая "ЭКЦ", обрабатываются стандартным способом
            else:
                logging.info(f"Значение '1' в Карте '{value_in_column_1}' обрабатывается как 2-ДНК, так как остался только ЭКЦ.")
                try:
                    excel_file_path = run_vba_macro_2_DNK(vba_path, txt_path, Path(card_data["4"]), card_data, loci_with_f=loci_with_f)
                    logging.info(f"Файл Excel успешно создан стандартным способом: {excel_file_path}")
                except Exception as e:
                    logging.error(f"Ошибка при выполнении макроса VBA стандартным способом: {e}")
                    raise
                    
            # Получаем путь к исходному файлу из папки "код 2"
            source_folder = Path(card_data["2"])
            source_file = source_folder / f"{card_data['НОМ']}-26.docx"

            # Перемещаем оба файла в финальную папку
            try:
                final_word_path, final_excel_path = move_to_final_location(card_data, word_file_path, excel_file_path)
                logging.info(f"Файлы перемещены в финальные папки: Word={final_word_path}, Excel={final_excel_path}")
            except Exception as e:
                logging.error(f"Ошибка при перемещении файлов: {e}")
                raise

            # Печать и сохранение заключения в папку под кодом 5
            try:
                archive_folder = Path(card_data["5"])
                
                # Проверяем существование исходного файла перед архивацией
                if source_file.exists():
                    archive_and_print_document(source_file, archive_folder, card_data=card_data)
                else:
                    logging.warning(f"Исходный файл не найден для архивации: {source_file}")
            except Exception as e:
                logging.error(f"Ошибка при печати или сохранении документа: {e}")
                raise
                
            print("Скрипт завершён!")
            
    except Exception as e:
        logging.error(f"Произошла ошибка: {e}")
        show_error_message(f"Произошла ошибка: {e}")
        
    finally:
        try:
            show_final_reminder()
        except Exception as e:
            logging.error(f"Не удалось показать финальное напоминание: {e}")
        finally:
            # Дополнительная гарантированная очистка
            import gc
            gc.collect()
            
# Запуск скрипта
if __name__ == "__main__":
    main()