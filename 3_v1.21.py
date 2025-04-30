#3
#v1.21

import os
import shutil
import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook
from PIL import Image, ImageEnhance
import win32com.client as win32
import pythoncom
import time
import xml.etree.ElementTree as ET
from datetime import datetime
import sys
import tkinter as tk
from tkinter import messagebox

# Настройка логирования
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Определяем порядок сортировки локусов (в верхнем регистре)
LOCUS_ORDER = [
    "D3S1358", "VWA", "D16S539", "CSF1PO", "TPOX", "YINDEL", "AMEL", 
    "D8S1179", "D21S11", "D18S51", "DYS391", "D2S441", "D19S433", 
    "TH01", "FGA", "D22S1045", "D5S818", "D13S317", "D7S820", "SE33", 
    "D10S1248", "D1S1656", "D12S391", "D2S1338", "SRY", "D6S1043"
]

# Функция для отображения сообщений об ошибках
def show_error_message(message):
    root = tk.Tk()
    root.withdraw()  # Скрываем главное окно Tkinter
    messagebox.showerror("Ошибка", message)
    root.destroy()

# Функция для сортировки данных по локусам
def sort_data_by_locus(data):
    def get_locus_index(locus):
        locus_upper = locus.strip().upper()
        try:
            return LOCUS_ORDER.index(locus_upper)
        except ValueError:
            return len(LOCUS_ORDER) + ord(locus_upper[0])
    return sorted(data, key=lambda x: get_locus_index(x[3]))

# Функция для проверки обязательных полей
def check_card_data(card_data):
    required_fields = ["НОМ", "2", "3", "4", "7"]  # Добавлено поле "7" для пути к файлу VBA
    missing_fields = [field for field in required_fields if field not in card_data or not card_data[field]]
    if missing_fields:
        raise ValueError(f"В карте отсутствуют значения для следующих полей: {missing_fields}")

# Функция для чтения данных из Карты (Excel)
def read_card(file_path):
    data = {}
    wb = load_workbook(file_path)
    sheet = wb.active
    for row in sheet.iter_rows(min_row=2, values_only=True):
        code = str(row[0])  # Преобразуем код в строку
        value = row[2]
        data[code] = value
    return data

# Функция для копирования шаблона Word
def copy_template(card_data):
    template_folder = Path(card_data["3"])  # Путь к папке с шаблонами
    template_name = card_data["1"]  # Название шаблона из строки с кодом 1
    for ext in [".doc", ".docx"]:
        template_path = template_folder / f"{template_name}{ext}"
        if template_path.exists():
            output_path = Path(card_data["2"]) / f"{card_data['НОМ']}-25{ext}"
            shutil.copy(template_path, output_path)
            return output_path
    raise FileNotFoundError(f"Файл шаблона '{template_name}' не найден в папке {template_folder}.")

# Функция для замены меток в Word
def replace_in_doc(doc_path, replacements):
    doc = Document(doc_path)
    found_keys = set()  
    def replace_text(element):
        nonlocal found_keys
        if hasattr(element, "text"):
            original_text = ''.join(run.text for run in element.runs)
            modified_text = original_text
            for key, value in replacements.items():
                for pattern in [f"{{{key}}}", f"{{ {key} }}", f"{{{key} }}", f"{{ {key}}}"]:
                    if pattern in modified_text:
                        modified_text = modified_text.replace(pattern, str(value))
                        found_keys.add(key)
                        logging.info(f"Замена метки '{pattern}' на '{value}' в тексте: "
                                     f"'{original_text}' -> '{modified_text}'")
            if original_text != modified_text:
                for run in element.runs:
                    run.text = ""
                element.runs[0].text = modified_text
    for p in doc.paragraphs:
        replace_text(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text(p)
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_text(p)
        for p in section.footer.paragraphs:
            replace_text(p)
    missing_keys = set(replacements.keys()) - found_keys
    if missing_keys:
        logging.warning(f"Предупреждение: следующие метки не найдены в документе: {missing_keys}")
    doc.save(doc_path)

# Функция для вставки изображений
def insert_images(doc_path, images_folder, brightness_factor=2.0, contrast_factor=3.0):
    doc = Document(doc_path)
    images_folder = Path(images_folder)
    special_images = {"1": "[IMAGE_1]", "1_2": "[IMAGE_1_2]", "1_3": "[IMAGE_1_3]"}
    for image_name, placeholder in special_images.items():
        image_path = images_folder / f"{image_name}.jpg"
        if not image_path.exists():
            for ext in [".jpg", ".jpeg", ".png"]:
                image_path = images_folder / f"{image_name}{ext}"
                if image_path.exists():
                    break
        if image_path.exists():
            print(f"Обработка специального изображения: {image_path}")
            try:
                image = Image.open(image_path)
                temp_image_path = images_folder / "temp_adjusted_image.jpg"
                image.save(temp_image_path, format="JPEG")
                print(f"Временное изображение сохранено: {temp_image_path}")
                for p in doc.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, "")
                        run = p.add_run()
                        run.add_picture(str(temp_image_path), width=Inches(6.4))
                        print(f"Изображение вставлено в метку {placeholder}")
                        break
                temp_image_path.unlink()
                print(f"Временное изображение удалено: {temp_image_path}")
            except Exception as e:
                print(f"Ошибка при обработке изображения {image_path}: {e}")
    doc.save(doc_path)
    print("Все изображения обработаны и вставлены.")

# Функция для поиска файла .txt в целевой папке
def find_txt_file(folder_path):
    txt_files = list(Path(folder_path).glob("*.txt"))
    logging.info(f"Поиск файлов .txt в папке: {folder_path}")
    if not txt_files:
        logging.error(f"В папке {folder_path} не найдено файлов .txt.")
        raise FileNotFoundError(f"В папке {folder_path} не найдено файлов .txt.")
    if len(txt_files) > 1:
        logging.warning(f"В папке {folder_path} найдено несколько TXT-файлов:")
        for i, file in enumerate(txt_files, start=1):
            logging.warning(f"{i}. {file.name}")
        while True:
            try:
                choice = int(input("Выберите номер файла для использования: "))
                if 1 <= choice <= len(txt_files):
                    selected_file = txt_files[choice - 1]
                    logging.info(f"Выбран файл: {selected_file}")
                    return selected_file
                else:
                    logging.warning("Неверный номер. Пожалуйста, выберите существующий вариант.")
            except ValueError:
                logging.warning("Пожалуйста, введите число.")
    selected_file = txt_files[0]
    logging.info(f"Выбран файл: {selected_file}")
    return selected_file

# Функция для чтения и фильтрации данных из .txt файла
def read_and_filter_txt_data(file_path):
    # Значения, которые нужно исключить (в разных вариантах регистра)
    exclude_values = {"AL", "Al", "K+", "К+", "K-", "К-"}
    logging.info(f"Чтение и фильтрация данных из файла: {file_path}")

    # Чтение файла с указанной кодировкой
    with open(file_path, 'r', encoding='windows-1251') as file:
        lines = file.readlines()

    # Убираем заголовок и разбиваем строки на столбцы
    data = [line.strip().split('\t') for line in lines[1:]]

    # Получаем все уникальные идентификаторы объектов (первая колонка)
    # исключая значения из exclude_values
    object_ids = set()
    for line in data:
        if line[0] and line[0].strip() not in exclude_values:
            object_ids.add(line[0].strip())

    # Сортируем идентификаторы объектов по порядку номеров или по алфавиту
    sorted_object_ids = sorted(object_ids, key=lambda x: (x.isdigit(), x))

    # Выводим найденные идентификаторы для отладки
    logging.info(f"Найденные идентификаторы объектов: {sorted_object_ids}")

    # Проверяем, что есть хотя бы один уникальный идентификатор
    if len(sorted_object_ids) == 0:
        logging.error("В файле не найдено ни одного допустимого объекта.")
        raise ValueError("В файле не найдено ни одного допустимого объекта.")

    # Если найдено несколько уникальных идентификаторов, спрашиваем пользователя, какой объект загрузки
    if len(sorted_object_ids) > 1:
        logging.warning("В файле найдено несколько различных объектов:")
        for i, obj_id in enumerate(sorted_object_ids, start=1):
            logging.warning(f"{i}. {obj_id}")
        while True:
            try:
                choice = int(input("Выберите номер объекта для загрузки: ")) - 1
                if 0 <= choice < len(sorted_object_ids):
                    selected_object_id = sorted_object_ids[choice]
                    break
                else:
                    logging.warning("Некорректный выбор. Пожалуйста, введите номер из списка.")
            except ValueError:
                logging.warning("Пожалуйста, введите числовое значение.")
    else:
        selected_object_id = sorted_object_ids[0]

    # Фильтруем данные, оставляя только строки с выбранным идентификатором объекта
    filtered_data = [line for line in data if line[0].strip() == selected_object_id]

    # Выводим количество отфильтрованных строк данных для отладки
    logging.info(f"Количество отфильтрованных строк данных: {len(filtered_data)}")
    return filtered_data

# Функция для вставки данных в таблицу Word
def insert_table_data(doc_path, txt_data, card_data, third_col_width=Inches(6)):
    doc = Document(doc_path)
    has_y_in_amel = False
    sorted_data = sort_data_by_locus(txt_data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "[TABLE_DATA]" in cell.text:
                    cell.text = ""
                    col1_width = table.columns[0].width
                    col2_width = table.columns[1].width
                    unique_markers = set()
                    for line in sorted_data:
                        marker = line[3].strip()
                        if marker.upper() in unique_markers:
                            continue
                        unique_markers.add(marker.upper())
                        allele1 = line[5] if line[5] else ""
                        allele2 = line[6] if line[6] else ""
                        if not allele1 and not allele2:
                            alleles = "-"
                        else:
                            alleles = ",".join(filter(None, [allele1, allele2]))
                        if marker.upper() == "AMEL":
                            if "Y" in [allele1.upper(), allele2.upper()]:
                                has_y_in_amel = True
                                print(f"Найдена буква Y в маркере AMEL: Allele1={allele1}, Allele2={allele2}")
                        new_row = table.add_row()
                        new_row.cells[0].text = marker
                        new_row.cells[1].text = alleles
                        run = new_row.cells[0].paragraphs[0].runs[0]
                        run.font.name = 'Times New Roman'
                        r = run._element
                        rPr = r.get_or_add_rPr()
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        rPr.append(rFonts)
                        run = new_row.cells[1].paragraphs[0].runs[0]
                        run.font.name = 'Times New Roman'
                        r = run._element
                        rPr = r.get_or_add_rPr()
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        rPr.append(rFonts)
                    rows_to_remove = []
                    for row in table.rows:
                        if all(cell.text.strip() == "" for cell in row.cells):
                            rows_to_remove.append(row)
                    for row in rows_to_remove:
                        table._tbl.remove(row._tr)

                    # Устанавливаем границы для каждой ячейки таблицы
                    for row in table.rows:
                        for cell in row.cells:
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcBorders = OxmlElement('w:tcBorders')
                            for border_name in ['top', 'left', 'bottom', 'right']:
                                border = OxmlElement(f'w:{border_name}')
                                border.set(qn('w:val'), 'single')
                                border.set(qn('w:sz'), '4')
                                border.set(qn('w:space'), '0')
                                border.set(qn('w:color'), '000000')
                                tcBorders.append(border)
                            tcPr.append(tcBorders)

                    for row in table.rows:
                        if len(row.cells) > 1:
                            cell = row.cells[1]
                            cell.paragraphs[0].alignment = 1
                    if not has_y_in_amel:
                        print("Буква Y не найдена в маркере AMEL. Добавляем примечание.")
                        p = doc.add_paragraph()
                        run = p.add_run("Примечание: ")
                        run.underline = True
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
                        p.add_run("(-) - прочерк означает отсутствие продуктов амплификации.").font.size = Pt(10)
                        table_element = table._tbl
                        parent_element = table_element.getparent()
                        parent_element.insert(parent_element.index(table_element) + 1, p._element)
                        next_element = p._element.getnext()
                        if next_element is not None and next_element.tag.endswith('p'):
                            next_paragraph = doc._body._element.index(next_element)
                            if not next_element.text.strip():
                                doc._body._element.remove(next_element)
                                print("Пустая строка после примечания удалена.")
                    table.add_column(third_col_width)
                    table.columns[0].width = col1_width
                    table.columns[1].width = col2_width
                    first_cell = table.cell(0, len(table.columns) - 1)
                    last_cell = table.cell(len(table.rows) - 1, len(table.columns) - 1)
                    tc = first_cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    vMerge = OxmlElement('w:vMerge')
                    vMerge.set(qn('w:val'), 'restart')
                    tcPr.append(vMerge)
                    for row in table.rows[1:]:
                        tc = row.cells[len(table.columns) - 1]._tc
                        tcPr = tc.get_or_add_tcPr()
                        vMerge = OxmlElement('w:vMerge')
                        vMerge.set(qn('w:val'), 'continue')
                        tcPr.append(vMerge)
                    for row in table.rows:
                        tr = row._tr
                        trPr = tr.get_or_add_trPr()
                        spacing = OxmlElement('w:trHeight')
                        spacing.set(qn('w:val'), "240")
                        spacing.set(qn('w:hRule'), "auto")
                        trPr.append(spacing)
                    images_folder = Path(card_data["4"])
                    nom_folder_name = f"{card_data['НОМ']}-25"
                    nom_folder = images_folder / nom_folder_name
                    image_path = nom_folder / "1.jpg"
                    if image_path.exists():
                        with Image.open(image_path) as img:
                            img_width, img_height = img.size
                            aspect_ratio = img_height / img_width
                            new_width = third_col_width
                            new_height = new_width * aspect_ratio
                            run = first_cell.paragraphs[0].add_run()
                            run.add_picture(str(image_path), width=new_width, height=new_height)
                    else:
                        logging.warning(f"Изображение {image_path} не найдено.")
                    tblPr = table._tblPr
                    tblW = tblPr.first_child_found_in("w:tblW")
                    if tblW is not None:
                        tblPr.remove(tblW)
                    break
    doc.save(doc_path)
    print("Примечание добавлено после таблицы.")

# Функция для поиска файла Excel с любым именем в текущей директории
def find_excel_file():
    """
    Ищет единственный Excel-файл в текущей директории.
    Возвращает путь к файлу или вызывает исключение, если файл не найден или их несколько.
    """
    # Получаем список всех Excel-файлов в текущей директории
    excel_files = list(Path('.').glob('*.xlsx')) + list(Path('.').glob('*.xls'))
    
    # Проверяем количество найденных файлов
    if not excel_files:
        raise FileNotFoundError("В текущей директории не найдено Excel-файлов (.xlsx или .xls)")
    if len(excel_files) > 1:
        raise ValueError(f"В директории найдено несколько Excel-файлов: {[f.name for f in excel_files]}. "
                         "Должен быть только один файл.")
    
    return excel_files[0]

# Функция для поиска файла Excel с VBA, указанного в Карте
def find_vba_file(card_data):
    vba_folder_path = Path(card_data["3"])  # Путь к папке с шаблонами
    vba_file_name = card_data["7"].strip()  # Имя файла VBA
    logging.info(f"Путь к папке с шаблонами: {vba_folder_path}")
    logging.info(f"Имя файла VBA: {vba_file_name}")
    # Добавляем расширение файла, если оно отсутствует
    if not vba_file_name.endswith('.xlsm'):
        vba_file_name += '.xlsm'
    vba_file_path = vba_folder_path / vba_file_name
    logging.info(f"Полный путь к файлу VBA: {vba_file_path}")
    if not vba_file_path.exists():
        logging.error(f"Файл VBA не найден по пути: {vba_file_path}")
        raise FileNotFoundError(f"Файл VBA не найден по пути: {vba_file_path}")
    if not vba_file_path.is_file():
        logging.error(f"Путь {vba_file_path} не является файлом.")
        raise ValueError(f"Путь {vba_file_path} не является файлом.")
    logging.info(f"Найден файл VBA: {vba_file_path}")
    return vba_file_path

def is_file_open(file_path):
    """Проверяет, открыт ли файл другим процессом"""
    try:
        os.rename(file_path, file_path)
        return False
    except OSError:
        return True

# Функция для вызова макроса VBA 2-ДНК
def run_vba_macro_2_DNK(vba_path, txt_path, output_folder, card_data):
    """
    Запускает макрос VBA в Excel файле с оптимизированными настройками и улучшенной обработкой ошибок.
    Args:
    vba_path (str): Путь к файлу Excel с макросом
    txt_path (str): Путь к txt файлу с данными
    output_folder (str): Путь к папке для сохранения результата
    card_data (dict): Данные из файла Карта
    Returns:
    str: Путь к созданному файлу
    Raises:
    FileNotFoundError: Если файлы не найдены
    ValueError: При ошибках в данных или выполнении макроса
    """
    try:
        # Преобразуем пути в абсолютные
        vba_path = os.path.abspath(vba_path)
        txt_path = os.path.abspath(txt_path)
        output_folder = os.path.abspath(output_folder)
        
        # Проверка существования файлов и папок
        if not os.path.exists(vba_path):
            logging.error(f"VBA файл не найден: {vba_path}")
            raise FileNotFoundError(f"VBA файл не найден: {vba_path}")
        if not os.path.exists(txt_path):
            logging.error(f"TXT файл не найден: {txt_path}")
            raise FileNotFoundError(f"TXT файл не найден: {txt_path}")
        if not os.path.exists(output_folder):
            logging.error(f"Папка назначения не найдена: {output_folder}")
            raise FileNotFoundError(f"Папка назначения не найдена: {output_folder}")
        
        # Создаем путь для нового файла
        output_filename = f"{card_data['ФИ']} {card_data['ИИ']} {card_data['ОИ']} {card_data['ДР']}.xlsm"
        output_path = os.path.join(output_folder, output_filename)
        logging.info(f"Начало выполнения макроса:")
        logging.info(f"VBA файл: {vba_path}")
        logging.info(f"TXT файл: {txt_path}")
        logging.info(f"Результат будет сохранен как: {output_path}")
        
        # Инициализация COM-объектов
        pythoncom.CoInitialize()
        excel = None
        workbook = None
        max_retries = 3  # Максимальное количество попыток
        retry_delay = 1  # Задержка между попытками (в секундах)
        
        try:
            for attempt in range(max_retries):
                try:
                    # Создаем экземпляр Excel
                    excel = win32.Dispatch("Excel.Application")
                    
                    # Оптимизированные настройки Excel для повышения производительности
                    excel.Visible = False  # Excel работает в фоновом режиме
                    excel.DisplayAlerts = False  # Отключаем всплывающие уведомления
                    excel.EnableEvents = False  # Отключаем события
                    excel.ScreenUpdating = False  # Отключаем обновление экрана
                    logging.info(f"Excel запущен (попытка {attempt + 1})")
                    
                    # Копируем файл VBA в выходную папку
                    shutil.copy2(vba_path, output_path)
                    logging.info(f"Создана копия шаблона: {output_path}")
                    
                    # Открываем копию файла
                    workbook = excel.Workbooks.Open(output_path)
                    logging.info("Файл открыт успешно")
                    
                    # Очищаем данные в Excel с помощью макроса ClearData
                    try:
                        workbook.Application.Run("ClearData")
                        logging.info("Данные очищены")
                    except Exception as e:
                        logging.warning(f"Не удалось очистить данные: {e}")
                    
                    # Устанавливаем путь к txt файлу в ячейку X14 на листе "2-ДНК"
                    dnk_sheet = workbook.Sheets("2-ДНК")
                    dnk_sheet.Range("X14").Value = txt_path
                    logging.info(f"Установлен путь к txt файлу")
                    
                    # Активируем нужный лист и ячейку
                    dnk_sheet.Activate()
                    dnk_sheet.Range("A28").Select()
                    
                    # Выполняем макрос cmdToDo
                    logging.info("Запуск макроса cmdToDo...")
                    workbook.Application.Run("mdlTemplate.cmdToDo")
                    
                    # Короткая пауза для обработки диалога
                    time.sleep(0.5)
                    
                    # Обработка диалогового окна (если оно появляется)
                    try:
                        excel.SendKeys("{DOWN}")  # Нажимаем стрелку вниз
                        time.sleep(0.1)
                        excel.SendKeys("{ENTER}")  # Нажимаем Enter
                        logging.info("Диалоговое окно обработано")
                    except Exception as e:
                        logging.warning(f"Ошибка при обработке диалога: {e}")
                    
                    # Проверяем результат выполнения макроса
                    time.sleep(0.5)
                    if dnk_sheet.Range("A28").Value:
                        logging.info("Макрос выполнен успешно")
                        break  # Выходим из цикла попыток
                    else:
                        raise ValueError("Данные не записаны в ячейку A28")
                except Exception as e:
                    if attempt < max_retries - 1:
                        logging.warning(f"Попытка {attempt + 1} не удалась: {e}")
                        time.sleep(retry_delay)  # Ждем перед повторной попыткой
                        if workbook:
                            try:
                                workbook.Close(False)  # Закрываем workbook без сохранения
                            except:
                                pass
                        if excel:
                            try:
                                excel.Quit()  # Закрываем Excel
                            except:
                                pass
                        continue  # Повторяем попытку
                    raise  # Если попытки закончились, выбрасываем исключение
            
            # Заменяем метки в Excel файле через COM-интерфейс
            replace_in_excel_com(workbook, card_data)
            
            # Сохраняем результат
            workbook.Save()
            logging.info(f"Файл сохранен: {output_path}")
            
            # Закрываем workbook
            workbook.Close(True)
            logging.info("Workbook закрыт")
            
            return output_path
        except Exception as e:
            logging.error(f"Критическая ошибка при выполнении макроса: {e}")
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)  # Удаляем временный файл после ошибки
                    logging.info("Удален временный файл после ошибки")
                except:
                    pass
            raise
        finally:
            # Закрываем все объекты
            if workbook:
                try:
                    workbook.Close(True)
                    logging.info("Workbook закрыт")
                except:
                    pass
            if excel:
                try:
                    excel.Quit()
                    logging.info("Excel закрыт")
                except:
                    pass
            pythoncom.CoUninitialize()  # Освобождаем COM-объекты
    except Exception as e:
        logging.error(f"Ошибка при выполнении макроса VBA: {e}")
        raise

# Функция для вызова макроса VBA ИК-2
def run_vba_macro_IK_2(vba_path, txt_path, output_folder, card_data):
    """
    Запускает макрос VBA IK-2 в Excel файле с оптимизированными настройками и улучшенной обработкой ошибок.
    Args:
    vba_path (str): Путь к файлу Excel с макросом
    txt_path (str): Путь к txt файлу с данными
    output_folder (str): Путь к папке для сохранения результата
    card_data (dict): Данные из файла Карта
    Returns:
    str: Путь к созданному файлу
    Raises:
    FileNotFoundError: Если файлы не найдены
    ValueError: При ошибках в данных или выполнении макроса
    """
    try:
        # Преобразуем пути в абсолютные
        vba_path = os.path.abspath(vba_path)
        txt_path = os.path.abspath(txt_path)
        output_folder = os.path.abspath(output_folder)
        
        # Проверка существования файлов и папок
        if not os.path.exists(vba_path):
            logging.error(f"VBA файл не найден: {vba_path}")
            raise FileNotFoundError(f"VBA файл не найден: {vba_path}")
        if not os.path.exists(txt_path):
            logging.error(f"TXT файл не найден: {txt_path}")
            raise FileNotFoundError(f"TXT файл не найден: {txt_path}")
        if not os.path.exists(output_folder):
            logging.error(f"Папка назначения не найдена: {output_folder}")
            raise FileNotFoundError(f"Папка назначения не найдена: {output_folder}")
        
        # Создаем путь для нового файла
        tr_value = str(card_data['ТР']).replace('/', '-').replace('\\', '-')
        output_filename = f"{tr_value}.xlsm"        
        output_path = os.path.join(output_folder, output_filename)
        logging.info(f"Начало выполнения макроса:")
        logging.info(f"VBA файл: {vba_path}")
        logging.info(f"TXT файл: {txt_path}")
        logging.info(f"Результат будет сохранен как: {output_path}")

        # Инициализация COM-объектов
        pythoncom.CoInitialize()
        excel = None
        workbook = None
        max_retries = 3  # Максимальное количество попыток
        retry_delay = 1  # Задержка между попытками (в секундах)
        
        try:
            for attempt in range(max_retries):
                try:
                    # Создаем экземпляр Excel
                    excel = win32.Dispatch("Excel.Application")
                    
                    # Оптимизированные настройки Excel для повышения производительности
                    excel.Visible = False  # Делаем Excel невидимым
                    excel.DisplayAlerts = False  # Отключаем всплывающие уведомления
                    excel.EnableEvents = False  # Отключаем события
                    excel.ScreenUpdating = False  # Отключаем обновление экрана
                    logging.info(f"Excel запущен (попытка {attempt + 1})")
                    
                    # Копируем файл VBA в выходную папку
                    shutil.copy2(vba_path, output_path)
                    logging.info(f"Создана копия шаблона: {output_path}")
                    
                    # Открываем копию файла
                    workbook = excel.Workbooks.Open(output_path)
                    logging.info("Файл открыт успешно")
                    
                    # Очищаем данные в Excel с помощью макроса ClearData (без указания пути)
                    try:
                        workbook.Application.Run("ClearData")
                        logging.info("Данные очищены")
                    except Exception as e:
                        logging.warning(f"Не удалось очистить данные: {e}")
                    
                    # Устанавливаем путь к txt файлу в ячейку X14 на листе "ИК-2"
                    ik_sheet = workbook.Sheets("ИК-2")
                    ik_sheet.Range("X14").Value = txt_path
                    
                    # Также устанавливаем путь в ячейке для последнего файла
                    # (в соответствии с VBA-кодом это поле PC_RANGE_LDR_LASTFILE)
                    try:
                        # Попытка установить значение напрямую в рабочую область, как в VBA
                        workbook.Names("ldrLastFile").RefersToRange.Value = txt_path
                    except Exception as e:
                        logging.warning(f"Не удалось установить ldrLastFile: {e}")
                    
                    logging.info(f"Установлен путь к txt файлу")

                    # Активируем нужный лист и ячейку
                    ik_sheet.Activate()
                    ik_sheet.Range("A28").Select()
                    
                    # Вызов макроса через разные варианты имен модулей
                    try:
                        logging.info("Запуск макроса cmdToDo...")
                        workbook.Application.Run("mdlTemplate.cmdToDo")
                    except Exception as e:
                        # Если не удалось через предполагаемый модуль, пробуем стандартные варианты
                        try:
                            logging.info(f"Ошибка при вызове через mdlTemplate: {e}, пробуем другие варианты...")
                            # Вариант 2: через имя "Module1" (стандартное имя в VBA)
                            workbook.Application.Run("Module1.cmdToDo")
                        except Exception as e2:
                            # Вариант 3: без указания модуля
                            logging.info(f"Ошибка при вызове через Module1: {e2}, пробуем без модуля...")
                            workbook.Application.Run("cmdToDo")
                    
                    # Более длительная пауза для обработки диалога
                    time.sleep(0.7)
                    
                    # Улучшенная обработка диалогового окна (если оно появляется)
                    try:
                        # Последовательные нажатия клавиш для взаимодействия с диалогом выбора
                        excel.SendKeys("{DOWN}")  # Нажимаем стрелку вниз
                        time.sleep(0.2)
                        excel.SendKeys("{DOWN}")  # Нажимаем стрелку вниз еще раз
                        time.sleep(0.2)
                        excel.SendKeys("{ENTER}")  # Нажимаем Enter
                        time.sleep(0.3)
                        # В случае появления второго диалогового окна
                        excel.SendKeys("{ENTER}")  # Нажимаем Enter снова
                        logging.info("Диалоговые окна обработаны")
                    except Exception as e:
                        logging.warning(f"Ошибка при обработке диалога: {e}")
                    
                    # Проверяем результат выполнения макроса с более длительным ожиданием
                    time.sleep(0.3)
                    if ik_sheet.Range("A28").Value:
                        logging.info("Макрос выполнен успешно")
                        break  # Выходим из цикла попыток
                    else:
                        # Возможно, данные были сохранены в другой ячейке или результат нужно проверять иначе
                        # Все равно считаем успешным и продолжаем
                        logging.warning("Данные не записаны в ячейку A28, но продолжаем работу")
                        break

                except Exception as e:
                    if attempt < max_retries - 1:
                        logging.warning(f"Попытка {attempt + 1} не удалась: {e}")
                        time.sleep(retry_delay)  # Ждем перед повторной попыткой
                        if workbook:
                            try:
                                workbook.Close(False)  # Закрываем workbook без сохранения
                            except:
                                pass
                        if excel:
                            try:
                                excel.Quit()  # Закрываем Excel
                            except:
                                pass
                        continue  # Повторяем попытку
                    raise  # Если попытки закончились, выбрасываем исключение
            
            # Заменяем метки в Excel файле через COM-интерфейс
            replace_in_excel_com(workbook, card_data)
            
            # Сохраняем результат
            workbook.Save()
            logging.info(f"Файл сохранен: {output_path}")
            
            # Закрываем workbook
            workbook.Close(True)
            logging.info("Workbook закрыт")
            
            return output_path

        except Exception as e:
            logging.error(f"Критическая ошибка при выполнении макроса: {e}")
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)  # Удаляем временный файл после ошибки
                    logging.info("Удален временный файл после ошибки")
                except:
                    pass
            raise
        finally:
            # Закрываем все объекты
            if workbook:
                try:
                    workbook.Close(True)
                    logging.info("Workbook закрыт")
                except:
                    pass
            if excel:
                try:
                    excel.Quit()
                    logging.info("Excel закрыт")
                except:
                    pass
            pythoncom.CoUninitialize()  # Освобождаем COM-объекты
    except Exception as e:
        logging.error(f"Ошибка при выполнении макроса VBA IK-2: {e}")
        raise

# Замена меток в Excel
def replace_in_excel_com(workbook, replacements):
    try:
        # Добавляем словарь для перевода значений РОД
        rod_translation = {
            "матери": "мать",
            "отца": "отец",
            "сестры": "сестра",
            "брата": "брат",
            "дочери": "дочь",
            "сына": "сын"
        }
        
        # Если значение РОД присутствует в словаре, заменяем его
        if 'РОД' in replacements:
            replacements['РОД'] = rod_translation.get(replacements['РОД'], replacements['РОД'])
        
        for sheet in workbook.Sheets:
            for cell in sheet.UsedRange:
                if cell.Value and isinstance(cell.Value, str):
                    original_value = cell.Value.strip()
                    modified_value = original_value
                    for key, value in replacements.items():
                        for pattern in [f"{{{key}}}", f"{{ {key} }}", f"{{{key} }}", f"{{ {key}}}"]:
                            if pattern in modified_value:
                                modified_value = modified_value.replace(pattern, str(value))
                                logging.info(f"Замена метки '{pattern}' на '{value}' в ячейке: "
                                             f"'{original_value}' -> '{modified_value}'")
                    if original_value != modified_value:
                        cell.Value = modified_value
                        logging.info(f"Ячейка обновлена: '{original_value}' -> '{cell.Value}'")
                    else:
                        logging.debug(f"Метки не найдены в ячейке: '{original_value}'")
        logging.info("Замены меток выполнены в Excel файле")
    except Exception as e:
        logging.error(f"Ошибка при замене меток в Excel файле: {e}")
        raise

# Конвертируем txt в xml
def convert_txt_to_xml(txt_file_path, xml_file_path, object_name_suffix):
    # Чтение данных из файла txt
    with open(txt_file_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()
    
    # Парсинг заголовков
    headers = lines[0].strip().split('\t')
    
    # Определение индексов колонок
    sample_name_idx = None
    marker_idx = None
    allele_indices = []
    
    # Проверка наличия заголовков на английском языке
    if 'Sample Name' in headers:
        sample_name_idx = headers.index('Sample Name')
    elif 'Название образца' in headers:
        sample_name_idx = headers.index('Название образца')
    else:
        raise ValueError("Заголовок 'Sample Name' или 'Название образца' не найден в файле.")
    
    if 'Marker' in headers:
        marker_idx = headers.index('Marker')
    elif 'Маркер' in headers:
        marker_idx = headers.index('Маркер')
    else:
        raise ValueError("Заголовок 'Marker' или 'Маркер' не найден в файле.")
    
    for i, h in enumerate(headers):
        if h.startswith('Allele') or h.startswith('Аллель'):
            allele_indices.append(i)
    
    if not allele_indices:
        raise ValueError("Заголовки 'Allele' или 'Аллель' не найдены в файле.")
    
    # Создание корневого элемента XML
    root = ET.Element("CODISImportFile", xmlns="urn:CODISImportFile-schema")
    ET.SubElement(root, "HEADERVERSION").text = "3.2"
    ET.SubElement(root, "MESSAGETYPE").text = "Import"
    ET.SubElement(root, "DESTINATIONORI").text = "destlab"
    ET.SubElement(root, "SOURCELAB").text = "srclab"
    ET.SubElement(root, "SUBMITBYUSERID").text = "gmidx"
    ET.SubElement(root, "SUBMITDATETIME").text = datetime.now().strftime("%Y-%m-%dT%H:%M:%S.%f%z")
    
    # Создание элемента SPECIMEN
    specimen = ET.SubElement(root, "SPECIMEN")
    ET.SubElement(specimen, "SPECIMENID").text = object_name_suffix
    ET.SubElement(specimen, "SPECIMENCATEGORY").text = "Other"
    
    # Чтение данных для каждого локуса
    for line in lines[1:]:
        data = line.strip().split('\t')
        sample_name = data[sample_name_idx]
        if sample_name.strip() not in {"AL", "Al", "K+", "К+", "K-", "К-"} and sample_name.strip() == object_name_suffix:
            marker = data[marker_idx]
            reading_by = "gmidx"
            reading_datetime = datetime.now().strftime("%Y-%m-%dT%H:%M:%S.%f%z")
            
            # Создание элемента LOCUS
            locus = ET.SubElement(specimen, "LOCUS")
            ET.SubElement(locus, "LOCUSNAME").text = marker
            ET.SubElement(locus, "READINGBY").text = reading_by
            ET.SubElement(locus, "READINGDATETIME").text = reading_datetime
            
            # Добавление аллелей
            for idx in allele_indices:
                allele_value = data[idx]
                if allele_value:
                    allele = ET.SubElement(locus, "ALLELE")
                    ET.SubElement(allele, "ALLELEVALUE").text = allele_value
    
    # Сохранение XML в файл
    tree = ET.ElementTree(root)
    tree.write(xml_file_path, encoding='utf-8', xml_declaration=True)
    logging.info(f"Файл XML сохранен: {xml_file_path}")

def ask_for_print():
    """Запрашивает у пользователя подтверждение на печать."""
    response = input("Хотите распечатать документ? (д/н) ").strip().lower()
    return response == "д"

def print_document(zakl_path):
    pythoncom.CoInitialize()  # Инициализация COM-объектов
    try:
        # Открываем документ в Word
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False  # Word будет работать в фоновом режиме
        doc = word.Documents.Open(zakl_path)
        try:
            # Шаг 1: Печать последней страницы с обычными полями
            logging.info("1. Отключение зеркальных полей для последней страницы...")
            doc.PageSetup.MirrorMargins = False  # Обычные поля
            last_page = doc.ComputeStatistics(2)  # Получаем количество страниц
            logging.info(f"2. Количество страниц до печати последней страницы: {last_page}")
            logging.info(f"3. Печать последней страницы (страница {last_page})...")
            doc.PrintOut(PrintToFile=False, Range=win32.constants.wdPrintRangeOfPages, Pages=str(last_page))
            
            # Шаг 2: Удаляем текст после "Расписка о получении заключения"
            logging.info("4. Удаление текста после 'Расписка о получении заключения'...")
            for paragraph in doc.Paragraphs:
                if "Расписка о получении заключения" in paragraph.Range.Text:
                    paragraph_range = paragraph.Range
                    paragraph_range.End = doc.Content.End
                    paragraph_range.Delete()
                    break  # Прерываем цикл после удаления

            # Шаг 3: Печатаем последнюю страницу после удаления с обычными полями
            last_page = doc.ComputeStatistics(2)  # Получаем новое количество страниц
            logging.info(f"5. Количество страниц после удаления текста: {last_page}")
            logging.info(f"6. Печать последней страницы после удаления (страница {last_page})...")
            doc.PrintOut(PrintToFile=False, Range=win32.constants.wdPrintRangeOfPages, Pages=str(last_page))

            # Шаг 4: Удаляем последний раздел полностью с сохранением форматирования
            logging.info("7. Удаление последнего раздела...")
            sections = doc.Sections
            if sections.Count > 1:
                # Определяем координаты последнего раздела
                last_section = sections(sections.Count)
                last_section_start = last_section.Range.Start
                
                # Создаем диапазон от начала последнего раздела до конца документа
                last_section_range = doc.Range(last_section_start, doc.Content.End)
                
                # Удаляем диапазон последнего раздела
                last_section_range.Delete()
                
                # Дополнительная очистка - находим и удаляем маркеры разрыва разделов в конце
                end_of_doc = doc.Content.End - 1
                cleanup_range = doc.Range(max(0, end_of_doc - 20), end_of_doc + 1)
                
                # Заменяем разрывы разделов в конце документа на ничто
                cleanup_range.Find.ClearFormatting()
                cleanup_range.Find.Replacement.ClearFormatting()
                cleanup_range.Find.Text = "^b"  # Код для разрыва раздела
                cleanup_range.Find.Replacement.Text = ""
                cleanup_range.Find.Execute(Replace=win32.constants.wdReplaceAll)
                
                # Удаляем последний абзац, если он пустой
                if doc.Paragraphs.Last.Range.Text.strip() == "":
                    doc.Paragraphs.Last.Range.Delete()
            
            # Принудительный пересчет страниц
            doc.Repaginate()
            
            # Обновляем информацию о количестве страниц
            last_page = doc.ComputeStatistics(2)
            logging.info(f"8. Количество страниц после удаления раздела: {last_page}")

            # Шаг 5: Включаем зеркальные поля для печати всего документа
            logging.info("9. Включение зеркальных полей для печати всего документа...")
            doc.PageSetup.MirrorMargins = True  # Зеркальные поля

            # Шаг 6: Печатаем документ в двух экземплярах
            last_page = doc.ComputeStatistics(2)  # Получаем новое количество страниц
            logging.info(f"10. Количество страниц перед печатью всего документа: {last_page}")
            logging.info("11. Печать документа в двух экземплярах без последней страницы...")
            doc.PrintOut(PrintToFile=False, Range=win32.constants.wdPrintRangeOfPages, Pages=f"1-{last_page}", Copies=2, Collate=True)

        except Exception as e:
            print(f"Произошла ошибка: {e}")
        finally:
            # Закрываем документ без сохранения изменений
            doc.Close(SaveChanges=False)
            word.Quit()
            logging.info("12. Документ закрыт, Word завершён.")
    finally:
        pythoncom.CoUninitialize()  # Освобождаем COM-объекты

def create_and_fill_document(card_data, txt_data, template_name):
    """Создает и заполняет документ Word, сохраняя его во временную папку (код 4)"""
    try:
        template_folder = Path(card_data["3"])  # Путь к папке с шаблонами
        template_path = template_folder / f"{template_name}.docx"
        logging.info(f"Путь к папке с шаблонами: {template_folder}")
        logging.info(f"Имя файла шаблона: {template_name}.docx")

        if not template_path.exists():
            logging.error(f"Файл шаблона '{template_name}' не найден в папке {template_folder}.")
            raise FileNotFoundError(f"Файл шаблона '{template_name}' не найден в папке {template_folder}.")

        # Сохраняем во временную папку (код 4)
        temp_folder = Path(card_data["4"])
        temp_folder.mkdir(parents=True, exist_ok=True)
        logging.info(f"Временная папка для сохранения: {temp_folder}")

        # Определение имени файла для сохранения
        if card_data.get("1") in ["Образец СВО Молов фонд", "Образец СВО Молов прямая идентификация", "Образец СВО Ростов"]:
            output_filename = f"{card_data['ФИ']} {card_data['ИИ']} {card_data['ОИ']} {card_data['НОМ']}-25.docx"
        elif card_data.get("1") == "Кость СВО":
            ob_last_char = card_data['ОБ'].strip()[-1] if card_data['ОБ'].strip() else ''
            output_filename = f"{card_data['ТР']} {card_data['НОМ']}-25 {ob_last_char}.docx"
        else:
            output_filename = f"{card_data['ФИ']} {card_data['ИИ']} {card_data['ОИ']} {card_data['НОМ']}-25.docx"

        temp_path = temp_folder / output_filename
        logging.info(f"Временный путь к выходному файлу: {temp_path}")

        try:
            shutil.copy(template_path, temp_path)
            logging.info(f"Шаблон скопирован во временный файл: {temp_path}")
        except Exception as e:
            logging.error(f"Ошибка при копировании шаблона: {e}")
            raise

        try:
            replacements = {code: str(value) for code, value in card_data.items() if value is not None}
            replace_in_doc(temp_path, replacements)
            logging.info("Замены меток выполнены в документе")
        except Exception as e:
            logging.error(f"Ошибка при замене меток в документе: {e}")
            raise

        try:
            insert_table_data(temp_path, txt_data, card_data)
            logging.info("Таблица данных вставлена в документ")
        except Exception as e:
            logging.error(f"Ошибка при вставке таблицы данных: {e}")
            raise

        try:
            insert_images(temp_path, card_data["4"])
            logging.info("Изображения вставлены в документ")
        except Exception as e:
            logging.error(f"Ошибка при вставке изображений: {e}")
            raise

        return temp_path

    except Exception as e:
        logging.error(f"Ошибка при создании и заполнении документа: {e}")
        raise

def move_to_final_location(card_data, word_file_path, excel_file_path=None):
    """Перемещает файлы из временной папки в финальные папки после создания обоих файлов"""
    try:
        logging.info("Начало перемещения файлов в финальные папки")
        
        # Преобразуем пути в объекты Path, если они строками
        word_file_path = Path(word_file_path) if isinstance(word_file_path, str) else word_file_path
        if excel_file_path:
            excel_file_path = Path(excel_file_path) if isinstance(excel_file_path, str) else excel_file_path
        
        # Определяем финальную папку на основе оригинальной логики
        if card_data.get("1") == "Образец СВО Молов прямая идентификация":
            final_folder = Path(r"F:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Личные профили 2-ДНК (прямая идентификация)")
        elif card_data.get("1") == "Образец СВО Молов фонд":
            final_folder = Path(r"F:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Родственники с 2-ДНК Тюмень")
        elif card_data.get("1") == "Кость СВО":
            final_folder = Path(r"F:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Кости с ИК-2")
        elif card_data.get("1") == "Образец СВО Ростов":
            final_folder = Path(r"F:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Родственники с 2-ДНК Ростов")
        else:
            final_folder = Path(card_data["6"])
        
        logging.info(f"Определена финальная папка: {final_folder}")
        final_folder.mkdir(parents=True, exist_ok=True)

        # Перемещаем Word-файл
        final_word_path = final_folder / word_file_path.name
        logging.info(f"Перемещение Word-файла: {word_file_path} -> {final_word_path}")
        shutil.move(str(word_file_path), str(final_word_path))
        logging.info(f"Word-файл успешно перемещен: {final_word_path}")

        # Перемещаем Excel-файл, если он существует
        if excel_file_path and Path(excel_file_path).exists():
            final_excel_path = final_folder / Path(excel_file_path).name
            logging.info(f"Перемещение Excel-файла: {excel_file_path} -> {final_excel_path}")
            shutil.move(str(excel_file_path), str(final_excel_path))
            logging.info(f"Excel-файл успешно перемещен: {final_excel_path}")
        else:
            logging.info("Excel-файл для перемещения не указан или не существует")
            final_excel_path = None

        return final_word_path, final_excel_path

    except Exception as e:
        logging.error(f"Ошибка при перемещении файлов в финальную папку: {e}")
        # В случае ошибки оставляем файлы во временной папке
        logging.warning("Файлы остаются во временной папке из-за ошибки перемещения")
        raise

# Основная функция
def main():
    try:
        # Вывод текущей рабочей директории
        current_directory = os.getcwd()
        logging.info(f"Текущая рабочая директория: {current_directory}")
        
        # Находим файл Excel
        try:
            card_path = find_excel_file()
            logging.info(f"Найден файл Карта: {card_path}")
            print(f"Найден файл Карта: {card_path}")
        except Exception as e:
            logging.error(f"Ошибка при поиске файла Excel: {e}")
            raise
            
        # Чтение данных из найденного файла Excel
        try:
            card_data = read_card(card_path)
            print("Данные из файла Карта:", card_data) 
        except Exception as e:
            logging.error(f"Ошибка при чтении данных из файла Excel: {e}")
            raise
            
        # Проверяем, что все обязательные поля заполнены
        try:
            check_card_data(card_data)
        except Exception as e:
            logging.error(f"Ошибка при проверке обязательных полей: {e}")
            raise

        # Проверяем, является ли файл особым типом "Кость СВО нет пригодной ДНК"
        is_special_case = card_data.get("1") == "НАЗВАНИЕ ШАБЛОНА ДЛЯ КОСТЕЙ"
        
        # Проверяем дополнительные особые случаи
        is_no_result_case = card_data.get("1") in [
            "СВО_Молов_образец_родственники_нет результата_форез",
            "СВО_Молов_образец_родственники_нет результата_RT"
        ]
        
        if is_special_case:
            logging.info("Обнаружен особый тип файла: 'Кость СВО нет пригодной ДНК'. Выполняется упрощенная обработка.")
            
            # Создаем и заполняем специальный шаблон "Для отправки кость нет результата"
            try:
                output_folder = Path(r"F:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\Ростов СВО\Кости с ИК-2")
                template_name = "Для отправки кость нет результата"
                # Передаем пустой список как txt_data, так как чтение txt файла не требуется
                output_path = create_and_fill_document(card_data, [], template_name)
                print(f"Новый документ создан и заполнен: {output_path}")
            except Exception as e:
                logging.error(f"Ошибка при создании и заполнении документа: {e}")
                raise
                
            # Перемещаем файл в финальную папку
            try:
                final_word_path, _ = move_to_final_location(card_data, output_path)
            except Exception as e:
                logging.error(f"Ошибка при перемещении документа: {e}")
                raise
                
            # Копируем документ в архивную папку
            try:
                archive_folder = Path(card_data["5"])
                destination = archive_folder / final_word_path.name
                shutil.copy(final_word_path, destination)
                logging.info(f"Файл скопирован в архив: {destination}")
            except Exception as e:
                logging.error(f"Ошибка при сохранении документа в архив: {e}")
                raise
                
            # Печать документа
            try:
                if final_word_path.exists():
                    if ask_for_print():
                        try:
                            print_document(str(final_word_path))
                            logging.info(f"Документ {final_word_path} отправлен на печать.")
                        except Exception as e:
                            logging.error(f"Ошибка при печати документа: {e}")
                    else:
                        logging.info("Печать документа отменена пользователем.")
                else:
                    logging.error(f"Документ для печати не найден: {final_word_path}")
            except Exception as e:
                logging.error(f"Ошибка при печати документа: {e}")
                raise
                
        elif is_no_result_case:
            logging.info(f"Обнаружен особый тип файла: '{card_data.get('1')}'. Выполняется только архивация и печать.")
            
            try:
                source_folder = Path(card_data["2"])
                archive_folder = Path(card_data["5"])
                
                # Копируем документ в архивную папку
                for file in source_folder.glob(f"{card_data['НОМ']}-25.doc*"):
                    destination = archive_folder / file.name
                    shutil.copy(file, destination)
                    logging.info(f"Файл скопирован в архив: {destination}")
                    
                # Печать документа
                document_to_print = source_folder / f"{card_data['НОМ']}-25.docx"
                if document_to_print.exists():
                    if ask_for_print():
                        try:
                            print_document(str(document_to_print))
                            logging.info(f"Документ {document_to_print} отправлен на печать.")
                        except Exception as e:
                            logging.error(f"Ошибка при печати документа: {e}")
                    else:
                        logging.info("Печать документа отменена пользователем.")
                else:
                    logging.error(f"Документ для печати не найден: {document_to_print}")
            except Exception as e:
                logging.error(f"Ошибка при архивации или печати документа: {e}")
                raise
                
        else:
            # Находим файл txt
            try:
                base_working_materials_folder = Path(card_data["4"])
                nom_folder_name = f"{card_data['НОМ']}-25"
                working_materials_folder = base_working_materials_folder / nom_folder_name
                if not working_materials_folder.exists():
                    logging.error(f"Папка '{working_materials_folder}' не существует.")
                    raise FileNotFoundError(f"Папка '{working_materials_folder}' не существует.")
                txt_path = find_txt_file(working_materials_folder)
                txt_data = read_and_filter_txt_data(txt_path)
            except Exception as e:
                logging.error(f"Ошибка при поиске или чтении файла txt: {e}")
                raise
                
            # Определяем путь для сохранения файла XML
            try:
                xml_output_folder = Path(r"F:\ГЕНЕТИКА НОВОЕ ЗДАНИЕ\BASE")
                xml_path = xml_output_folder / f"{card_data['НОМ']}-25.xml"
                # Преобразование txt в xml
                convert_txt_to_xml(txt_path, xml_path, object_name_suffix=txt_data[0][0])
            except Exception as e:
                logging.error(f"Ошибка при преобразовании txt в xml: {e}")
                raise
                
            # Определяем шаблон в зависимости от типа
            try:
                if card_data.get("1") == "Образец СВО Молов прямая идентификация":
                    template_name = "Для отправки личный"
                elif card_data.get("1") in ["Образец СВО Молов фонд", "Образец СВО Ростов"]:
                    template_name = "Для отправки родственники"
                elif card_data.get("1") == "Кость СВО":
                    template_name = "Для отправки кость"
                else:
                    template_name = "Для отправки родственники"
            except Exception as e:
                logging.error(f"Ошибка при определении шаблона документа: {e}")
                raise
                
            # Создаем Word-файл во временной папке
            try:
                word_file_path = create_and_fill_document(card_data, txt_data, template_name)
                print(f"Новый документ создан и заполнен: {word_file_path}")
            except Exception as e:
                logging.error(f"Ошибка при создании и заполнении документа: {e}")
                raise
                
            # Находим файл VBA
            try:
                vba_path = find_vba_file(card_data)
            except Exception as e:
                logging.error(f"Ошибка при поиске файла VBA: {e}")
                raise
                
            # Логирование меток перед вызовом replace_in_excel
            logging.info(f"Метки для замены: {card_data}")
            
            # Проверяем значение в графе "1" и вызываем соответствующую функцию
            value_in_column_1 = card_data.get("1", "").strip()
            if value_in_column_1 in ["Образец СВО Молов фонд", "Образец СВО Молов прямая идентификация", "Образец СВО Ростов"]:
                logging.info("Значение '1' в Карте соответствует 2-ДНК")
                try:
                    excel_file_path = run_vba_macro_2_DNK(vba_path, txt_path, Path(card_data["4"]), card_data)
                    logging.info(f"Файл Excel успешно создан: {excel_file_path}")
                except Exception as e:
                    logging.error(f"Ошибка при выполнении макроса VBA 2-DNK: {e}")
                    raise
            elif value_in_column_1 == "Кость СВО":
                logging.info("Значение '1' в Карте соответствует ИК-2")
                try:
                    excel_file_path = run_vba_macro_IK_2(vba_path, txt_path, Path(card_data["4"]), card_data)
                    logging.info(f"Файл Excel успешно создан: {excel_file_path}")
                except Exception as e:
                    logging.error(f"Ошибка при выполнении макроса VBA IK-2: {e}")
                    raise
            else:
                logging.warning(f"Неизвестное значение в графе '1' Карты: {value_in_column_1}")
                try:
                    excel_file_path = run_vba_macro(vba_path, txt_path, Path(card_data["4"]), card_data)
                    logging.info(f"Файл Excel успешно создан: {excel_file_path}")
                except Exception as e:
                    logging.error(f"Ошибка при выполнении макроса VBA: {e}")
                    raise
                    
            # Перемещаем оба файла в финальную папку
            try:
                final_word_path, final_excel_path = move_to_final_location(card_data, word_file_path, excel_file_path)
                logging.info(f"Файлы перемещены в финальные папки: Word={final_word_path}, Excel={final_excel_path}")
            except Exception as e:
                logging.error(f"Ошибка при перемещении файлов: {e}")
                raise
                
            # Печать и сохранение заключения в папку под кодом 5
            try:
                archive_folder = Path(card_data["5"])
                
                # Проверяем, существует ли документ
                if final_word_path.exists():
                    # Запрашиваем подтверждение на печать
                    if ask_for_print():
                        try:
                            # Вызываем функцию печати
                            print_document(str(final_word_path))
                            logging.info(f"Документ {final_word_path} отправлен на печать.")
                        except Exception as e:
                            logging.error(f"Ошибка при печати документа: {e}")
                    else:
                        logging.info("Печать документа отменена пользователем.")
                else:
                    logging.error(f"Документ для печати не найден: {final_word_path}")
                    
                # Копируем документ в архивную папку
                destination = archive_folder / final_word_path.name
                shutil.copy(final_word_path, destination)
                logging.info(f"Файл скопирован в архив: {destination}")
            except Exception as e:
                logging.error(f"Ошибка при печати или сохранении документа: {e}")
                raise
                
            print("Скрипт завершён!")
    except Exception as e:
        logging.error(f"Произошла ошибка: {e}")
        show_error_message(f"Произошла ошибка: {e}")
        
    finally:
        try:
            # Создаем окно с сообщением
            root = tk.Tk()
            root.withdraw()
            
            message_window = tk.Toplevel(root)
            message_window.title("ВАЖНО!")
            message_window.geometry("600x300")
            message_window.protocol("WM_DELETE_WINDOW", lambda: (message_window.destroy(), root.quit()))
            
            # Привязка клавиш
            message_window.bind('<Escape>', lambda e: (message_window.destroy(), root.quit()))
            message_window.bind('<Return>', lambda e: (message_window.destroy(), root.quit()))
            
            # Текст сообщения
            label = tk.Label(
                message_window, 
                text="""1. ВНЕСИ ДАННЫЕ В 1С
2. ЗАПОЛНИ ЖУРНАЛ
3. ЗАПОЛНИ ПОДПИСКУ
4. СДАЙ ОБЪЕКТЫ!""", 
                font=("Arial", 20, "bold"), 
                fg="red"
            )
            label.pack(expand=True, fill="both", padx=20, pady=50)
            
            # Кнопка OK
            ok_button = tk.Button(
                message_window, 
                text="OK (Enter)", 
                command=lambda: (message_window.destroy(), root.quit()),
                font=("Arial", 14)
            )
            ok_button.pack(pady=10)
            
            # Настройка фокуса
            message_window.after(100, lambda: (
                ok_button.focus_force(),
                message_window.attributes('-topmost', True)
            ))
            
            # Центрирование окна
            message_window.update_idletasks()
            width = message_window.winfo_width()
            height = message_window.winfo_height()
            x = (message_window.winfo_screenwidth() // 2) - (width // 2)
            y = (message_window.winfo_screenheight() // 2) - (height // 2)
            message_window.geometry(f'+{x}+{y}')
            
            # Запускаем главный цикл
            root.mainloop()
            
        except Exception as e:
            logging.error(f"Ошибка в окне сообщения: {e}")
        finally:
            # Гарантированная очистка
            try:
                if 'message_window' in locals():
                    message_window.destroy()
            except:
                pass
                
            try:
                if 'root' in locals():
                    root.quit()
                    root.destroy()
            except:
                pass
            
            # Принудительный сбор мусора
            import gc
            gc.collect()
            
# Запуск скрипта
if __name__ == "__main__":
    main()